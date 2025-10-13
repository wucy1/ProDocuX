#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ProDocuX Web應用
Flask Web介面
"""

import os
import sys
import json
import logging
from pathlib import Path
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, session
from werkzeug.utils import secure_filename
import uuid
from typing import Optional






def _analyze_learning_differences(original_data: dict, corrected_data: dict) -> list:
    """
    分析學習差異
    
    Args:
        original_data: 原始資料
        corrected_data: 修正後資料
        
    Returns:
        差異列表
    """
    differences = []
    
    def compare_dicts(orig, corr, path=""):
        for key in set(orig.keys()) | set(corr.keys()):
            current_path = f"{path}.{key}" if path else key
            
            if key not in orig:
                differences.append({
                    "type": "added",
                    "path": current_path,
                    "original": None,
                    "corrected": corr[key]
                })
            elif key not in corr:
                differences.append({
                    "type": "removed",
                    "path": current_path,
                    "original": orig[key],
                    "corrected": None
                })
            elif isinstance(orig[key], dict) and isinstance(corr[key], dict):
                compare_dicts(orig[key], corr[key], current_path)
            elif orig[key] != corr[key]:
                differences.append({
                    "type": "modified",
                    "path": current_path,
                    "original": orig[key],
                    "corrected": corr[key]
                })
    
    compare_dicts(original_data, corrected_data)
    return differences
from core.profile_manager import ProfileManager
from core.extractor import DocumentExtractor
from core.transformer import DocumentTransformer
from utils.file_handler import FileHandler
from utils.cost_calculator import CostCalculator
from utils.settings_manager import SettingsManager
from utils.pricing_manager import get_pricing_manager
from utils.workflow_preferences import get_preferences_manager

# 設定日誌
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def get_default_model(provider):
    """獲取預設模型 - 僅用於系統內部，不強制用戶使用"""
    # 這個函數僅用於系統內部處理，不應該強制用戶使用任何預設模型
    # 用戶應該完全自由選擇他們想要的AI提供者和模型
    return None  # 不提供任何預設值

def is_first_run():
    """檢查是否為首次運行"""
    try:
        from utils.desktop_manager import DesktopManager
        desktop_manager = DesktopManager()
        
        # 檢查設定檔案
        config_file = desktop_manager.workspace_dir / "startup_config.json"
        return not config_file.exists()
        
    except Exception:
        return True

def create_app():
    """創建Flask應用"""
    app = Flask(__name__)
    app.config['SECRET_KEY'] = 'prodocux-secret-key'
    app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
    
    # 國際化配置
    app.config['LANGUAGES'] = {
        'zh_TW': '繁體中文',
        'en': 'English'
    }
    app.config['DEFAULT_LOCALE'] = 'zh_TW'
    
    # 初始化組件變數
    settings_manager = None
    file_handler = None
    profile_manager = None
    cost_calculator = None
    
    def get_components():
        """獲取或初始化組件"""
        nonlocal settings_manager, file_handler, profile_manager, cost_calculator
        
        if settings_manager is None:
            # 嘗試從環境變數或啟動設定獲取工作空間路徑
            workspace_path = os.getenv('PRODOCUX_WORKSPACE_PATH')
            if not workspace_path:
                # 嘗試從可能的啟動設定檔案載入
                from utils.desktop_manager import DesktopManager
                temp_dm = DesktopManager()
                app_dir = temp_dm.app_dir
                
                # 檢查多個可能的位置
                possible_configs = [
                    Path.home() / "Documents" / "ProDocuX_Workspace" / "startup_config.json",
                    Path.home() / "文檔" / "ProDocuX_Workspace" / "startup_config.json",
                    Path.cwd() / "startup_config.json"
                ]
                
                # 只在開發環境中檢查應用程式目錄
                if not getattr(sys, 'frozen', False):
                    possible_configs.insert(0, app_dir / "ProDocuX_Workspace" / "startup_config.json")
                
                for config_path in possible_configs:
                    if config_path.exists():
                        try:
                            with open(config_path, 'r', encoding='utf-8') as f:
                                import json
                                config = json.load(f)
                                workspace_path = config.get('workspace_path')
                                if workspace_path:
                                    break
                        except Exception:
                            continue
            
            try:
                # 載入工作空間的.env文件
                if workspace_path:
                    env_file = Path(workspace_path) / ".env"
                    if env_file.exists():
                        from dotenv import load_dotenv
                        load_dotenv(env_file)
                        logger.info(f"Loaded workspace .env file: {env_file}")
                
                settings_manager = SettingsManager(workspace_path=workspace_path)
                file_handler = FileHandler()
                
                # 獲取工作空間的 profiles 目錄
                workspace_profiles_dir = None
                if workspace_path:
                    workspace_profiles_dir = Path(workspace_path) / "profiles"
                else:
                    # 如果沒有工作空間路徑，使用設定管理器的目錄路徑
                    dir_paths = settings_manager.get_directory_paths()
                    workspace_profiles_dir = dir_paths.get('profiles')
                
                profile_manager = ProfileManager(workspace_profiles_dir=workspace_profiles_dir)
                cost_calculator = CostCalculator()
                
                # 確保目錄存在
                settings_manager.ensure_directories()
                
                logger.info("All components initialized successfully")
            except Exception as e:
                logger.error(f"Component initialization failed: {e}", exc_info=True)
                raise
        
        return settings_manager, file_handler, profile_manager, cost_calculator

    def save_template_file(template_file):
        """保存模板檔案"""
        try:
            # 獲取模板目錄
            settings_manager, _, _, _ = get_components()
            template_dir = settings_manager.get_directory_paths()['template']
            template_dir.mkdir(parents=True, exist_ok=True)
            
            # 使用原始文件名
            filename = template_file.filename
            file_path = template_dir / filename
            
            # 如果文件已存在，添加時間戳
            if file_path.exists():
                name_part = file_path.stem
                ext_part = file_path.suffix
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{name_part}_{timestamp}{ext_part}"
                file_path = template_dir / filename
                logger.info(f"File already exists, using new name: {filename}")
            
            # 保存檔案
            template_file.save(str(file_path))
            
            return filename  # 返回文件名而不是完整路徑
        except Exception as e:
            logger.error(f"Failed to save template file: {e}")
            return None
    
    @app.route('/')
    def index():
        """主頁"""
        return render_template('index.html')
    
    @app.route('/upload', methods=['POST'])
    def upload_file():
        """上傳檔案"""
        try:
            logger.info(f"Received file upload request, form data: {request.form}")
            logger.info(f"File list: {list(request.files.keys())}")
            
            if 'file' not in request.files:
                logger.warning("No 'file' field in request")
                return jsonify({'error': '沒有選擇檔案'}), 400
            
            file = request.files['file']
            if file.filename == '':
                logger.warning("File name is empty")
                return jsonify({'error': '沒有選擇檔案'}), 400
            
            logger.info(f"Processing file: {file.filename}")
            
            # 保存檔案
            _, file_handler, _, _ = get_components()
            filename = secure_filename(file.filename)
            file_id = str(uuid.uuid4())
            
            logger.info(f"File ID: {file_id}, secure filename: {filename}")
            
            file_path = file_handler.save_uploaded_file(
                file.read(), f"{file_id}_{filename}"
            )
            
            logger.info(f"File saved to: {file_path}")
            
            return jsonify({
                'success': True,
                'file_id': file_id,
                'filename': filename,
                'file_path': str(file_path)
            })
            
        except Exception as e:
            logger.error(f"File upload failed: {e}", exc_info=True)
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/pages/count', methods=['POST'])
    def count_pages():
        """快速獲取檔案頁面數量"""
        try:
            data = request.get_json()
            file_id = data.get('file_id')
            
            if not file_id:
                return jsonify({'error': '缺少檔案ID'}), 400
            
            # 查找檔案
            _, file_handler, _, _ = get_components()
            upload_dir = file_handler.upload_dir
            
            file_path = None
            for file_name in os.listdir(upload_dir):
                if file_name.startswith(file_id):
                    file_path = Path(upload_dir) / file_name
                    break
            
            if not file_path or not file_path.exists():
                return jsonify({'error': '檔案不存在'}), 404
            
            # 快速獲取頁面數量
            extractor = DocumentExtractor()
            
            # 根據檔案類型選擇讀取方法
            if file_path.suffix.lower() == '.pdf':
                pages_data = extractor._read_pdf_by_pages(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                pages_data = extractor._read_docx_by_pages(file_path)
            else:
                return jsonify({'error': '不支援的檔案格式'}), 400
            
            return jsonify({
                'success': True,
                'file_type': file_path.suffix.lower().lstrip('.'),
                'total_pages': len(pages_data)
            })
            
        except Exception as e:
            logger.error(f"Failed to get page count: {e}")
            return jsonify({'error': str(e)}), 500

    @app.route('/api/pages/preview', methods=['POST'])
    def preview_pages():
        """預覽檔案頁面，讓用戶選擇要處理的頁面"""
        try:
            data = request.get_json()
            file_id = data.get('file_id')
            
            if not file_id:
                return jsonify({'error': '缺少檔案ID'}), 400
            
            # 查找檔案
            _, file_handler, _, _ = get_components()
            upload_dir = file_handler.upload_dir
            
            file_path = None
            for file_name in os.listdir(upload_dir):
                if file_name.startswith(file_id):
                    file_path = Path(upload_dir) / file_name
                    break
            
            if not file_path or not file_path.exists():
                return jsonify({'error': '檔案不存在'}), 404
            
            # 讀取頁面資訊（不再依賴 DocumentExtractor，避免要求 Profile/提示詞）
            pages_data = []
            suffix = file_path.suffix.lower()
            if suffix == '.pdf':
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page_num, page in enumerate(pdf.pages, 1):
                            page_text = page.extract_text() or ''
                            tables = page.extract_tables() or []
                            images = getattr(page, 'images', []) or []
                            pages_data.append({
                                'page_number': page_num,
                                'content': page_text,
                                'tables': tables,
                                'images': images
                            })
                except Exception as e:
                    logger.error(f"PDF page reading failed: {e}")
                    return jsonify({'error': f'PDF 頁面讀取失敗: {str(e)}'}), 500
            elif suffix in ['.docx', '.doc']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    # 模擬分頁：每 10 段落視為一頁，與後端 extractor 的邏輯一致
                    paragraphs_per_page = 10
                    current_page = 1
                    current_content = []
                    for i, paragraph in enumerate(doc.paragraphs):
                        text = paragraph.text.strip()
                        if text:
                            current_content.append(text)
                        if (i + 1) % paragraphs_per_page == 0:
                            page_text = '\n'.join(current_content)
                            pages_data.append({
                                'page_number': current_page,
                                'content': page_text,
                                'tables': [],
                                'images': []
                            })
                            current_page += 1
                            current_content = []
                    # 收尾
                    if current_content:
                        page_text = '\n'.join(current_content)
                        pages_data.append({
                            'page_number': current_page,
                            'content': page_text,
                            'tables': [],
                            'images': []
                        })
                    if not pages_data:
                        pages_data = [{'page_number': 1, 'content': '', 'tables': [], 'images': []}]
                except Exception as e:
                    logger.error(f"DOCX page reading failed: {e}")
                    return jsonify({'error': f'DOCX 頁面讀取失敗: {str(e)}'}), 500
            else:
                return jsonify({'error': '不支援的檔案格式'}), 400
            
            # 準備頁面預覽資料
            page_previews = []
            for page_data in pages_data:
                content = page_data.get('content', '')
                preview = {
                    'page_number': page_data.get('page_number'),
                    'content_preview': content[:200] + '...' if len(content) > 200 else content,
                    'content_length': len(content),
                    'has_tables': bool(page_data.get('tables')),
                    'table_count': len(page_data.get('tables', [])),
                    'keywords_found': []
                }
                
                # 檢查關鍵詞
                content_lower = content.lower()
                keywords = ['產品名稱', '成分', '製造業者', '使用方法', '安全', 'cas', 'inci']
                for keyword in keywords:
                    if keyword in content_lower:
                        preview['keywords_found'].append(keyword)
                
                page_previews.append(preview)
            
            return jsonify({
                'success': True,
                'total_pages': len(pages_data),
                'pages': page_previews
            })
            
        except Exception as e:
            logger.error(f"Page preview failed: {e}", exc_info=True)
            return jsonify({'error': f'頁面預覽失敗: {str(e)}'}), 500

    @app.route('/api/pages/thumbnail', methods=['POST'])
    def generate_page_thumbnail():
        """生成PDF頁面縮圖"""
        try:
            data = request.get_json()
            file_id = data.get('file_id')
            page_number = data.get('page_number', 1)
            
            if not file_id:
                return jsonify({'error': '缺少檔案ID'}), 400
            
            # 查找檔案
            _, file_handler, _, _ = get_components()
            upload_dir = file_handler.upload_dir
            
            file_path = None
            for file_name in os.listdir(upload_dir):
                if file_name.startswith(file_id):
                    file_path = Path(upload_dir) / file_name
                    break
            
            if not file_path or not file_path.exists():
                return jsonify({'error': '檔案不存在'}), 404
            
            # 生成縮圖
            thumbnail_path = generate_pdf_thumbnail(file_path, page_number, file_handler.cache_dir)
            
            if thumbnail_path and thumbnail_path.exists():
                # 將縮圖轉換為base64
                import base64
                with open(thumbnail_path, 'rb') as f:
                    thumbnail_data = base64.b64encode(f.read()).decode('utf-8')
                
                return jsonify({
                    'success': True,
                    'thumbnail': f"data:image/png;base64,{thumbnail_data}"
                })
            else:
                return jsonify({'error': '縮圖生成失敗'}), 500
                
        except Exception as e:
            logger.error(f"Thumbnail generation failed: {e}")
            return jsonify({'error': str(e)}), 500

    @app.route('/process', methods=['POST'])
    def process_document():
        """處理文檔（支援單檔和批量處理）"""
        try:
            data = request.get_json()
            file_ids = data.get('file_ids', [])
            file_id = data.get('file_id')  # 單檔相容
            profile_name = data.get('profile', 'default')
            output_format = data.get('format', 'json')
            selected_template = data.get('template', '')
            work_id = data.get('work_id')
            work_data = data.get('work_data', {})
            user_prompt = data.get('user_prompt', '')
            
            # 嚴格要求使用者提供提示詞
            if not user_prompt or user_prompt.strip() == '':
                return jsonify({'error': '必須提供使用者提示詞。系統不會使用預設提示詞。'}), 400
            selected_pages = data.get('selected_pages', [])  # 用戶選擇的頁面（全域）
            selected_pages_map = data.get('selected_pages_map', {})  # 每個檔案的頁面選擇
            
            # AI設定（優先使用工作流程指定，否則使用預設）
            ai_provider = data.get('ai_provider', '')
            ai_model = data.get('ai_model', '')
            
            # 如果沒有指定AI設定，要求用戶必須選擇
            if not ai_provider or not ai_model:
                return jsonify({'error': '必須指定AI提供者和模型。請在處理前選擇您要使用的AI服務。'}), 400
            
            logger.info(f"Processing document with AI settings: {ai_provider}/{ai_model}")
            
            # 統一處理：單檔轉為列表
            if file_id and not file_ids:
                file_ids = [file_id]
            
            if not file_ids:
                return jsonify({'error': '缺少檔案ID'}), 400
            
            # 批量處理檔案
            _, file_handler, _, _ = get_components()
            results = []
            errors = []
            
            # 初始化提取器（所有檔案共用）
            try:
                resolved_profile_path = None
                inline_profile = None
                work_profile = None

                # 1) 優先使用工作流程內附的 Profile 和 Prompt
                if work_id and work_data:
                    work_profile = (work_data or {}).get('profile')
                    work_prompt = (work_data or {}).get('prompt')
                    
                    # 如果工作流程提供了 prompt，優先使用工作流程的 prompt
                    if work_prompt and work_prompt.strip():
                        user_prompt = work_prompt
                        logger.info(f"Using workflow-provided prompt, length: {len(user_prompt)} characters")
                    
                    if work_profile is not None:
                        logger.info("Received Profile from workflow, will strictly use this Profile without fallback to default")
                        pm = ProfileManager()

                        # dict 直接保存為臨時yml
                        if isinstance(work_profile, dict):
                            inline_profile = work_profile
                            tmp_path = pm.work_profiles_dir / f"work_{work_id}_inline.yml"
                            pm.save_profile(inline_profile, tmp_path)
                            resolved_profile_path = str(tmp_path)
                        elif isinstance(work_profile, str):
                            possible_path = Path(work_profile)
                            if possible_path.exists():
                                resolved_profile_path = str(possible_path)
                            else:
                                # 嘗試解析文字內容（YAML/JSON）
                                try:
                                    import yaml, json as _json
                                    try:
                                        inline_profile = yaml.safe_load(work_profile)
                                    except Exception:
                                        inline_profile = _json.loads(work_profile)
                                    if isinstance(inline_profile, dict):
                                        tmp_path = pm.work_profiles_dir / f"work_{work_id}_inline.yml"
                                        pm.save_profile(inline_profile, tmp_path)
                                        resolved_profile_path = str(tmp_path)
                                except Exception:
                                    pass

                        if not resolved_profile_path:
                            raise ValueError("工作流程的Profile無法識別或保存。請檢查格式。")
                    else:
                        # 工作流程必須提供Profile，不允許沒有Profile的情況
                        raise ValueError("工作流程必須提供Profile。系統不會使用分層載入邏輯。")

                # 2) 舊版傳統Profile（僅在沒有work_data時使用，且不允許default回退）
                if not 'extractor' in locals():
                    if resolved_profile_path:
                        logger.info(f"Using resolved Profile path: {resolved_profile_path}")
                        extractor = DocumentExtractor(resolved_profile_path, ai_provider=ai_provider, ai_model=ai_model, user_prompt=user_prompt)
                    else:
                        if not profile_name or profile_name.strip() == '':
                            raise ValueError("缺少Profile。請在工作流程提供Profile或在請求中指定profile名稱。")
                        if profile_name.endswith('.yml'):
                            profile_path = f"profiles/{profile_name}"
                        else:
                            profile_path = f"profiles/{profile_name}.yml"
                        logger.info(f"Using traditional Profile path: {profile_path}")
                        extractor = DocumentExtractor(profile_path, ai_provider=ai_provider, ai_model=ai_model, user_prompt=user_prompt)
                logger.info("DocumentExtractor initialized successfully")
            except Exception as e:
                logger.error(f"DocumentExtractor initialization failed: {e}", exc_info=True)
                return jsonify({'error': f'提取器初始化失敗: {str(e)}'}), 500
            
            # 處理每個檔案
            total_files = len(file_ids)
            results = []
            errors = []
            
            for i, current_file_id in enumerate(file_ids):
                try:
                    # 進度日誌
                    progress = f"({i+1}/{total_files})"
                    logger.info(f"{progress} Starting to process file: {current_file_id}")
                    
                    # 查找檔案
                    file_path = None
                    logger.info(f"Looking for file ID: {current_file_id}")
                    logger.info(f"Upload directory: {file_handler.upload_dir}")
                    
                    for file_info in file_handler.list_files(file_handler.upload_dir):
                        logger.info(f"Checking file: {file_info['name']} -> {file_info['path']}")
                        # 修復：檢查檔案名是否以file_id開頭
                        if file_info['name'].startswith(current_file_id + '_'):
                            file_path = Path(file_info['path'])
                            logger.info(f"Found matching file: {file_path}")
                            break
                    
                    if not file_path:
                        logger.error(f"File ID not found: {current_file_id}")
                        errors.append(f"檔案ID {current_file_id} 不存在")
                        continue
                    
                    # 檢查檔案是否存在
                    if not file_path.exists():
                        logger.error(f"File does not exist: {file_path}")
                        readable_filename = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                        errors.append(f"檔案「{readable_filename}」不存在於 {file_path}")
                        continue
                    
                    # 提取文檔（支援智能提示詞）
                    logger.info(f"{progress} Starting document extraction: {file_path}")
                    logger.info(f"{progress} Using prompt: {user_prompt}")
                    per_file_pages = selected_pages_map.get(current_file_id) if isinstance(selected_pages_map, dict) else None
                    logger.info(f"{progress} Selected pages: {per_file_pages or selected_pages}")
                    try:
                        structured_data = extractor.extract(file_path, "dict", user_prompt, per_file_pages or selected_pages)
                        
                        # 檢查是否為檔案輸出
                        if isinstance(structured_data, dict) and structured_data.get('_file_output', False):
                            logger.info(f"{progress} AI directly output file, skipping conversion step")
                            file_info = structured_data['_file_info']
                            
                            # 保存檔案
                            if 'file_data' in file_info:
                                file_extension = file_info['file_type']
                                if file_extension == 'unknown':
                                    file_extension = 'bin'
                                
                                original_name = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                                original_stem = Path(original_name).stem
                                output_filename = f"{original_stem}_ai_output.{file_extension}"
                                output_path = file_handler.output_dir / output_filename
                                
                                with open(output_path, 'wb') as f:
                                    f.write(file_info['file_data'])
                                
                                logger.info(f"{progress} File saved: {output_filename}")
                                results.append({
                                    'success': True,
                                    'filename': filename,
                                    'output_file': output_filename,
                                    'file_type': file_info['file_type'],
                                    'file_size': file_info.get('file_size', 0),
                                    'extraction_method': 'file_output'
                                })
                                continue
                            elif 'file_path' in file_info:
                                logger.info(f"{progress} File download link: {file_info['file_path']}")
                                results.append({
                                    'success': True,
                                    'filename': filename,
                                    'download_link': file_info['file_path'],
                                    'file_type': file_info['file_type'],
                                    'extraction_method': 'download_link'
                                })
                                continue
                        
                        logger.info(f"{progress} Document extraction successful, data size: {len(str(structured_data))} characters")
                    except Exception as e:
                        logger.error(f"{progress} Document extraction failed: {e}", exc_info=True)
                        error_msg = str(e)
                        # 從檔案路徑提取可讀的檔案名稱
                        readable_filename = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                        if "安全過濾器阻止" in error_msg:
                            errors.append(f"檔案「{readable_filename}」: AI模型安全過濾器阻止了內容處理。請手動選擇其他AI模型或調整提示詞內容")
                        else:
                            errors.append(f"檔案「{readable_filename}」: 文檔提取失敗 - {error_msg}")
                        continue
                    
                    # 轉換文檔（工作空間模板）
                    logger.info(f"{progress} Starting document conversion, output format: {output_format}")
                    try:
                        transformer = DocumentTransformer(resolved_profile_path)
                        # 輸出以原檔名命名
                        original_name = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                        original_stem = Path(original_name).stem
                        output_filename = f"{original_stem}_結果.{output_format}"
                        output_path = file_handler.output_dir / output_filename
                        logger.info(f"{progress} Output file path: {output_path}")
                    except Exception as e:
                        logger.error(f"{progress} Transformer initialization failed: {e}", exc_info=True)
                        readable_filename = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                        errors.append(f"檔案「{readable_filename}」: 轉換器初始化失敗 - {str(e)}")
                        continue
                    
                    if output_format == "json":
                        success = file_handler.save_json_data(structured_data, output_filename)
                    else:
                        # 模板尋址：嚴格使用工作流程指定的模板
                        settings_manager, _, _, _ = get_components()
                        template_dir = settings_manager.get_directory_paths()['template']
                        
                        # 工作流程必須提供模板，不允許回退
                        if not selected_template:
                            readable_filename = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                            errors.append(f"檔案「{readable_filename}」: 工作流程必須指定模板。系統不會使用自動猜測。")
                            continue
                        
                        # 直接使用工作流程指定的模板檔案，不做任何智能比對
                        candidate = template_dir / selected_template
                        
                        if not candidate.exists():
                            readable_filename = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                            errors.append(f"檔案「{readable_filename}」: 找不到指定的模板檔案: {selected_template}")
                            continue
                        
                        # 嘗試 docxtpl 渲染；允許回退到關鍵詞替換模式
                        logger.info(f"{progress} Starting template conversion: {candidate} -> {output_path}")
                        try:
                            success = transformer.transform(structured_data, candidate, output_path, output_format, allow_fallback=True)
                            logger.info(f"{progress} Template conversion result: {success}")
                        except Exception as e:
                            logger.error(f"{progress} Template conversion failed: {e}", exc_info=True)
                            readable_filename = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                            errors.append(f"檔案「{readable_filename}」: 模板轉換失敗 - {str(e)}")
                            continue
                    
                    if success:
                        logger.info(f"{progress} File processing completed: {current_file_id}")
                        results.append({
                            'file_id': current_file_id,
                            'data': structured_data,
                            'download_url': f'/download/{output_filename}',
                            'status': 'success'
                        })
                    else:
                        logger.error(f"{progress} File processing failed: {current_file_id}")
                        readable_filename = file_path.name.split('_', 1)[1] if '_' in file_path.name else file_path.name
                        errors.append(f"檔案「{readable_filename}」: 文檔處理失敗")
                        
                except Exception as e:
                    logger.error(f"{progress} Processing file {current_file_id} failed: {e}")
                    # 嘗試從檔案ID獲取檔案名稱，如果失敗則使用檔案ID
                    try:
                        for file_info in file_handler.list_files(file_handler.upload_dir):
                            if file_info['name'].startswith(current_file_id + '_'):
                                readable_filename = file_info['name'].split('_', 1)[1] if '_' in file_info['name'] else file_info['name']
                                errors.append(f"檔案「{readable_filename}」: {str(e)}")
                                break
                        else:
                            errors.append(f"檔案ID {current_file_id}: {str(e)}")
                    except:
                        errors.append(f"檔案ID {current_file_id}: {str(e)}")
            
            # 更新工作流程調用次數（如果有成功處理的檔案）
            if results and work_id:
                try:
                    settings_manager, _, _, _ = get_components()
                    # 獲取當前工作流程
                    work = settings_manager.get_work(work_id)
                    if work:
                        # 更新調用次數
                        work['processed_count'] = work.get('processed_count', 0) + len(results)
                        work['last_used'] = datetime.now().isoformat()
                        # 保存更新後的工作流程
                        settings_manager.save_work(work)
                        logger.info(f"Workflow {work_id} call count updated: {work['processed_count']}")
                except Exception as e:
                    logger.warning(f"Failed to update workflow call count: {e}")
            
            # 返回批量處理結果
            logger.info(f"Batch processing completed: Total files={len(file_ids)}, Success={len(results)}, Failed={len(errors)}")
            if results:
                return jsonify({
                    'success': True,
                    'batch_processing': True,
                    'total_files': len(file_ids),
                    'successful_files': len(results),
                    'failed_files': len(errors),
                    'results': results,
                    'errors': errors
                })
            else:
                return jsonify({
                    'success': False,
                    'error': '所有檔案處理失敗',
                    'errors': errors
                }), 500
                
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/download/<filename>')
    def download_file(filename):
        """下載檔案"""
        try:
            _, file_handler, _, _ = get_components()
            file_path = file_handler.output_dir / filename
            if file_path.exists():
                # 以原樣檔名回傳（瀏覽器顯示此檔名）
                return send_file(str(file_path), as_attachment=True, download_name=filename)
            else:
                return jsonify({'error': '檔案不存在'}), 404
        except Exception as e:
            logger.error(f"File download failed: {e}")
            return jsonify({'error': str(e)}), 500

    @app.route('/api/input/list', methods=['GET'])
    def api_list_input():
        """列出工作空間 input 目錄的檔案"""
        try:
            settings_manager, file_handler, _, _ = get_components()
            input_dir = settings_manager.get_directory_paths()['input']
            files = file_handler.list_files(input_dir, pattern='*')
            # 僅保留支援的副檔名
            files = [f for f in files if f.get('extension', '').lower() in ['.pdf', '.doc', '.docx']]
            return jsonify({'success': True, 'files': files})
        except Exception as e:
            logger.error(f"Failed to list input files: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/input/import', methods=['POST'])
    def api_import_input():
        """將 input 目錄檔案匯入為上傳（複製到 uploads）"""
        try:
            data = request.get_json(silent=True) or {}
            selected = set(data.get('filenames') or [])

            settings_manager, file_handler, _, _ = get_components()
            input_dir = Path(settings_manager.get_directory_paths()['input'])

            imported = []
            for p in input_dir.iterdir():
                if not p.is_file():
                    continue
                if p.suffix.lower() not in ['.pdf', '.doc', '.docx']:
                    continue
                if selected and p.name not in selected:
                    continue

                file_id = str(uuid.uuid4())
                safe_name = secure_filename(p.name)
                target = file_handler.upload_dir / f"{file_id}_{safe_name}"
                shutil.copy2(str(p), str(target))
                imported.append({'file_id': file_id, 'filename': safe_name, 'file_path': str(target), 'size': target.stat().st_size})

            return jsonify({'success': True, 'files': imported})
        except Exception as e:
            logger.error(f"Input file import failed: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500
    
    @app.route('/profiles')
    def list_profiles():
        """列出可用的Profile"""
        try:
            _, _, profile_manager, _ = get_components()
            profiles = profile_manager.list_profiles()
            return jsonify({'success': True, 'profiles': profiles})
        except Exception as e:
            logger.error(f"Failed to get Profile list: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500
    
    @app.route('/templates')
    def list_templates():
        """列出可用的模板"""
        try:
            transformer = DocumentTransformer()
            templates = transformer.get_available_templates()
            return jsonify({'templates': templates})
        except Exception as e:
            logger.error(f"Failed to get template list: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/learn', methods=['POST'])
    def learn_from_corrections():
        """從修正中學習（支援分層Profile）"""
        try:
            data = request.get_json()
            original_data = data.get('original_data')
            corrected_data = data.get('corrected_data')
            source_file = data.get('source_file')
            profile_name = data.get('profile', 'default')
            work_id = data.get('work_id')
            work_data = data.get('work_data', {})
            
            if not all([original_data, corrected_data, source_file]):
                return jsonify({'error': '缺少必要參數'}), 400
            
            # 分析學習差異
            differences = _analyze_learning_differences(original_data, corrected_data)
            learning_data = {
                'differences': differences,
                'original_data': original_data,
                'corrected_data': corrected_data,
                'source_file': source_file
            }
            
            success = False
            if work_id and work_data:
                # 新版：分層Profile學習
                _, _, profile_manager, _ = get_components()
                
                # 1. 更新品牌Profile
                brand = work_data.get('brand', '')
                work_type = work_data.get('type', '')
                if brand and work_type:
                    brand_success = profile_manager.update_brand_profile_from_learning(
                        brand, work_type, learning_data
                    )
                    if brand_success:
                        logger.info(f"Brand Profile updated: {brand}_{work_type}")
                
                # 2. 如果學習次數多，創建工作專屬Profile
                work_learning_count = work_data.get('learning_count', 0)
                if work_learning_count >= 5:  # 閾值：5次學習後創建工作專屬Profile
                    work_success = profile_manager.update_work_profile_from_learning(
                        work_id, work_data, learning_data
                    )
                    if work_success:
                        logger.info(f"Work-specific Profile updated: work_{work_id}")
                
                success = True
            else:
                # 舊版：傳統Profile學習
                learner = ProfileLearner()
                success = learner.learn_from_corrections(
                    original_data, corrected_data, source_file, profile_name
                )
            
            if success:
                return jsonify({'success': True, 'message': '學習完成'})
            else:
                return jsonify({'error': '學習失敗'}), 500
                
        except Exception as e:
            logger.error(f"Learning failed: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/learn/word', methods=['POST'])
    def learn_from_word_document():
        """從Word文檔學習"""
        try:
            # 檢查是否有檔案上傳
            if 'file' not in request.files:
                return jsonify({'error': '沒有上傳檔案'}), 400
            
            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': '沒有選擇檔案'}), 400
            
            # 檢查檔案格式
            if not file.filename.lower().endswith('.docx'):
                return jsonify({'error': '只支援DOCX格式'}), 400
            
            # 獲取其他參數
            original_data = request.form.get('original_data')
            source_file = request.form.get('source_file')
            profile = request.form.get('profile', 'default')
            
            if not all([original_data, source_file]):
                return jsonify({'error': '缺少必要參數'}), 400
            
            # 保存上傳的Word文檔
            # 對於打包版本，使用工作空間的 uploads 目錄
            if getattr(sys, 'frozen', False):
                try:
                    from utils.desktop_manager import DesktopManager
                    dm = DesktopManager()
                    upload_dir = dm.workspace_dir / "uploads"
                except:
                    upload_dir = Path("uploads")
            else:
                upload_dir = Path("uploads")
            upload_dir.mkdir(exist_ok=True)
            
            import time
            filename = f"corrected_{int(time.time())}_{file.filename}"
            file_path = upload_dir / filename
            file.save(str(file_path))
            
            # 解析原始資料
            try:
                original_data_dict = json.loads(original_data)
            except json.JSONDecodeError:
                return jsonify({'error': '原始資料格式錯誤'}), 400
            
            # 執行學習
            learner = ProfileLearner()
            success = learner.learn_from_word_document(
                original_data_dict, str(file_path), source_file, profile
            )
            
            if success:
                return jsonify({'success': True, 'message': '學習完成'})
            else:
                return jsonify({'error': '學習失敗'}), 500
                
        except Exception as e:
            logger.error(f"Word document learning failed: {e}")
            return jsonify({'error': '學習失敗'}), 500
    
    @app.route('/cost/estimate', methods=['POST'])
    def estimate_cost():
        """估算處理成本"""
        try:
            logger.info("Received cost estimation request")
            data = request.get_json()
            file_id = data.get('file_id')
            profile_name = data.get('profile', 'default')
            
            logger.info(f"Cost estimation parameters: file_id={file_id}, profile={profile_name}")
            
            if not file_id:
                logger.warning("Cost estimation request missing file ID")
                return jsonify({'error': '缺少檔案ID'}), 400
            
            # 獲取組件
            try:
                _, file_handler, _, cost_calculator = get_components()
                logger.info("Components retrieved successfully")
            except Exception as e:
                logger.error(f"Failed to retrieve components: {e}")
                return jsonify({'error': f'組件初始化失敗: {str(e)}'}), 500
            
            # 查找檔案
            file_path = None
            logger.info(f"Looking for file in upload directory: {file_handler.upload_dir}")
            for file_info in file_handler.list_files(file_handler.upload_dir):
                logger.info(f"Checking file: {file_info['name']}")
                if file_info['name'].startswith(file_id):
                    file_path = file_info['path']
                    logger.info(f"Found matching file: {file_path}")
                    break
            
            if not file_path:
                logger.error(f"File ID not found: {file_id}")
                return jsonify({'error': '檔案不存在'}), 404
            
            # 估算成本
            from pathlib import Path
            logger.info(f"Starting cost estimation: {file_path}")
            cost_info = cost_calculator.estimate_cost(Path(file_path), profile_name)
            logger.info(f"Cost estimation completed: {cost_info}")
            
            return jsonify(cost_info)
            
        except Exception as e:
            logger.error(f"Cost estimation failed: {e}", exc_info=True)
            return jsonify({'error': str(e)}), 500
    
    @app.route('/setup')
    def setup():
        """初始設定頁面"""
        try:
            # 檢查是否已完成設定
            if not is_first_run():
                return redirect(url_for('index'))
            
            return render_template('setup.html', user_name=os.getenv('USERNAME', 'User'))
        except Exception as e:
            logger.error(f"Settings page loading failed: {e}")
            return render_template('error.html', error=str(e))
    
    @app.route('/settings')
    def settings_page():
        """設定頁面"""
        return render_template('settings.html')
    
    @app.route('/api/i18n/<locale>')
    def get_i18n(locale):
        """獲取國際化翻譯"""
        try:
            # 使用相對於web目錄的路徑
            import os
            web_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(web_dir)
            locale_file = os.path.join(project_root, "locale", f"{locale}.json")
            
            logger.info(f"Attempting to load internationalization file: {locale_file}")
            logger.info(f"File exists: {os.path.exists(locale_file)}")
            
            if os.path.exists(locale_file):
                with open(locale_file, 'r', encoding='utf-8') as f:
                    translations = json.load(f)
                logger.info(f"Successfully loaded {locale} translation file, containing {len(translations)} translation entries")
                return jsonify(translations)
            else:
                logger.error(f"Internationalization file does not exist: {locale_file}")
                # 列出locale目錄內容用於調試
                locale_dir = os.path.join(project_root, "locale")
                if os.path.exists(locale_dir):
                    files = os.listdir(locale_dir)
                    logger.info(f"locale directory contents: {files}")
                return jsonify({'error': f'Language file not found: {locale_file}'}), 404
        except Exception as e:
            logger.error(f"Internationalization file loading failed: {e}")
            import traceback
            logger.error(f"Detailed error: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/language', methods=['GET', 'POST'])
    def language_api():
        """獲取或設置語言"""
        try:
            if request.method == 'GET':
                # 獲取當前語言
                current_lang = session.get('language', app.config['DEFAULT_LOCALE'])
                return jsonify({'language': current_lang})
            else:
                # 設置語言
                data = request.get_json()
                language = data.get('language')
                if language in app.config['LANGUAGES']:
                    session['language'] = language
                    return jsonify({'success': True, 'message': 'Language updated'})
                else:
                    return jsonify({'error': 'Unsupported language'}), 400
        except Exception as e:
            logger.error(f"Language API failed: {e}")
            return jsonify({'error': str(e)}), 500

    @app.route('/api/settings', methods=['GET'])
    def get_settings():
        """獲取設定"""
        try:
            settings_manager, _, _, _ = get_components()
            settings = settings_manager.get_all_settings()
            return jsonify({'success': True, 'settings': settings})
        except Exception as e:
            logger.error(f"Failed to get settings: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/settings', methods=['POST'])
    def update_settings():
        """更新設定"""
        try:
            settings_manager, _, _, _ = get_components()
            data = request.get_json()
            success = settings_manager.update_settings(data)
            
            if success:
                return jsonify({'success': True, 'message': '設定已更新'})
            else:
                return jsonify({'error': '設定更新失敗'}), 500
                
        except Exception as e:
            logger.error(f"Settings update failed: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/safety/check', methods=['POST'])
    def check_content_safety():
        """檢查內容安全性"""
        try:
            data = request.get_json()
            content = data.get('content', '')
            context_type = data.get('context', 'cosmetic_context')
            check_models = data.get('check_models', [])  # 要檢查的模型列表
            
            if not content:
                return jsonify({'error': '內容不能為空'}), 400
            
            # 本機預檢查
            local_check_result = perform_local_safety_check(content, context_type)
            
            # AI 模型檢查
            ai_check_results = {}
            if check_models:
                ai_check_results = perform_ai_safety_check(content, context_type, check_models)
            
            # 合併結果
            overall_result = {
                'is_safe': local_check_result['is_safe'] and all(
                    result.get('is_safe', False) for result in ai_check_results.values()
                ),
                'local_check': local_check_result,
                'ai_checks': ai_check_results,
                'passed_models': [
                    model for model, result in ai_check_results.items() 
                    if result.get('is_safe', False)
                ],
                'failed_models': [
                    model for model, result in ai_check_results.items() 
                    if not result.get('is_safe', True)
                ],
                'suggestions': local_check_result.get('suggestions', []) + 
                              [suggestion for result in ai_check_results.values() 
                               for suggestion in result.get('suggestions', [])]
            }
            
            return jsonify({
                'success': True,
                'result': overall_result
            })
            
        except Exception as e:
            logger.error(f"安全檢查失敗: {e}")
            return jsonify({'error': str(e)}), 500

    def perform_local_safety_check(content, context_type):
        """執行本機安全檢查"""
        try:
            from utils.safety_precheck import GeminiSafetyPrechecker
            prechecker = GeminiSafetyPrechecker()
            result = prechecker.check_content_safety(content, context_type)
            
            return {
                'is_safe': result.is_safe,
                'risk_level': result.risk_level.value,
                'risk_factors': result.risk_factors,
                'suggestions': result.suggestions,
                'confidence': result.confidence
            }
        except Exception as e:
            logger.error(f"本機安全檢查失敗: {e}")
            return {
                'is_safe': True,  # 預設安全
                'risk_level': 'low',
                'risk_factors': [],
                'suggestions': [],
                'confidence': 0.0,
                'error': str(e)
            }

    def perform_ai_safety_check(content, context_type, check_models):
        """執行 AI 模型安全檢查"""
        results = {}
        
        try:
            settings_manager, _, _, _ = get_components()
            settings = settings_manager.get_all_settings()
            
            for model_info in check_models:
                provider = model_info.get('provider')
                model = model_info.get('model')
                
                try:
                    # 創建對應的 AI 客戶端
                    if provider == 'openai':
                        from utils.multi_ai_client import OpenAIClient
                        api_key = settings.get('openai_api_key')
                        if api_key:
                            client = OpenAIClient(api_key, model)
                            result = client.check_content_safety(content)
                        else:
                            result = {'is_safe': True, 'error': 'API 金鑰未設定'}
                    elif provider == 'gemini':
                        from utils.multi_ai_client import GeminiClient
                        api_key = settings.get('gemini_api_key')
                        if api_key:
                            client = GeminiClient(api_key, model)
                            result = client.check_content_safety(content)
                        else:
                            result = {'is_safe': True, 'error': 'API 金鑰未設定'}
                    elif provider == 'claude':
                        from utils.multi_ai_client import ClaudeClient
                        api_key = settings.get('claude_api_key')
                        if api_key:
                            client = ClaudeClient(api_key, model)
                            result = client.check_content_safety(content)
                        else:
                            result = {'is_safe': True, 'error': 'API 金鑰未設定'}
                    elif provider == 'grok':
                        from utils.multi_ai_client import GrokClient
                        api_key = settings.get('grok_api_key')
                        if api_key:
                            client = GrokClient(api_key, model)
                            result = client.check_content_safety(content)
                        else:
                            result = {'is_safe': True, 'error': 'API 金鑰未設定'}
                    elif provider == 'copilot':
                        from utils.multi_ai_client import CopilotClient
                        api_key = settings.get('copilot_api_key')
                        if api_key:
                            client = CopilotClient(api_key, model)
                            result = client.check_content_safety(content)
                        else:
                            result = {'is_safe': True, 'error': 'API 金鑰未設定'}
                    else:
                        result = {'is_safe': True, 'error': f'不支援的提供者: {provider}'}
                    
                    results[f"{provider}-{model}"] = result
                    
                except Exception as e:
                    logger.error(f"AI 安全檢查失敗 ({provider}-{model}): {e}")
                    results[f"{provider}-{model}"] = {
                        'is_safe': True,
                        'error': str(e)
                    }
                    
        except Exception as e:
            logger.error(f"AI 安全檢查初始化失敗: {e}")
            
        return results

    @app.route('/api/settings/reset', methods=['POST'])
    def reset_settings():
        """重置設定"""
        try:
            settings_manager, _, _, _ = get_components()
            success = settings_manager.reset_settings()
            
            if success:
                return jsonify({'success': True, 'message': '設定已重置'})
            else:
                return jsonify({'error': '設定重置失敗'}), 500
                
        except Exception as e:
            logger.error(f"設定重置失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/settings/test', methods=['POST'])
    def test_settings():
        """測試設定"""
        try:
            settings_manager, _, _, _ = get_components()
            validation = settings_manager.validate_settings()
            
            if validation['valid']:
                return jsonify({'success': True, 'message': '設定驗證通過'})
            else:
                return jsonify({'success': False, 'errors': validation['errors']})
                
        except Exception as e:
            logger.error(f"設定測試失敗: {e}")
            return jsonify({'error': str(e)}), 500

    
    def create_env_file_for_setup(workspace_dir, provider, api_key):
        """為setup創建.env文件並設置API金鑰"""
        try:
            # 複製env_example.txt到工作空間
            env_example_file = Path(__file__).parent.parent / "env_example.txt"
            env_file = workspace_dir / ".env"
            
            if env_example_file.exists():
                # 讀取範例檔案
                with open(env_example_file, 'r', encoding='utf-8') as f:
                    env_content = f.read()
                
                # 替換API金鑰
                if provider == 'openai':
                    env_content = env_content.replace('your_openai_api_key_here', api_key)
                    env_content = env_content.replace('your_iopenai_api_key_here', api_key)
                elif provider == 'claude':
                    env_content = env_content.replace('your_claude_api_key_here', api_key)
                elif provider == 'gemini':
                    env_content = env_content.replace('your_gemini_api_key_here', api_key)
                elif provider == 'grok':
                    env_content = env_content.replace('your_grok_api_key_here', api_key)
                elif provider == 'microsoft':
                    env_content = env_content.replace('your_microsoft_api_key_here', api_key)
                
                # 寫入.env檔案
                with open(env_file, 'w', encoding='utf-8') as f:
                    f.write(env_content)
                
                logger.info(f"已創建.env文件並設置{provider.upper()} API金鑰")
            else:
                logger.warning("找不到env_example.txt，創建基本.env文件")
                # 創建基本.env文件
                with open(env_file, 'w', encoding='utf-8') as f:
                    f.write(f"""# ProDocuX 環境變數設定
# 此檔案由ProDocuX Web設定自動生成

# OpenAI API設定
OPENAI_API_KEY={'sk-your-key-here' if provider != 'openai' else api_key}
IOPENAI_API_KEY={'sk-your-key-here' if provider != 'openai' else api_key}

# Claude API設定
CLAUDE_API_KEY={'sk-ant-your-key-here' if provider != 'claude' else api_key}

# Gemini API設定
GEMINI_API_KEY={'AI-your-key-here' if provider != 'gemini' else api_key}

# Grok API設定
GROK_API_KEY={'grok-your-key-here' if provider != 'grok' else api_key}

# Microsoft API設定
MICROSOFT_API_KEY={'ms-your-key-here' if provider != 'microsoft' else api_key}

# 模型設定
OPENAI_MODEL=gpt-4
IOPENAI_MODEL=gpt-4
CLAUDE_MODEL=claude-3-sonnet-20240229
GEMINI_MODEL=gemini-pro
GROK_MODEL=grok-beta

# 其他設定
MAX_CHUNK_SIZE=8000
CONFIDENCE_THRESHOLD=0.7
""")
            
            return True
            
        except Exception as e:
            logger.error(f"創建.env文件失敗: {e}")
            return False

    @app.route('/api/setup', methods=['POST'])
    def complete_setup():
        """完成初始設定"""
        try:
            settings_manager, _, _, _ = get_components()
            data = request.get_json()
            
            # 保存設定
            settings = {
                'workspace_path': data.get('workspace_path'),
                'ai_provider': data.get('ai_provider'),
                'ai_model': get_default_model(data.get('ai_provider')),
                'openai_api_key': data.get('api_key') if data.get('ai_provider') == 'openai' else '',
                'claude_api_key': data.get('api_key') if data.get('ai_provider') == 'claude' else '',
                'gemini_api_key': data.get('api_key') if data.get('ai_provider') == 'gemini' else '',
                'grok_api_key': data.get('api_key') if data.get('ai_provider') == 'grok' else '',
                'microsoft_api_key': data.get('api_key') if data.get('ai_provider') == 'microsoft' else '',
            }
            
            # 更新設定
            success = settings_manager.update_settings(settings)
            
            if success:
                # 創建工作空間和桌面快捷方式
                from utils.desktop_manager import DesktopManager
                
                # 準備快捷方式選項
                shortcuts = data.get('shortcuts', {})
                selected_shortcuts = []
                if shortcuts.get('workspace'):
                    selected_shortcuts.append('workspace')
                if shortcuts.get('input'):
                    selected_shortcuts.append('input')
                if shortcuts.get('output'):
                    selected_shortcuts.append('output')
                if shortcuts.get('template'):
                    selected_shortcuts.append('template')
                
                # 使用自定義工作空間路徑創建桌面管理器
                try:
                    workspace_path = data.get('workspace_path')
                    desktop_manager = DesktopManager(workspace_path=workspace_path)
                    desktop_manager.setup_workspace(selected_shortcuts=selected_shortcuts)
                    
                    # 創建.env文件並設置API金鑰
                    create_env_file_for_setup(desktop_manager.workspace_dir, data.get('ai_provider'), data.get('api_key'))
                    
                    # 創建啟動設定檔案
                    startup_config = {
                        'workspace_path': workspace_path,
                        'selected_shortcuts': selected_shortcuts,
                        'ai_provider': data.get('ai_provider'),
                        'api_key': data.get('api_key')
                    }
                    
                    startup_config_file = desktop_manager.workspace_dir / "startup_config.json"
                    with open(startup_config_file, 'w', encoding='utf-8') as f:
                        import json
                        json.dump(startup_config, f, ensure_ascii=False, indent=2)
                        
                except Exception as e:
                    logger.error(f"工作空間設置失敗: {e}")
                    return jsonify({'error': f'工作空間設置失敗: {str(e)}'}), 500
                
                return jsonify({'success': True, 'message': '設定完成'})
            else:
                return jsonify({'error': '設定保存失敗'}), 500
                
        except Exception as e:
            logger.error(f"設定完成失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/open-folder/<folder_type>')
    def open_folder(folder_type):
        """開啟資料夾"""
        try:
            import platform
            import subprocess
            
            # 獲取資料夾路徑
            settings_manager, _, _, _ = get_components()
            dir_paths = settings_manager.get_directory_paths()
            
            if folder_type not in dir_paths:
                return jsonify({'error': '無效的資料夾類型'}), 400
            
            folder_path = dir_paths[folder_type]
            
            if not folder_path.exists():
                return jsonify({'error': '資料夾不存在'}), 404
            
            # 根據作業系統開啟資料夾
            if platform.system() == "Windows":
                # 使用 os.startfile 更可靠，不會有退出狀態問題
                import os
                try:
                    os.startfile(str(folder_path))
                except Exception as e:
                    # 如果 os.startfile 失敗，嘗試使用 explorer
                    result = subprocess.run(['explorer', str(folder_path)], 
                                          capture_output=True, text=True)
                    if result.returncode != 0:
                        raise Exception(f"無法開啟資料夾: {result.stderr}")
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(['open', str(folder_path)], check=True)
            else:  # Linux
                subprocess.run(['xdg-open', str(folder_path)], check=True)
            
            return jsonify({
                'success': True, 
                'message': f'已開啟 {folder_type} 資料夾',
                'path': str(folder_path)
            })
            
        except Exception as e:
            logger.error(f"開啟資料夾失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/workspace-info')
    def get_workspace_info():
        """獲取工作空間資訊"""
        try:
            settings_manager, _, _, _ = get_components()
            workspace_info = settings_manager.desktop_manager.get_workspace_info()
            return jsonify({
                'success': True,
                'workspace_info': workspace_info
            })
        except Exception as e:
            logger.error(f"獲取工作空間資訊失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/templates', methods=['GET'])
    def get_templates():
        """獲取模板列表"""
        try:
            settings_manager, _, _, _ = get_components()
            template_dir = settings_manager.get_directory_paths()['template']
            templates = []
            
            if template_dir.exists():
                for file_path in template_dir.iterdir():
                    if file_path.is_file() and file_path.suffix.lower() in ['.docx', '.doc', '.pdf']:
                        templates.append({
                            'filename': file_path.name,
                            'name': file_path.stem,
                            'path': str(file_path),
                            'size': file_path.stat().st_size,
                            'modified': file_path.stat().st_mtime
                        })
            
            return jsonify({
                'success': True,
                'templates': templates
            })
        except Exception as e:
            logger.error(f"獲取模板列表失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/templates/upload', methods=['POST'])
    def upload_template():
        """上傳模板檔案"""
        try:
            logger.info(f"收到模板上傳請求，檔案數量: {len(request.files)}")
            if 'template' not in request.files:
                logger.error("沒有上傳檔案")
                return jsonify({'error': '沒有上傳檔案'}), 400
            
            file = request.files['template']
            if file.filename == '':
                logger.error("沒有選擇檔案")
                return jsonify({'error': '沒有選擇檔案'}), 400
            
            logger.info(f"上傳檔案: {file.filename}")
            
            # 檢查檔案格式
            if not file.filename.lower().endswith(('.docx', '.doc', '.pdf')):
                logger.error(f"不支援的檔案格式: {file.filename}")
                return jsonify({'error': '只支援 DOCX、DOC 與 PDF 格式'}), 400
            
            # 保存檔案到模板目錄
            settings_manager, _, _, _ = get_components()
            dir_paths = settings_manager.get_directory_paths()
            template_dir = dir_paths['template']
            logger.info(f"模板目錄: {template_dir}")
            template_dir.mkdir(parents=True, exist_ok=True)
            
            # 處理重複檔案名，添加時間戳
            file_path = template_dir / file.filename
            if file_path.exists():
                name_part = file_path.stem
                ext_part = file_path.suffix
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{name_part}_{timestamp}{ext_part}"
                file_path = template_dir / filename
                logger.info(f"File already exists, using new name: {filename}")
            else:
                filename = file.filename
            
            file.save(str(file_path))
            logger.info(f"檔案已保存到: {file_path}")

            # 預處理：占位符/標籤分析，輸出 meta 與 labels
            try:
                import re, json as _json
                from pathlib import Path as _Path
                meta_dir = _Path(dir_paths['cache']) / 'templates_meta'
                meta_dir.mkdir(parents=True, exist_ok=True)
                meta = {
                    'filename': filename,  # 使用實際保存的檔案名
                    'path': str(file_path),
                    'has_jinja': False,
                    'placeholders': [],
                    'labels': [],
                    'tables': [],
                    'language_hints': []
                }

                suffix = file_path.suffix.lower()
                text_chunks = []

                if suffix == '.docx':
                    try:
                        from docx import Document as _Doc
                        from docx.shared import RGBColor
                        _doc = _Doc(str(file_path))
                        
                        # 著色欄位偵測
                        colored_fields = []
                        
                        # paragraphs
                        for p in _doc.paragraphs:
                            if p.text:
                                text_chunks.append(p.text)
                                # 檢查段落中的著色文字
                                for run in p.runs:
                                    if run.font.color and run.font.color.rgb:
                                        # 檢查是否為非黑色（著色）
                                        if run.font.color.rgb != RGBColor(0, 0, 0):
                                            colored_fields.append({
                                                'text': run.text.strip(),
                                                'color': str(run.font.color.rgb),
                                                'type': 'paragraph'
                                            })
                        
                        # tables
                        for t in _doc.tables:
                            table_txt = []
                            for row in t.rows:
                                row_vals = []
                                for cell in row.cells:
                                    if cell.text:
                                        text_chunks.append(cell.text)
                                        row_vals.append(cell.text)
                                        # 檢查表格中的著色文字
                                        for p in cell.paragraphs:
                                            for run in p.runs:
                                                if run.font.color and run.font.color.rgb:
                                                    if run.font.color.rgb != RGBColor(0, 0, 0):
                                                        colored_fields.append({
                                                            'text': run.text.strip(),
                                                            'color': str(run.font.color.rgb),
                                                            'type': 'table'
                                                        })
                                if row_vals:
                                    table_txt.append(row_vals)
                            if table_txt:
                                meta['tables'].append({'rows': len(table_txt), 'cols': max(len(r) for r in table_txt)})
                        
                        # 添加著色欄位資訊
                        meta['colored_fields'] = colored_fields
                        
                    except Exception as e:
                        logger.warning(f"模板預處理失敗: {e}")
                        pass
                else:
                    # 非 docx：只做淺層分析（不打開原格式）
                    text_chunks = []

                all_text = "\n".join(text_chunks)

                # 檢測 jinja
                if '{{' in all_text and '}}' in all_text:
                    meta['has_jinja'] = True

                # 占位符樣式
                ph_patterns = [
                    r"\{\{\s*([\w\.\-]+)\s*\}\}",  # jinja
                    r"\[\[\s*([\w\.\-]+)\s*\]\]",
                    r"<<\s*([\w\.\-]+)\s*>>",
                    r"\{\s*([\w\.\-]+)\s*\}",
                    r"《\s*([\w\.\-]+)\s*》",
                ]
                placeholders = set()
                for pat in ph_patterns:
                    for m in re.findall(pat, all_text):
                        placeholders.add(m)
                meta['placeholders'] = sorted(placeholders)

                # 標籤：值 偵測（中英冒號）
                label_pat = r"(^|\n)\s*([\w\u4e00-\u9fa5\uff08\uff09\u3001\-\/\s]{1,40})[\:：]\s*"
                labels = set()
                for m in re.finditer(label_pat, all_text):
                    label = m.group(2).strip()
                    if label:
                        labels.add(label)
                meta['labels'] = sorted(labels)

                # 簡單語言線索
                if re.search(r"[\u4e00-\u9fa5]", all_text):
                    meta['language_hints'].append('zh')
                if re.search(r"[A-Za-z]", all_text):
                    meta['language_hints'].append('en')

                # 寫入 meta 檔
                meta_path = meta_dir / f"{file_path.stem}.json"
                with open(meta_path, 'w', encoding='utf-8') as mf:
                    _json.dump(meta, mf, ensure_ascii=False, indent=2)

                # 產出/合併 labels.yaml（草稿）：只存偵測到的標籤供人工調整
                try:
                    import yaml
                    labels_yaml = _Path(dir_paths['prompts']).parent / 'labels.yaml'
                    existing = {}
                    if labels_yaml.exists():
                        with open(labels_yaml, 'r', encoding='utf-8') as lf:
                            try:
                                existing = yaml.safe_load(lf) or {}
                            except Exception:
                                existing = {}
                    detected = sorted(labels)
                    existing_detected = set((existing.get('detected_labels') or []))
                    combined = sorted(set(detected) | existing_detected)
                    existing['detected_labels'] = combined
                    with open(labels_yaml, 'w', encoding='utf-8') as lf:
                        yaml.safe_dump(existing, lf, allow_unicode=True)
                except Exception as _e:
                    logger.warning(f"labels.yaml 產出失敗: {_e}")

                analysis = {
                    'placeholders': meta['placeholders'],
                    'labels': meta['labels'],
                    'has_jinja': meta['has_jinja'],
                    'meta_path': str(meta_path)
                }
            except Exception as e:
                logger.warning(f"模板預處理失敗: {e}")
                analysis = None
            
            return jsonify({
                'success': True,
                'message': '模板上傳成功',
                'filename': filename,  # 使用實際保存的檔案名
                'path': str(file_path),
                'analysis': analysis
            })
            
        except Exception as e:
            logger.error(f"模板上傳失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/prompts', methods=['GET'])
    def get_prompts():
        """獲取提示詞列表"""
        try:
            settings_manager, _, _, _ = get_components()
            prompts_dir = settings_manager.get_directory_paths()['prompts']
            prompts = []
            
            if prompts_dir.exists():
                for file_path in prompts_dir.iterdir():
                    if file_path.is_file() and file_path.suffix.lower() == '.yaml':
                        try:
                            import yaml
                            with open(file_path, 'r', encoding='utf-8') as f:
                                prompt_data = yaml.safe_load(f)
                            
                            prompts.append({
                                'filename': file_path.name,
                                'name': prompt_data.get('name', file_path.stem),
                                'description': prompt_data.get('description', ''),
                                'path': str(file_path),
                                'modified': file_path.stat().st_mtime
                            })
                        except Exception as e:
                            logger.warning(f"讀取提示詞檔案失敗 {file_path}: {e}")
                            continue
            
            return jsonify({
                'success': True,
                'prompts': prompts
            })
        except Exception as e:
            logger.error(f"獲取提示詞列表失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/prompts/<filename>', methods=['GET'])
    def get_prompt(filename):
        """獲取單個提示詞內容"""
        try:
            settings_manager, _, _, _ = get_components()
            prompts_dir = settings_manager.get_directory_paths()['prompts']
            prompt_path = prompts_dir / filename
            
            if not prompt_path.exists():
                return jsonify({'error': '提示詞檔案不存在'}), 404
            
            import yaml
            with open(prompt_path, 'r', encoding='utf-8') as f:
                prompt_data = yaml.safe_load(f)
            
            return jsonify({
                'success': True,
                'prompt': prompt_data
            })
        except Exception as e:
            logger.error(f"獲取提示詞內容失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/prompts', methods=['POST'])
    def save_prompt():
        """保存提示詞"""
        try:
            data = request.get_json()
            name = data.get('name')
            description = data.get('description', '')
            content = data.get('content')
            fields = data.get('fields', [])
            
            if not name or not content:
                return jsonify({'error': '提示詞名稱和內容不能為空'}), 400
            
            # 創建提示詞資料
            prompt_data = {
                'name': name,
                'description': description,
                'content': content,
                'fields': fields,
                'created_date': datetime.now().strftime('%Y-%m-%d %H:%M'),
                'modified_date': datetime.now().strftime('%Y-%m-%d %H:%M')
            }
            
            # 保存到檔案
            settings_manager, _, _, _ = get_components()
            prompts_dir = settings_manager.get_directory_paths()['prompts']
            prompts_dir.mkdir(parents=True, exist_ok=True)
            
            filename = f"{name.replace(' ', '_')}.yaml"
            prompt_path = prompts_dir / filename
            
            import yaml
            with open(prompt_path, 'w', encoding='utf-8') as f:
                yaml.dump(prompt_data, f, default_flow_style=False, allow_unicode=True)
            
            return jsonify({
                'success': True,
                'message': '提示詞保存成功',
                'filename': filename,
                'path': str(prompt_path)
            })
            
        except Exception as e:
            logger.error(f"保存提示詞失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/profiles', methods=['GET'])
    def get_profiles():
        """獲取配置列表"""
        try:
            settings_manager, _, _, _ = get_components()
            profiles_dir = settings_manager.get_directory_paths()['profiles']
            profiles = []
            
            if profiles_dir.exists():
                for file_path in profiles_dir.iterdir():
                    if file_path.is_file() and file_path.suffix.lower() == '.yml':
                        try:
                            import yaml
                            with open(file_path, 'r', encoding='utf-8') as f:
                                profile_data = yaml.safe_load(f)
                            
                            profiles.append({
                                'filename': file_path.name,
                                'name': profile_data.get('name', file_path.stem),
                                'description': profile_data.get('description', ''),
                                'path': str(file_path),
                                'modified': file_path.stat().st_mtime
                            })
                        except Exception as e:
                            logger.warning(f"讀取配置檔案失敗 {file_path}: {e}")
                            continue
            
            return jsonify({
                'success': True,
                'profiles': profiles
            })
        except Exception as e:
            logger.error(f"獲取配置列表失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/works', methods=['GET'])
    def get_works():
        """獲取工作列表"""
        try:
            # 從設定管理器獲取工作列表
            settings_manager, _, _, _ = get_components()
            works = settings_manager.get_works()
            return jsonify({
                'success': True,
                'works': works
            })
        except Exception as e:
            logger.error(f"獲取工作列表失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/works', methods=['POST'])
    def create_work():
        """創建新工作（支援AI生成的Profile和提示詞）"""
        try:
            # 檢查請求類型
            if request.is_json:
                # JSON請求（舊格式，向後相容）
                data = request.get_json()
                work_name = data.get('name')
                work_description = data.get('description', '')
                work_type = data.get('type')
                brand = data.get('brand', '')
                profile = data.get('profile', '')
                template = data.get('template', '')
                prompt = data.get('prompt', '')
                ai_generated = False
            else:
                # FormData請求（新格式，AI引導式）
                work_name = request.form.get('name')
                work_description = request.form.get('description', '')
                work_type = request.form.get('type')
                profile_json = request.form.get('profile')
                prompt_text = request.form.get('prompt')
                template_file = request.files.get('template')
                ai_generated = True
                
                # 處理模板檔案
                template_path = None
                if template_file and template_file.filename:
                    template_path = save_template_file(template_file)
                
                # 解析Profile
                try:
                    profile = json.loads(profile_json) if profile_json else {}
                except json.JSONDecodeError:
                    return jsonify({'error': 'Profile JSON格式錯誤'}), 400
                
                # 驗證Profile結構
                settings_manager, _, profile_manager, _ = get_components()
                if not profile_manager.validate_ai_profile_structure(profile):
                    return jsonify({'error': 'Profile結構不正確'}), 400
                
                # 設置其他變數
                brand = request.form.get('brand', '').strip()
                template = template_path or ''
                prompt = prompt_text or ''
            
            # 驗證基本資訊
            if not work_name or not work_description:
                return jsonify({'error': '工作名稱和描述不能為空'}), 400
            
            # 創建工作
            work_id = str(uuid.uuid4())
            work = {
                'id': work_id,
                'name': work_name,
                'description': work_description,
                'type': work_type,
                'brand': brand,
                'profile': profile,
                'template': template,
                'prompt': prompt,
                'ai_generated': ai_generated,  # 標記是否為AI生成
                'processed_count': 0,
                'learning_count': 0,
                'created_at': datetime.now().isoformat(),
                'status': 'active'
            }
            
            # 保存工作到設定管理器
            settings_manager, _, profile_manager, _ = get_components()
            success = settings_manager.save_work(work)
            
            if success:
                # 如果是AI生成的工作，保存Profile到專用目錄
                if ai_generated and profile:
                    try:
                        profile_manager.save_ai_generated_profile(work_id, profile)
                        logger.info(f"已保存AI生成的Profile for work {work_id}")
                    except Exception as e:
                        logger.warning(f"保存AI生成Profile失敗: {e}")
                
                # 品牌作為資訊欄位，不進行自動優化
                
                return jsonify({
                    'success': True,
                    'work': work,
                    'message': '工作創建成功'
                })
            else:
                return jsonify({'error': '工作保存失敗'}), 500
                
        except Exception as e:
            logger.error(f"創建工作失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/works/<work_id>', methods=['PUT'])
    def update_work(work_id):
        """更新工作"""
        try:
            settings_manager, _, profile_manager, _ = get_components()
            data = request.get_json()
            
            # 更新工作數據
            success = settings_manager.update_work(work_id, data)
            
            if success:
                return jsonify({'success': True, 'message': '工作更新成功'})
            else:
                return jsonify({'error': '工作更新失敗'}), 500
                
        except Exception as e:
            logger.error(f"更新工作失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/works/<work_id>', methods=['DELETE'])
    def delete_work(work_id):
        """刪除工作"""
        try:
            settings_manager, _, _, _ = get_components()
            success = settings_manager.delete_work(work_id)
            
            if success:
                return jsonify({'success': True, 'message': '工作刪除成功'})
            else:
                return jsonify({'error': '工作刪除失敗'}), 500
                
        except Exception as e:
            logger.error(f"刪除工作失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/works/<work_id>/versions', methods=['GET'])
    def get_work_versions(work_id):
        """獲取工作版本歷史"""
        try:
            settings_manager, _, _, _ = get_components()
            version_history = settings_manager.get_work_version_history(work_id)
            
            return jsonify({
                'success': True, 
                'versions': version_history,
                'count': len(version_history)
            })
                
        except Exception as e:
            logger.error(f"獲取工作版本歷史失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/works/<work_id>/rollback/<int:version_index>', methods=['POST'])
    def rollback_work_version(work_id, version_index):
        """回滾工作到指定版本"""
        try:
            settings_manager, _, _, _ = get_components()
            success = settings_manager.rollback_work_version(work_id, version_index)
            
            if success:
                return jsonify({'success': True, 'message': f'工作已回滾到版本 {version_index}'})
            else:
                return jsonify({'error': '版本回滾失敗'}), 500
                
        except Exception as e:
            logger.error(f"版本回滾失敗: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/health')
    def health_check():
        """健康檢查"""
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'version': '1.0.0'
        })

    @app.route('/api/pricing', methods=['GET'])
    def get_pricing_info():
        """獲取AI供應商定價資訊"""
        try:
            pricing_manager = get_pricing_manager()
            
            # 獲取查詢參數
            provider = request.args.get('provider')
            model = request.args.get('model')
            
            if model:
                # 獲取特定模型的定價
                pricing = pricing_manager.get_model_pricing(model)
                return jsonify({'success': True, 'pricing': pricing})
            elif provider:
                # 獲取特定供應商的定價
                pricing = pricing_manager.get_pricing_info(provider)
                return jsonify({'success': True, 'pricing': pricing})
            else:
                # 獲取所有定價資訊
                pricing = pricing_manager.get_pricing_info()
                return jsonify({'success': True, 'pricing': pricing})
                
        except Exception as e:
            logger.error(f"獲取定價資訊失敗: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/pricing/calculate', methods=['POST'])
    def calculate_cost():
        """計算成本"""
        try:
            data = request.get_json()
            input_tokens = data.get('input_tokens', 0)
            output_tokens = data.get('output_tokens', 0)
            model_name = data.get('model', 'gpt-4o')
            
            pricing_manager = get_pricing_manager()
            cost_info = pricing_manager.calculate_cost(input_tokens, output_tokens, model_name)
            
            return jsonify({'success': True, 'cost_info': cost_info})
            
        except Exception as e:
            logger.error(f"計算成本失敗: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500

    # 工作流程偏好設定API
    @app.route('/api/workflow-preferences/<work_id>', methods=['GET'])
    def get_workflow_preferences(work_id):
        """獲取工作流程偏好設定"""
        try:
            preferences_manager = get_preferences_manager()
            preferences = preferences_manager.get_workflow_preferences(work_id)
            
            return jsonify({
                'success': True, 
                'preferences': preferences,
                'work_id': work_id
            })
            
        except Exception as e:
            logger.error(f"獲取工作流程偏好設定失敗: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/workflow-preferences/<work_id>', methods=['POST'])
    def save_workflow_preferences(work_id):
        """保存工作流程偏好設定"""
        try:
            data = request.get_json()
            preferences = data.get('preferences', {})
            
            preferences_manager = get_preferences_manager()
            success = preferences_manager.save_workflow_preferences(work_id, preferences)
            
            if success:
                return jsonify({
                    'success': True, 
                    'message': '偏好設定保存成功',
                    'work_id': work_id
                })
            else:
                return jsonify({'success': False, 'error': '保存偏好設定失敗'}), 500
                
        except Exception as e:
            logger.error(f"保存工作流程偏好設定失敗: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/workflow-preferences/<work_id>', methods=['PUT'])
    def update_workflow_preference(work_id):
        """更新工作流程偏好設定"""
        try:
            data = request.get_json()
            key = data.get('key')
            value = data.get('value')
            
            if not key:
                return jsonify({'success': False, 'error': '缺少偏好設定鍵'}), 400
            
            preferences_manager = get_preferences_manager()
            success = preferences_manager.update_workflow_preference(work_id, key, value)
            
            if success:
                return jsonify({
                    'success': True, 
                    'message': '偏好設定更新成功',
                    'work_id': work_id,
                    'key': key,
                    'value': value
                })
            else:
                return jsonify({'success': False, 'error': '更新偏好設定失敗'}), 500
                
        except Exception as e:
            logger.error(f"更新工作流程偏好設定失敗: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/workflow-preferences/<work_id>', methods=['DELETE'])
    def delete_workflow_preferences(work_id):
        """刪除工作流程偏好設定"""
        try:
            preferences_manager = get_preferences_manager()
            success = preferences_manager.delete_workflow_preferences(work_id)
            
            if success:
                return jsonify({
                    'success': True, 
                    'message': '偏好設定刪除成功',
                    'work_id': work_id
                })
            else:
                return jsonify({'success': False, 'error': '刪除偏好設定失敗'}), 500
                
        except Exception as e:
            logger.error(f"刪除工作流程偏好設定失敗: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/workflow-preferences', methods=['GET'])
    def get_all_workflow_preferences():
        """獲取所有工作流程偏好設定"""
        try:
            preferences_manager = get_preferences_manager()
            all_preferences = preferences_manager.get_all_workflow_preferences()
            
            return jsonify({
                'success': True, 
                'preferences': all_preferences
            })
            
        except Exception as e:
            logger.error(f"獲取所有工作流程偏好設定失敗: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500
    
    return app

def generate_pdf_thumbnail(pdf_path, page_number, cache_dir):
    """生成PDF頁面縮圖"""
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        import io
        
        # 創建縮圖目錄
        thumbnail_dir = cache_dir / "thumbnails"
        thumbnail_dir.mkdir(exist_ok=True)
        
        # 生成縮圖檔案名
        pdf_name = pdf_path.stem
        thumbnail_name = f"{pdf_name}_page_{page_number}.png"
        thumbnail_path = thumbnail_dir / thumbnail_name
        
        # 如果縮圖已存在，直接返回
        if thumbnail_path.exists():
            return thumbnail_path
        
        # 打開PDF檔案
        doc = fitz.open(pdf_path)
        if page_number > len(doc) or page_number < 1:
            return None
        
        # 獲取指定頁面
        page = doc[page_number - 1]
        
        # 設置縮圖參數
        zoom = 0.5  # 縮放比例
        mat = fitz.Matrix(zoom, zoom)
        
        # 渲染頁面為圖片
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        
        # 保存縮圖
        with open(thumbnail_path, 'wb') as f:
            f.write(img_data)
        
        doc.close()
        return thumbnail_path
        
    except ImportError:
        logger.warning("PyMuPDF 未安裝，無法生成PDF縮圖")
        return None
    except Exception as e:
        logger.error(f"生成PDF縮圖失敗: {e}")
        return None

# 創建應用實例
app = create_app()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
