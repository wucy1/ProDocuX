#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Profile 學習模組
提供從修正資料與更正後的 Word 檔學習的接口。
實現真正的規則學習邏輯。
"""

import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional
from docx import Document
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

class ProfileLearner:
    """Profile 學習器"""

    def __init__(self) -> None:
        self.learning_history = []
        logger.info("ProfileLearner 已初始化")

    def learn_from_corrections(
        self,
        original_data: Dict[str, Any],
        corrected_data: Dict[str, Any],
        source_file: str,
        profile_name: str,
    ) -> bool:
        """
        從 JSON 修正資料學習
        
        Args:
            original_data: 原始提取的資料
            corrected_data: 修正後的資料
            source_file: 來源檔案路徑
            profile_name: Profile名稱
            
        Returns:
            學習是否成功
        """
        try:
            logger.info(
                "執行 learn_from_corrections | profile=%s | source=%s",
                profile_name,
                source_file,
            )
            
            # 分析差異
            differences = self._analyze_differences(original_data, corrected_data)
            
            if not differences:
                logger.info("沒有發現差異，跳過學習")
                return True
            
            # 分析修正模式
            patterns = self._analyze_correction_patterns(differences)
            
            # 記錄學習歷史
            learning_record = {
                'type': 'json_correction',
                'profile': profile_name,
                'source_file': source_file,
                'differences': differences,
                'patterns': patterns,
                'timestamp': self._get_timestamp()
            }
            self.learning_history.append(learning_record)
            
            # 更新Profile規則
            success = self._update_profile_rules(profile_name, patterns)
            
            logger.info(f"JSON學習完成，發現 {len(differences)} 個差異，生成 {len(patterns)} 個模式")
            return success
            
        except Exception as e:
            logger.error("learn_from_corrections 失敗: %s", e)
            return False

    def learn_from_word_document(
        self,
        original_data: Dict[str, Any],
        corrected_docx_path: str,
        source_file: str,
        profile_name: str,
    ) -> bool:
        """
        從更正後的 Word 檔學習
        
        Args:
            original_data: 原始提取的資料
            corrected_docx_path: 修正後的Word檔案路徑
            source_file: 來源檔案路徑
            profile_name: Profile名稱
            
        Returns:
            學習是否成功
        """
        try:
            logger.info(
                "執行 learn_from_word_document | profile=%s | source=%s | corrected=%s",
                profile_name,
                source_file,
                corrected_docx_path,
            )
            
            if not Path(corrected_docx_path).exists():
                logger.warning("找不到更正後的 Word 檔: %s", corrected_docx_path)
                return False
            
            # 解析Word文檔內容
            word_content = self._parse_word_document(corrected_docx_path)
            
            if not word_content:
                logger.warning("無法解析Word文檔內容")
                return False
            
            # 比對原始JSON與Word內容的差異
            differences = self._compare_json_with_word(original_data, word_content)
            
            if not differences:
                logger.info("沒有發現差異，跳過學習")
                return True
            
            # 分析修正模式
            patterns = self._analyze_correction_patterns(differences)
            
            # 記錄學習歷史
            learning_record = {
                'type': 'word_correction',
                'profile': profile_name,
                'source_file': source_file,
                'corrected_file': corrected_docx_path,
                'differences': differences,
                'patterns': patterns,
                'timestamp': self._get_timestamp()
            }
            self.learning_history.append(learning_record)
            
            # 更新Profile規則
            success = self._update_profile_rules(profile_name, patterns)
            
            logger.info(f"Word學習完成，發現 {len(differences)} 個差異，生成 {len(patterns)} 個模式")
            return success
            
        except Exception as e:
            logger.error("learn_from_word_document 失敗: %s", e)
            return False
    
    def _parse_word_document(self, docx_path: str) -> Dict[str, Any]:
        """
        解析Word文檔內容
        
        Args:
            docx_path: Word檔案路徑
            
        Returns:
            解析後的文檔內容
        """
        try:
            doc = Document(docx_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'colored_text': [],
                'structured_data': {}
            }
            
            # 解析段落
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_data = {
                        'index': i,
                        'text': paragraph.text.strip(),
                        'runs': []
                    }
                    
                    # 檢查著色文字
                    for run in paragraph.runs:
                        if run.font.color and run.font.color.rgb != RGBColor(0, 0, 0):
                            para_data['runs'].append({
                                'text': run.text,
                                'color': str(run.font.color.rgb),
                                'is_colored': True
                            })
                        else:
                            para_data['runs'].append({
                                'text': run.text,
                                'is_colored': False
                            })
                    
                    content['paragraphs'].append(para_data)
            
            # 解析表格
            for i, table in enumerate(doc.tables):
                table_data = {
                    'index': i,
                    'rows': []
                }
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        table_data['rows'].append(row_data)
                
                if table_data['rows']:
                    content['tables'].append(table_data)
            
            # 提取著色文字
            for para in content['paragraphs']:
                for run in para['runs']:
                    if run.get('is_colored', False):
                        content['colored_text'].append({
                            'text': run['text'],
                            'color': run.get('color', ''),
                            'paragraph_index': para['index']
                        })
            
            # 嘗試從文檔中提取結構化資料
            content['structured_data'] = self._extract_structured_data_from_word(content)
            
            return content
            
        except Exception as e:
            logger.error(f"解析Word文檔失敗: {e}")
            return {}
    
    def _extract_structured_data_from_word(self, word_content: Dict[str, Any]) -> Dict[str, Any]:
        """
        從Word內容中提取結構化資料
        
        Args:
            word_content: Word文檔內容
            
        Returns:
            提取的結構化資料
        """
        structured_data = {}
        
        # 從段落中提取資料
        for para in word_content['paragraphs']:
            text = para['text']
            
            # 尋找標籤：值模式
            label_patterns = [
                r'([^:：]+)[:：]\s*(.+)',
                r'([^：]+)[：]\s*(.+)'
            ]
            
            for pattern in label_patterns:
                match = re.search(pattern, text)
                if match:
                    label = match.group(1).strip()
                    value = match.group(2).strip()
                    structured_data[label] = value
                    break
        
        # 從表格中提取資料
        for table in word_content['tables']:
            for row in table['rows']:
                if len(row) >= 2:
                    # 假設第一欄是標籤，第二欄是值
                    label = row[0].strip()
                    value = row[1].strip()
                    if label and value:
                        structured_data[label] = value
        
        return structured_data
    
    def _compare_json_with_word(self, original_data: Dict[str, Any], 
                               word_content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        比對原始JSON與Word內容的差異
        
        Args:
            original_data: 原始JSON資料
            word_content: Word文檔內容
            
        Returns:
            差異列表
        """
        differences = []
        word_data = word_content.get('structured_data', {})
        
        # 比較所有欄位
        all_keys = set(original_data.keys()) | set(word_data.keys())
        
        for key in all_keys:
            original_value = original_data.get(key, '')
            word_value = word_data.get(key, '')
            
            if original_value != word_value:
                differences.append({
                    'type': 'modified',
                    'path': key,
                    'original': original_value,
                    'corrected': word_value,
                    'source': 'word_document'
                })
        
        return differences
    
    def _analyze_differences(self, original: Dict[str, Any], 
                           corrected: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        分析兩個資料字典的差異
        
        Args:
            original: 原始資料
            corrected: 修正後資料
            
        Returns:
            差異列表
        """
        differences = []
        
        def compare_dicts(orig, corr, path=""):
            for key in set(orig.keys()) | set(corr.keys()):
                current_path = f"{path}.{key}" if path else key
                
                if key not in orig:
                    differences.append({
                        "type": "added",
                        "path": current_path,
                        "original": None,
                        "corrected": corr[key],
                        "source": "json_correction"
                    })
                elif key not in corr:
                    differences.append({
                        "type": "removed",
                        "path": current_path,
                        "original": orig[key],
                        "corrected": None,
                        "source": "json_correction"
                    })
                elif isinstance(orig[key], dict) and isinstance(corr[key], dict):
                    compare_dicts(orig[key], corr[key], current_path)
                elif orig[key] != corr[key]:
                    differences.append({
                        "type": "modified",
                        "path": current_path,
                        "original": orig[key],
                        "corrected": corr[key],
                        "source": "json_correction"
                    })
        
        compare_dicts(original, corrected)
        return differences
    
    def _analyze_correction_patterns(self, differences: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        分析修正模式
        
        Args:
            differences: 差異列表
            
        Returns:
            修正模式列表
        """
        patterns = []
        
        for diff in differences:
            if diff['type'] == 'modified':
                pattern = {
                    'field': diff['path'],
                    'original_pattern': self._extract_pattern(diff['original']),
                    'corrected_pattern': self._extract_pattern(diff['corrected']),
                    'transformation_rule': self._generate_transformation_rule(diff),
                    'confidence': self._calculate_confidence(diff)
                }
                patterns.append(pattern)
            elif diff['type'] == 'added':
                pattern = {
                    'field': diff['path'],
                    'action': 'add_field',
                    'value_pattern': self._extract_pattern(diff['corrected']),
                    'confidence': 0.8
                }
                patterns.append(pattern)
        
        return patterns
    
    def _extract_pattern(self, value: Any) -> str:
        """
        從值中提取模式
        
        Args:
            value: 要分析的值
            
        Returns:
            提取的模式
        """
        if not isinstance(value, str):
            return str(value)
        
        # 數字模式
        if re.match(r'^\d+$', value):
            return 'NUMBER'
        elif re.match(r'^\d+\.\d+$', value):
            return 'DECIMAL'
        elif re.match(r'^\d+%$', value):
            return 'PERCENTAGE'
        
        # 日期模式
        if re.match(r'\d{4}-\d{2}-\d{2}', value):
            return 'DATE_YYYY-MM-DD'
        elif re.match(r'\d{2}/\d{2}/\d{4}', value):
            return 'DATE_MM/DD/YYYY'
        
        # 郵箱模式
        if re.match(r'[^@]+@[^@]+\.[^@]+', value):
            return 'EMAIL'
        
        # 電話模式
        if re.match(r'[\d\-\+\(\)\s]+', value) and len(value) >= 7:
            return 'PHONE'
        
        # 長度模式（考慮中文字符）
        # 中文字符通常比英文字符佔用更多空間，需要調整長度判斷
        char_count = len(value)
        # 粗略估算：中文字符約佔2個字符寬度
        estimated_width = sum(2 if '\u4e00' <= char <= '\u9fff' else 1 for char in value)
        
        if estimated_width > 60:  # 降低閾值，因為中文字符較短
            return 'LONG_TEXT'
        elif estimated_width > 20:  # 降低閾值
            return 'MEDIUM_TEXT'
        else:
            return 'SHORT_TEXT'
    
    def _generate_transformation_rule(self, diff: Dict[str, Any]) -> str:
        """
        生成轉換規則
        
        Args:
            diff: 差異資料
            
        Returns:
            轉換規則描述
        """
        original = diff['original']
        corrected = diff['corrected']
        
        # 大小寫轉換
        if str(original).lower() == str(corrected).lower():
            return 'CASE_CONVERSION'
        
        # 格式轉換
        if re.match(r'^\d+$', str(original)) and re.match(r'^\d+\.\d+$', str(corrected)):
            return 'INT_TO_DECIMAL'
        
        # 語言轉換（簡單檢測）
        if self._is_chinese(str(original)) and not self._is_chinese(str(corrected)):
            return 'TRANSLATE_TO_ENGLISH'
        elif not self._is_chinese(str(original)) and self._is_chinese(str(corrected)):
            return 'TRANSLATE_TO_CHINESE'
        
        # 預設規則
        return 'CUSTOM_TRANSFORMATION'
    
    def _is_chinese(self, text: str) -> bool:
        """檢測是否包含中文字符"""
        return bool(re.search(r'[\u4e00-\u9fff]', text))
    
    def _calculate_confidence(self, diff: Dict[str, Any]) -> float:
        """
        計算修正的可信度
        
        Args:
            diff: 差異資料
            
        Returns:
            可信度分數 (0-1)
        """
        confidence = 0.5  # 基礎分數
        
        # 基於變化程度調整
        original_len = len(str(diff['original']))
        corrected_len = len(str(diff['corrected']))
        
        if original_len > 0:
            change_ratio = abs(corrected_len - original_len) / original_len
            if change_ratio < 0.1:  # 變化很小
                confidence += 0.3
            elif change_ratio < 0.5:  # 變化中等
                confidence += 0.1
        
        # 基於模式匹配調整
        if self._extract_pattern(diff['original']) == self._extract_pattern(diff['corrected']):
            confidence += 0.2
        
        return min(confidence, 1.0)
    
    def _update_profile_rules(self, profile_name: str, patterns: List[Dict[str, Any]]) -> bool:
        """
        更新Profile規則
        
        Args:
            profile_name: Profile名稱
            patterns: 修正模式列表
            
        Returns:
            更新是否成功
        """
        try:
            # 這裡可以實現具體的Profile更新邏輯
            # 目前先記錄到學習歷史中
            logger.info(f"更新Profile規則: {profile_name}, 模式數量: {len(patterns)}")
            
            for pattern in patterns:
                logger.info(f"  - 欄位: {pattern['field']}, 規則: {pattern.get('transformation_rule', 'N/A')}")
            
            return True
            
        except Exception as e:
            logger.error(f"更新Profile規則失敗: {e}")
            return False
    
    def _get_timestamp(self) -> str:
        """獲取當前時間戳"""
        from datetime import datetime
        return datetime.now().isoformat()
    
    def get_learning_history(self) -> List[Dict[str, Any]]:
        """獲取學習歷史"""
        return self.learning_history
    
    def clear_learning_history(self):
        """清除學習歷史"""
        self.learning_history = []
    
    def learn_from_repeated_processing(self, work_id: str, processing_history: List[Dict[str, Any]]) -> bool:
        """
        從反覆處理中學習
        
        Args:
            work_id: 工作ID
            processing_history: 處理歷史記錄
            
        Returns:
            學習是否成功
        """
        try:
            logger.info(f"開始反覆處理學習: {work_id}, 歷史記錄數量: {len(processing_history)}")
            
            if len(processing_history) < 2:
                logger.info("處理歷史不足，跳過反覆學習")
                return True
            
            # 分析重複的修正模式
            repeated_patterns = self._analyze_repeated_patterns(processing_history)
            
            if not repeated_patterns:
                logger.info("沒有發現重複模式，跳過學習")
                return True
            
            # 分析學習趨勢
            learning_trends = self._analyze_learning_trends(processing_history)
            
            # 記錄學習歷史
            learning_record = {
                'type': 'repeated_processing',
                'work_id': work_id,
                'processing_count': len(processing_history),
                'repeated_patterns': repeated_patterns,
                'learning_trends': learning_trends,
                'timestamp': self._get_timestamp()
            }
            self.learning_history.append(learning_record)
            
            # 更新Profile規則
            success = self._update_profile_rules(f"work_{work_id}", repeated_patterns)
            
            logger.info(f"反覆處理學習完成，發現 {len(repeated_patterns)} 個重複模式")
            return success
            
        except Exception as e:
            logger.error(f"反覆處理學習失敗: {e}")
            return False
    
    def _analyze_repeated_patterns(self, processing_history: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        分析重複的修正模式
        
        Args:
            processing_history: 處理歷史記錄
            
        Returns:
            重複模式列表
        """
        pattern_counts = {}
        field_patterns = {}
        
        # 統計每個欄位的修正模式
        for record in processing_history:
            if 'patterns' in record:
                for pattern in record['patterns']:
                    field = pattern['field']
                    if field not in field_patterns:
                        field_patterns[field] = []
                    
                    pattern_key = f"{pattern.get('transformation_rule', 'UNKNOWN')}_{pattern.get('original_pattern', '')}_{pattern.get('corrected_pattern', '')}"
                    field_patterns[field].append(pattern_key)
        
        # 找出重複的模式
        repeated_patterns = []
        for field, patterns in field_patterns.items():
            pattern_freq = {}
            for pattern in patterns:
                pattern_freq[pattern] = pattern_freq.get(pattern, 0) + 1
            
            # 找出出現次數超過1次的模式
            for pattern, count in pattern_freq.items():
                if count > 1:
                    repeated_patterns.append({
                        'field': field,
                        'pattern': pattern,
                        'frequency': count,
                        'confidence': min(count / len(processing_history), 1.0)
                    })
        
        return repeated_patterns
    
    def _analyze_learning_trends(self, processing_history: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        分析學習趨勢
        
        Args:
            processing_history: 處理歷史記錄
            
        Returns:
            學習趨勢分析
        """
        trends = {
            'total_corrections': 0,
            'average_corrections_per_processing': 0,
            'improvement_rate': 0.0,
            'most_corrected_fields': [],
            'learning_curve': []
        }
        
        if not processing_history:
            return trends
        
        # 統計總修正數
        total_corrections = sum(len(record.get('differences', [])) for record in processing_history)
        trends['total_corrections'] = total_corrections
        trends['average_corrections_per_processing'] = total_corrections / len(processing_history)
        
        # 分析最常修正的欄位
        field_correction_counts = {}
        for record in processing_history:
            for diff in record.get('differences', []):
                field = diff['path']
                field_correction_counts[field] = field_correction_counts.get(field, 0) + 1
        
        trends['most_corrected_fields'] = sorted(
            field_correction_counts.items(), 
            key=lambda x: x[1], 
            reverse=True
        )[:5]
        
        # 分析學習曲線（修正數量是否隨時間減少）
        corrections_over_time = [len(record.get('differences', [])) for record in processing_history]
        if len(corrections_over_time) > 1:
            first_half = corrections_over_time[:len(corrections_over_time)//2]
            second_half = corrections_over_time[len(corrections_over_time)//2:]
            
            avg_first = sum(first_half) / len(first_half)
            avg_second = sum(second_half) / len(second_half)
            
            if avg_first > 0:
                trends['improvement_rate'] = (avg_first - avg_second) / avg_first
        
        trends['learning_curve'] = corrections_over_time
        
        return trends

