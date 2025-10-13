#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文檔轉換器
將結構化資料轉換為不同格式的文檔
"""

import os
import sys
import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional, Union, List
from docxtpl import DocxTemplate
import pdfplumber
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

class DocumentTransformer:
    """文檔轉換器"""
    
    def __init__(self, profile_path: Optional[str] = None):
        """初始化文檔轉換器"""
        # 以使用者工作空間為優先
        try:
            from utils.settings_manager import SettingsManager
            sm = SettingsManager()
            self.template_dir = Path(sm.get_directory_paths()['template'])
        except Exception:
            # 後備：對於打包版本，使用工作空間的 templates
            if getattr(sys, 'frozen', False):
                try:
                    from utils.desktop_manager import DesktopManager
                    dm = DesktopManager()
                    self.template_dir = dm.workspace_dir / "templates"
                except:
                    self.template_dir = Path("templates")
            else:
                self.template_dir = Path("templates")
        
        # 載入Profile資訊
        self.profile_fields = {}
        if profile_path:
            self._load_profile(profile_path)
        
        logger.info("Document transformer initialized")
    
    def _load_profile(self, profile_path: str):
        """載入Profile資訊"""
        try:
            import yaml
            with open(profile_path, 'r', encoding='utf-8') as f:
                profile_data = yaml.safe_load(f)
            
            # 提取欄位名稱
            if 'fields' in profile_data:
                for field in profile_data['fields']:
                    if isinstance(field, dict) and 'name' in field:
                        self.profile_fields[field['name']] = field
                        logger.debug(f"Loading profile field: {field['name']}")
            
            logger.info(f"Profile loaded: {profile_path}, {len(self.profile_fields)} fields")
        except Exception as e:
            logger.warning(f"Failed to load profile: {e}")
    
    def transform(self, data: Dict[str, Any], 
                 template_path: Union[str, Path],
                 output_path: Union[str, Path],
                 output_format: str = "docx",
                 allow_fallback: bool = False) -> bool:
        """
        轉換結構化資料為文檔
        
        Args:
            data: 結構化資料
            template_path: 模板檔案路徑
            output_path: 輸出檔案路徑
            output_format: 輸出格式 (docx, pdf, json)
            
        Returns:
            轉換是否成功
        """
        try:
            template_path = Path(template_path)
            output_path = Path(output_path)
            
            # 確保輸出目錄存在
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            if output_format == "docx":
                return self._transform_to_docx(data, template_path, output_path, allow_fallback)
            elif output_format == "pdf":
                return self._transform_to_pdf(data, template_path, output_path, allow_fallback)
            elif output_format == "json":
                return self._transform_to_json(data, output_path)
            else:
                raise ValueError(f"不支援的輸出格式: {output_format}")
                
        except Exception as e:
            logger.error(f"Document transformation failed: {e}")
            # 直接拋出異常，不使用回退機制
            raise Exception(f"文檔轉換失敗: {e}")
    
    def _transform_to_docx(self, data: Dict[str, Any], 
                          template_path: Path, output_path: Path, allow_fallback: bool = False) -> bool:
        """轉換為Word文檔"""
        try:
            # 檢查是否為 docx
            if template_path.suffix.lower() != '.docx':
                # 對於非 docx，先嘗試直接複製（無渲染）
                import shutil
                shutil.copy2(template_path, output_path)
                logger.warning("Non-.docx template, copied directly; recommend using .docx for rendering")
                return True

            # 檢查模板元數據以判斷是否含 jinja 變數
            has_jinja = False
            try:
                # 首先檢查元數據文件
                meta_path = template_path.parent / f"{template_path.stem}.json"
                if meta_path.exists():
                    import json
                    with open(meta_path, 'r', encoding='utf-8') as f:
                        meta = json.load(f)
                    has_jinja = meta.get('has_jinja', False)
                    placeholders = meta.get('placeholders', [])
                    logger.info(f"Read from metadata: has_jinja={has_jinja}, placeholders={len(placeholders)}")
                else:
                    # 如果沒有元數據文件，回退到原始檢查
                    # 更準確的Jinja2變數檢測：檢查是否包含真正的模板變數
                    try:
                        from docx import Document
                        doc = Document(str(template_path))
                        all_text = ""
                        for para in doc.paragraphs:
                            all_text += para.text + "\n"
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    all_text += cell.text + "\n"
                        
                        # 檢查是否包含真正的Jinja2變數（字母數字組合）
                        import re
                        jinja_pattern = r'\{\{\s*[a-zA-Z_][a-zA-Z0-9_.]*\s*\}\}'
                        has_jinja = bool(re.search(jinja_pattern, all_text))
                        logger.info(f"Template variable check (fallback): has_jinja={has_jinja}")
                        if has_jinja:
                            matches = re.findall(jinja_pattern, all_text)
                            logger.info(f"Found Jinja2 variables: {matches[:5]}")
                    except Exception as e:
                        logger.warning(f"Template variable check failed: {e}")
                        has_jinja = False
            except Exception as e:
                logger.warning(f"Template variable check failed: {e}")
                has_jinja = True

            if has_jinja:
                template = DocxTemplate(template_path)
                
                # 添加調試信息
                logger.info(f"Starting Word template rendering: {template_path}")
                logger.info(f"Data structure: {list(data.keys())}")
                if '成分表' in data:
                    logger.info(f"Ingredients data: {data['成分表'][:2] if isinstance(data['成分表'], list) and len(data['成分表']) > 0 else 'No ingredients data'}")
                
                # 轉換資料結構以匹配模板期望的格式
                template_data = self._prepare_template_data(data)
                logger.info(f"Transformed data structure: {list(template_data.keys())}")
                
                try:
                    template.render(template_data)
                    template.save(output_path)
                    logger.info("Template rendering successful")
                except Exception as render_error:
                    logger.error(f"Template rendering failed: {render_error}")
                    # 嘗試保存原始模板作為備份
                    import shutil
                    backup_path = output_path.with_suffix('.backup.docx')
                    shutil.copy2(template_path, backup_path)
                    logger.info(f"Original template backup saved: {backup_path}")
                    raise
                
                # Jinja2渲染後，還需要進行成分表替換
                logger.info("Performing ingredient table replacement after Jinja2 rendering")
                self._replace_ingredients_tables_in_doc(output_path, data)
                
                # 添加淺藍底色標記到插入的資料
                self._highlight_inserted_data(output_path, data)
                
                logger.info(f"Word document generated: {output_path}")
                return True
            else:
                if not allow_fallback:
                    raise ValueError("模板不含變數且未允許回退替換模式")
                # 關鍵詞替換模式：以 docx 文段遍歷做簡單覆寫（標籤：值 / [[key]] / <<key>> / {key}）
                from docx import Document as DocxDocument
                doc = DocxDocument(str(template_path))

                # 扁平化 data 的鍵（a.b.c -> value）
                flat: Dict[str, Any] = {}
                def flatten(prefix: str, obj: Any):
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            flatten(f"{prefix}.{k}" if prefix else k, v)
                    else:
                        flat[prefix] = obj
                flatten("", data)
                
                # 特別處理成分表資料
                if '成分表' in data and isinstance(data['成分表'], list):
                    logger.info(f"Processing ingredients data, {len(data['成分表'])} ingredients")
                    # 為成分表創建表格替換資料
                    flat['成分表_表格'] = self._create_ingredients_table(data['成分表'])

                # 可匹配的占位樣式
                def variations(key: str) -> List[str]:
                    k = key
                    names = [k, k.split(".")[-1]]
                    outs = []
                    for n in names:
                        outs.extend([
                            f"[[{n}]]", f"<<{n}>>", f"{{{n}}}", f"《{n}》",
                            f"{n}：", f"{n}: "
                        ])
                    return outs

                # 段落替換
                for p in doc.paragraphs:
                    text = p.text
                    new_text = text
                    for k, v in flat.items():
                        for token in variations(k):
                            if token.endswith('：') or token.endswith(': '):
                                # 標籤：值 → 覆寫後方內容
                                label = token.rstrip()
                                if label in new_text:
                                    # 簡單策略：label 之後整行替換為 label + 值
                                    parts = new_text.split(label, 1)
                                    new_text = parts[0] + label + str(v)
                            else:
                                new_text = new_text.replace(token, str(v))
                    if new_text != text:
                        p.text = new_text

                # 表格儲存格替換
                for table_idx, table in enumerate(doc.tables):
                    logger.info(f"Checking table {table_idx}, rows: {len(table.rows)}")
                    if table.rows:
                        first_row_text = " ".join([cell.text.strip() for cell in table.rows[0].cells]).lower()
                        logger.info(f"Table {table_idx} first row content: {first_row_text}")
                    
                    # 檢查是否為成分表表格
                    if self._is_ingredients_table(table):
                        logger.info(f"Found ingredient table {table_idx}, performing replacement")
                        ingredients = self._extract_ingredients_from_data(data)
                        logger.info(f"Ingredients data: {ingredients}")
                        self._replace_ingredients_table(table, ingredients)
                    else:
                        # 一般表格替換
                        for row in table.rows:
                            for cell in row.cells:
                                txt = cell.text
                                new_txt = txt
                                for k, v in flat.items():
                                    for token in variations(k):
                                        if token.endswith('：') or token.endswith(': '):
                                            label = token.rstrip()
                                            if label in new_txt:
                                                parts = new_txt.split(label, 1)
                                                new_txt = parts[0] + label + str(v)
                                        else:
                                            new_txt = new_txt.replace(token, str(v))
                                if new_txt != txt:
                                    cell.text = new_txt

                doc.save(str(output_path))
                logger.info(f"Word document (keyword replacement mode) generated: {output_path}")
                
                # 添加淺藍底色標記
                self._highlight_inserted_data(Path(output_path), data)
                
                return True

        except Exception as e:
            logger.error(f"Word conversion failed: {e}")
            # 直接拋出異常，不使用回退機制
            raise Exception(f"Word轉換失敗: {e}")
    
    def _prepare_template_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """準備模板資料，確保變數名稱匹配"""
        # 首先處理嵌套結構，將其扁平化為Profile期望的格式
        flattened_data = self._flatten_nested_data(data)
        template_data = flattened_data.copy()
        
        # 處理成分表資料結構轉換
        if '成分表' in template_data:
            # 如果模板期望的是 成分.成分表 格式
            if '成分' not in template_data:
                template_data['成分'] = {}
            template_data['成分']['成分表'] = template_data['成分表']
            
            # 同時保持原有的成分表格式
            # template_data['成分表'] = template_data['成分表']
        
        # 處理其他可能的結構轉換
        # 基本資訊
        if '基本資訊' in template_data:
            basic_info = template_data['基本資訊']
            # 如果模板期望的是 產品基本資訊 格式
            if '產品基本資訊' not in template_data:
                template_data['產品基本資訊'] = basic_info
        
        # 製造資訊
        if '製造資訊' in template_data:
            manufacturer_info = template_data['製造資訊']
            # 如果模板期望的是 製造商資訊 格式
            if '製造商資訊' not in template_data:
                template_data['製造商資訊'] = manufacturer_info
        
        # 安全資訊
        if '安全資訊' in template_data:
            safety_info = template_data['安全資訊']
            # 如果模板期望的是 危害資訊 格式
            if '危害資訊' not in template_data:
                template_data['危害資訊'] = safety_info
        
        # 使用方法
        if '使用方法' in template_data:
            usage_info = template_data['使用方法']
            # 如果模板期望的是 使用說明 格式
            if '使用說明' not in template_data:
                template_data['使用說明'] = usage_info
        
        # 其他資訊
        if '其他資訊' in template_data:
            other_info = template_data['其他資訊']
            # 如果模板期望的是 理化性質 格式
            if '理化性質' not in template_data:
                template_data['理化性質'] = other_info
        
        return template_data
    
    def _flatten_nested_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """將嵌套的JSON結構扁平化為Profile期望的格式 - 通用版本"""
        flattened = {}
        
        # 定義欄位映射規則 - 支援多種可能的欄位名稱
        field_mappings = {
            # 產品基本資訊映射
            '產品名稱': [
                '產品名稱', 'product_name', 'name', 'title', '產品標題',
                '基本資訊.產品名稱', '產品檔案.基本資訊.產品名稱',
                '產品基本資訊.產品名稱', '處理結果.產品檔案.基本資訊.產品名稱'
            ],
            '產品類別': [
                '產品類別', 'product_category', 'category', 'type', '產品類型',
                '基本資訊.產品類別', '產品檔案.基本資訊.產品類型',
                '產品基本資訊.產品類型', '處理結果.產品檔案.基本資訊.產品類型'
            ],
            '產品劑型': [
                '產品劑型', 'product_form', 'form', '劑型', '物理形態', '物理狀態',
                '基本資訊.產品劑型', '產品檔案.基本資訊.物理形態',
                '產品基本資訊.產品劑型', '處理結果.產品檔案.物理化學特性.物理狀態'
            ],
            '產品用途': [
                '產品用途', 'product_use', 'use', '用途', 'intended_use',
                '基本資訊.產品用途', '產品檔案.基本資訊.產品用途',
                '產品基本資訊.產品用途', '處理結果.產品檔案.基本資訊.產品用途'
            ],
            '容量': [
                '容量', 'volume', 'size', 'content', '包裝容量',
                '基本資訊.容量', '產品檔案.穩定性與包裝.容量',
                '產品基本資訊.容量', '處理結果.產品檔案.穩定性與包裝.容量'
            ],
            '原產地': [
                '原產地', 'country_of_origin', 'origin', '原產國', '產地', 'country',
                '基本資訊.原產地', '產品檔案.基本資訊.原產國',
                '產品基本資訊.原產地', '處理結果.產品檔案.基本資訊.原產國',
                'result.product_info.basic.country', 'result.product_info.basic.origin'
            ],
            # 製造商資訊映射
            '製造商名稱': [
                '製造商名稱', 'manufacturer_name', 'manufacturer', '製造商', '公司名稱',
                '製造商資訊.公司名稱', '產品檔案.製造商資訊.公司名稱',
                '廠商與負責人資訊.製造商.名稱', '處理結果.產品檔案.製造商資訊.公司名稱'
            ],
            '製造商地址': [
                '製造商地址', 'manufacturer_address', 'manufacturer_addr', '地址',
                '製造商資訊.地址', '產品檔案.製造商資訊.地址',
                '廠商與負責人資訊.製造商.地址', '處理結果.產品檔案.製造商資訊.地址'
            ],
            '製造商聯絡方式': [
                '製造商聯絡方式', 'manufacturer_contact', 'manufacturer_phone', '電話', '聯絡方式',
                '製造商資訊.電話', '產品檔案.製造商資訊.電話',
                '廠商與負責人資訊.製造商.電話', '處理結果.產品檔案.製造商資訊.電話'
            ],
            # 輸入商資訊映射
            '輸入商名稱': [
                '輸入商名稱', 'importer_name', 'importer', '輸入商', '責任人', '歐盟負責人',
                '責任人資訊.公司名稱', '產品檔案.責任人資訊.公司名稱',
                '廠商與負責人資訊.歐盟負責人.名稱', '處理結果.產品檔案.責任人資訊.公司名稱'
            ],
            '輸入商地址': [
                '輸入商地址', 'importer_address', 'importer_addr',
                '責任人資訊.地址', '產品檔案.責任人資訊.地址',
                '廠商與負責人資訊.歐盟負責人.地址', '處理結果.產品檔案.責任人資訊.地址'
            ],
            '輸入商電話': [
                '輸入商電話', 'importer_phone', 'importer_contact',
                '責任人資訊.電話', '產品檔案.責任人資訊.電話',
                '廠商與負責人資訊.歐盟負責人.電話', '處理結果.產品檔案.責任人資訊.電話'
            ],
            # 成分表映射
            '成分表': [
                '成分表', 'ingredients', 'ingredient_list', '成分', '完整成分表', '完整成分表(INCI)',
                '主要成分', '成分資訊', 'ingredient_info', 'composition'
            ],
            # 其他資訊映射
            '有效期限': [
                '有效期限', 'expiry_date', 'shelf_life', '保質期', '產品保質期', '開封後保質期(PAO)',
                '穩定性與包裝.產品保質期', '產品檔案.穩定性與包裝.產品保質期',
                '穩定性與儲存.最短保質期', '處理結果.產品檔案.穩定性與包裝.產品保質期'
            ],
            '安全評估結果': [
                '安全評估結果', 'safety_assessment', 'safety_evaluation', '安全評估結論',
                '安全資訊.安全評估結論', '產品檔案.安全資訊.安全評估結論',
                '處理結果.產品檔案.安全資訊.安全評估結論'
            ],
            '成分安全性': [
                '成分安全性', 'ingredient_safety', 'safety_info',
                '安全資訊.成分安全性', '產品檔案.安全資訊.成分安全性',
                '處理結果.產品檔案.安全資訊.成分安全性'
            ],
            '使用限制': [
                '使用限制', 'usage_restrictions', 'restrictions', '使用限制和禁忌',
                '安全資訊.使用限制', '產品檔案.安全資訊.使用限制',
                '處理結果.產品檔案.安全資訊.使用限制'
            ],
            '使用方式': [
                '使用方式', 'usage_method', 'how_to_use', '使用方法', '使用注意事項',
                '安全資訊.標籤警語.使用注意事項', '產品檔案.安全資訊.標籤警語.使用注意事項',
                '處理結果.產品檔案.安全資訊.標籤警語.使用注意事項'
            ],
            '注意事項': [
                '注意事項', 'precautions', 'warnings', '使用注意事項',
                '安全資訊.標籤警語.使用注意事項', '產品檔案.安全資訊.標籤警語.使用注意事項',
                '處理結果.產品檔案.安全資訊.標籤警語.使用注意事項'
            ],
            '使用部位': [
                '使用部位', 'application_site', 'where_to_use', '適用部位'
            ]
        }
        
        # 遞歸搜尋並提取欄位值
        def extract_field_value(data: Dict[str, Any], field_paths: List[str]) -> Any:
            """遞歸搜尋欄位值"""
            for path in field_paths:
                try:
                    # 處理點號分隔的嵌套路徑
                    if '.' in path:
                        keys = path.split('.')
                        current_data = data
                        for key in keys:
                            if isinstance(current_data, dict) and key in current_data:
                                current_data = current_data[key]
                            else:
                                current_data = None
                                break
                        if current_data is not None:
                            return current_data
                    # 處理直接欄位名稱
                    elif path in data:
                        return data[path]
                except (KeyError, TypeError, AttributeError):
                    continue
            return None
        
        # 遞歸搜尋所有嵌套結構
        def recursive_search(data: Dict[str, Any], target_field: str) -> Any:
            """遞歸搜尋目標欄位"""
            if target_field in data:
                return data[target_field]
            
            for key, value in data.items():
                if isinstance(value, dict):
                    result = recursive_search(value, target_field)
                    if result is not None:
                        return result
                elif isinstance(value, list):
                    for item in value:
                        if isinstance(item, dict):
                            result = recursive_search(item, target_field)
                            if result is not None:
                                return result
            
            return None
        
        # 應用欄位映射
        for target_field, possible_paths in field_mappings.items():
            value = None
            
            # 首先嘗試直接路徑搜尋
            value = extract_field_value(data, possible_paths)
            
            # 如果沒找到，嘗試遞歸搜尋
            if value is None:
                for path in possible_paths:
                    if '.' not in path:  # 只對簡單欄位名稱進行遞歸搜尋
                        value = recursive_search(data, path)
                        if value is not None:
                            break
            
            if value is not None:
                flattened[target_field] = value
        
        # 保留原始資料中的其他欄位（非系統欄位）
        system_fields = ['_raw_response', '處理結果']
        for key, value in data.items():
            if key not in system_fields and key not in flattened:
                flattened[key] = value
        
        logger.info(f"🔧 Flattened fields: {list(flattened.keys())}")
        logger.info(f"🔧 Ingredients field exists: {'成分表' in flattened}")
        if '成分表' in flattened:
            logger.info(f"🔧 Ingredients type: {type(flattened['成分表'])}")
            if isinstance(flattened['成分表'], list):
                logger.info(f"🔧 Ingredients length: {len(flattened['成分表'])}")
        
        return flattened
    
    def _create_ingredients_table(self, ingredients: List[Dict[str, Any]]) -> str:
        """創建成分表表格字串"""
        if not ingredients:
            return ""
        
        # 創建表格標題行
        table_lines = ["INCI Name\tCas. No\tW/V%\t功能"]
        
        # 添加成分資料行
        for ingredient in ingredients:
            inci_name = ingredient.get('INCI名稱', '')
            cas_number = ingredient.get('CAS號碼', '') or '-'
            content = ingredient.get('含量', '')
            function = ingredient.get('功能', '')
            
            table_lines.append(f"{inci_name}\t{cas_number}\t{content}\t{function}")
        
        return "\n".join(table_lines)
    
    def _is_ingredients_table(self, table) -> bool:
        """檢查表格是否為成分表"""
        if not table.rows:
            return False
        
        # 檢查第一行是否包含成分表標題
        first_row = table.rows[0]
        first_row_text = " ".join([cell.text.strip() for cell in first_row.cells]).lower()
        
        ingredients_keywords = ['inci', 'cas', 'w/v', '功能', 'ingredient', 'content']
        return any(keyword in first_row_text for keyword in ingredients_keywords)
    
    def _extract_ingredients_from_data(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """從數據中提取成分表資訊，支援多種格式"""
        ingredients = []
        
        # 添加調試日誌
        logger.info(f"🔍 Starting ingredient extraction, data structure: {list(data.keys())}")
        logger.info(f"🔍 Original response length: {len(data.get('_raw_response', ''))}")
        
        # 1. 嘗試從Profile定義的成分表欄位提取
        ingredient_field_names = []
        
        # 從Profile中尋找可能的成分表欄位
        for field_name, field_info in self.profile_fields.items():
            if any(keyword in field_name.lower() for keyword in ['成分', 'ingredient', '全成分']):
                ingredient_field_names.append(field_name)
                logger.info(f"🔍 Found possible ingredient field: {field_name}")
        
        # 如果沒有找到，使用預設的欄位名稱
        if not ingredient_field_names:
            ingredient_field_names = ['成分表', '全成分表', 'ingredients']
            logger.info("🔍 使用預設成分表欄位名稱")
        
        # 嘗試從找到的欄位中提取成分
        for field_name in ingredient_field_names:
            # 首先嘗試在根層級尋找
            if field_name in data and data[field_name]:
                logger.info(f"🔍 在根層級找到成分表欄位 '{field_name}': {type(data[field_name])}")
                ingredients_found = self._extract_ingredients_from_field(data[field_name], field_name)
                if ingredients_found:
                    ingredients.extend(ingredients_found)
                    logger.info(f"✅ 從根層級提取到 {len(ingredients_found)} 個成分")
                    break
            # 如果根層級沒有，嘗試在 'data' 子物件中尋找
            elif 'data' in data and isinstance(data['data'], dict) and field_name in data['data'] and data['data'][field_name]:
                logger.info(f"🔍 在data子物件中找到成分表欄位 '{field_name}': {type(data['data'][field_name])}")
                ingredients_found = self._extract_ingredients_from_field(data['data'][field_name], field_name)
                if ingredients_found:
                    ingredients.extend(ingredients_found)
                    logger.info(f"✅ 從data子物件提取到 {len(ingredients_found)} 個成分")
                    break
        else:
            logger.warning(f"❌ 未找到成分表欄位，嘗試的欄位名稱: {ingredient_field_names}")
            logger.info(f"🔍 數據結構: {list(data.keys())}")
            if 'data' in data and isinstance(data['data'], dict):
                logger.info(f"🔍 data子物件結構: {list(data['data'].keys())}")
        
        # 2. 嘗試從主要成分欄位提取
        if '主要成分' in data and data['主要成分']:
            logger.info(f"🔍 找到主要成分欄位: {type(data['主要成分'])}")
            main_ingredients = data['主要成分']
            if isinstance(main_ingredients, str):
                # 如果是字串，嘗試解析
                ingredients.append({
                    'INCI名稱': main_ingredients,
                    'CAS號碼': '',
                    '含量': '',
                    '功能': ''
                })
                logger.info(f"✅ 添加主要成分字串: {main_ingredients[:50]}...")
            elif isinstance(main_ingredients, list):
                logger.info(f"🔍 主要成分列表長度: {len(main_ingredients)}")
                for i, ingredient in enumerate(main_ingredients):
                    if isinstance(ingredient, str):
                        ingredients.append({
                            'INCI名稱': ingredient,
                            'CAS號碼': '',
                            '含量': '',
                            '功能': ''
                        })
                        logger.info(f"✅ 添加主要成分 {i+1}: {ingredient[:50]}...")
        else:
            logger.info("ℹ️ 未找到主要成分欄位")
        
        # 3. 嘗試從原始響應中提取 Markdown 表格
        raw_response = data.get('_raw_response', '')
        if raw_response and '|' in raw_response:
            logger.info(f"🔍 原始響應包含Markdown表格標記，長度: {len(raw_response)}")
            markdown_ingredients = self._parse_markdown_ingredients_table(raw_response)
            ingredients.extend(markdown_ingredients)
            logger.info(f"✅ 從Markdown表格提取到 {len(markdown_ingredients)} 個成分")
        else:
            logger.info("ℹ️ 原始響應不包含Markdown表格標記")
        
        # 4. 檢查其他可能的成分欄位（增強通用性）
        possible_ingredient_fields = [
            'ingredients', '成分', 'ingredient_list', 'composition',
            'ingredient', 'ingredients_list', 'ingredient_table',
            '成分列表', '成分表', '主要成分', '成分資訊', '完整成分表',
            'ingredient_info', 'composition_list', 'ingredient_data',
            'ingredients_data', 'ingredient_details', 'composition_details',
            '標籤成分列表', '成分清單', 'ingredient_composition'
        ]
        
        for field in possible_ingredient_fields:
            if field in data and data[field]:
                logger.info(f"🔍 找到可能的成分欄位 '{field}': {type(data[field])}")
                # 嘗試從這個欄位提取成分
                if isinstance(data[field], list):
                    logger.info(f"🔍 欄位 '{field}' 是列表格式，長度: {len(data[field])}")
                    for i, item in enumerate(data[field]):
                        if isinstance(item, dict):
                            # 檢查是否包含成分相關的鍵
                            if any(key in item for key in ['INCI名稱', 'inci_name', 'name', '成分名稱', 'ingredient_name']):
                                # 標準化欄位名稱
                                standardized_item = self._standardize_ingredient_fields(item)
                                ingredients.append(standardized_item)
                                logger.info(f"✅ 從欄位 '{field}' 添加成分 {i+1}: {standardized_item.get('INCI名稱', 'Unknown')}")
                        elif isinstance(item, str) and len(item.strip()) > 0:
                            # 嘗試解析字串格式的成分
                            parsed_ingredient = self._parse_ingredient_string(item)
                            if parsed_ingredient:
                                ingredients.append(parsed_ingredient)
                                logger.info(f"✅ 從欄位 '{field}' 解析成分 {i+1}: {parsed_ingredient.get('INCI名稱', 'Unknown')}")
                elif isinstance(data[field], str) and len(data[field].strip()) > 0:
                    logger.info(f"🔍 欄位 '{field}' 是字串格式")
                    # 嘗試解析字串格式的成分
                    parsed_ingredient = self._parse_ingredient_string(data[field])
                    if parsed_ingredient:
                        ingredients.append(parsed_ingredient)
                        logger.info(f"✅ 從欄位 '{field}' 解析成分: {parsed_ingredient.get('INCI名稱', 'Unknown')}")
        
        logger.info(f"🎯 最終提取到 {len(ingredients)} 個成分")
        return ingredients
    
    def _standardize_ingredient_fields(self, ingredient: Dict[str, Any]) -> Dict[str, Any]:
        """標準化成分欄位名稱，確保與模板期望的格式一致"""
        standardized = {}
        
        # 名稱欄位映射
        name_fields = ['INCI名稱', 'inci_name', 'name', '成分名稱', 'ingredient_name', '成分名稱']
        for field in name_fields:
            if field in ingredient:
                standardized['INCI名稱'] = ingredient[field]
                break
        
        # CAS號碼欄位映射
        cas_fields = ['CAS號碼', 'cas_number', 'cas', 'CAS', 'cas_no']
        for field in cas_fields:
            if field in ingredient:
                standardized['CAS號碼'] = ingredient[field]
                break
        
        # 含量欄位映射
        content_fields = ['含量', 'content', 'percentage', '百分比', 'concentration', 'w/v%']
        for field in content_fields:
            if field in ingredient:
                standardized['含量'] = ingredient[field]
                break
        
        # 功能欄位映射
        function_fields = ['功能', 'function', 'purpose', '功能說明', 'role']
        for field in function_fields:
            if field in ingredient:
                standardized['功能'] = ingredient[field]
                break
        
        # 如果沒有找到標準欄位，保留原始欄位
        if not standardized:
            standardized = ingredient.copy()
        
        logger.info(f"🔧 標準化成分欄位: {list(standardized.keys())}")
        return standardized
    
    def _parse_ingredient_string(self, ingredient_str: str) -> Dict[str, Any]:
        """解析成分字串格式"""
        try:
            # 格式：Alcohol denat. (CAS: 64-17-5, 含量: 76.80000%, 功能: 抗泡劑、抗菌劑、收斂劑、遮蔽劑、溶劑、黏度控制劑)
            import re
            
            # 提取INCI名稱（括號前的部分）
            inci_match = re.match(r'^([^(]+)', ingredient_str.strip())
            if not inci_match:
                return None
            
            inci_name = inci_match.group(1).strip()
            
            # 提取CAS號碼
            cas_match = re.search(r'CAS:\s*([^,)]+)', ingredient_str)
            cas_number = cas_match.group(1).strip() if cas_match else ''
            
            # 提取含量
            content_match = re.search(r'含量:\s*([^,)]+)', ingredient_str)
            content = content_match.group(1).strip() if content_match else ''
            
            # 提取功能
            function_match = re.search(r'功能:\s*([^)]+)', ingredient_str)
            function = function_match.group(1).strip() if function_match else ''
            
            return {
                'INCI名稱': inci_name,
                'CAS號碼': cas_number,
                '含量': content,
                '功能': function
            }
            
        except Exception as e:
            logger.warning(f"解析成分字串失敗: {ingredient_str}, 錯誤: {e}")
            return None
    
    def _extract_ingredients_from_field(self, field_data: Any, field_name: str) -> List[Dict[str, Any]]:
        """從成分表欄位中提取成分"""
        ingredients = []
        
        if isinstance(field_data, list):
            logger.info(f"🔍 成分表列表長度: {len(field_data)}")
            for i, ingredient in enumerate(field_data):
                logger.info(f"🔍 成分 {i+1}: {type(ingredient)} - {str(ingredient)[:100]}...")
                if isinstance(ingredient, dict):
                    # 已經是字典格式，直接添加
                    ingredients.append(ingredient)
                    logger.info(f"✅ 添加字典格式成分: {ingredient.get('INCI名稱', 'Unknown')}")
                elif isinstance(ingredient, str):
                    # 字串格式，需要解析
                    parsed_ingredient = self._parse_ingredient_string(ingredient)
                    if parsed_ingredient:
                        ingredients.append(parsed_ingredient)
                        logger.info(f"✅ 解析字串格式成分: {parsed_ingredient.get('INCI名稱', 'Unknown')}")
                    else:
                        logger.warning(f"❌ 字串格式成分解析失敗: {ingredient[:50]}...")
        elif isinstance(field_data, str):
            # 字串格式的成分表，嘗試解析
            logger.info(f"🔍 成分表為字串格式，長度: {len(field_data)}")
            parsed_ingredients = self._parse_tab_separated_ingredients(field_data)
            if parsed_ingredients:
                ingredients.extend(parsed_ingredients)
                logger.info(f"✅ 解析Tab分隔成分表，共 {len(parsed_ingredients)} 個成分")
            else:
                logger.warning(f"❌ Tab分隔成分表解析失敗")
        else:
            logger.warning(f"❌ 成分表欄位不是列表或字串格式: {type(field_data)}")
        
        return ingredients
    
    def _parse_tab_separated_ingredients(self, ingredient_table_str: str) -> List[Dict[str, Any]]:
        """解析Tab分隔的成分表格式"""
        try:
            ingredients = []
            lines = ingredient_table_str.strip().split('\n')
            
            if len(lines) < 2:
                logger.warning("成分表格式不正確，至少需要標題行和數據行")
                return []
            
            # 第一行是標題行
            header_line = lines[0]
            headers = [h.strip() for h in header_line.split('\t')]
            logger.info(f"🔍 成分表標題: {headers}")
            
            # 處理數據行
            for i, line in enumerate(lines[1:], 1):
                if not line.strip():
                    continue
                
                parts = [p.strip() for p in line.split('\t')]
                if len(parts) < len(headers):
                    logger.warning(f"第 {i} 行數據不完整: {line}")
                    continue
                
                # 創建成分字典
                ingredient = {}
                for j, header in enumerate(headers):
                    if j < len(parts):
                        # 將標題映射到標準欄位名稱
                        if 'inci' in header.lower() or '名稱' in header:
                            ingredient['INCI名稱'] = parts[j]
                        elif 'cas' in header.lower() or '號碼' in header:
                            ingredient['CAS號碼'] = parts[j]
                        elif 'w/v' in header.lower() or '含量' in header or '%' in header:
                            ingredient['含量'] = parts[j]
                        elif '功能' in header.lower() or 'function' in header.lower():
                            ingredient['功能'] = parts[j]
                        else:
                            # 使用原始標題作為鍵名
                            ingredient[header] = parts[j]
                
                if ingredient:
                    ingredients.append(ingredient)
                    logger.debug(f"✅ 解析成分 {i}: {ingredient.get('INCI名稱', 'Unknown')}")
            
            logger.info(f"✅ 成功解析Tab分隔成分表，共 {len(ingredients)} 個成分")
            return ingredients
            
        except Exception as e:
            logger.error(f"解析Tab分隔成分表失敗: {e}")
            return []
    
    def _parse_markdown_ingredients_table(self, text: str) -> List[Dict[str, Any]]:
        """解析 Markdown 格式的成分表"""
        ingredients = []
        lines = text.split('\n')
        
        in_table = False
        headers = []
        
        for line in lines:
            line = line.strip()
            
            # 檢測表格開始（必須是表頭行）
            if '|' in line and ('成分' in line or 'inci' in line.lower() or '序號' in line) and not in_table:
                in_table = True
                # 解析表頭
                headers = [h.strip() for h in line.split('|') if h.strip()]
                continue
            
            # 檢測表格結束（分隔行）
            if in_table and line.startswith('|') and line.count('-') >= 3:
                continue
            
            # 解析表格行
            if in_table and line.startswith('|') and not line.startswith('|---'):
                # 分割表格行，保留空字串
                raw_cells = line.split('|')
                # 移除首尾的空字串，但保留中間的空字串
                if raw_cells and not raw_cells[0].strip():
                    raw_cells = raw_cells[1:]
                if raw_cells and not raw_cells[-1].strip():
                    raw_cells = raw_cells[:-1]
                cells = [c.strip() for c in raw_cells]
                if len(cells) >= 2:  # 至少要有成分名稱
                    ingredient = {}
                    for i, cell in enumerate(cells):
                        if i < len(headers):
                            header = headers[i].lower()
                            if '成分' in header or 'name' in header:
                                ingredient['INCI名稱'] = cell
                            elif 'cas' in header:
                                ingredient['CAS號碼'] = cell
                            elif '含量' in header or 'content' in header or 'w/v' in header:
                                ingredient['含量'] = cell
                            elif '功能' in header or 'function' in header or '說明' in header:
                                ingredient['功能'] = cell
                    
                    # 確保有成分名稱才添加
                    if ingredient.get('INCI名稱') and ingredient['INCI名稱'] != '成分名稱':
                        ingredients.append(ingredient)
        
        return ingredients
    
    def _replace_ingredients_table(self, table, ingredients: List[Dict[str, Any]]) -> None:
        """替換成分表表格內容"""
        if not ingredients:
            return
        
        # 保存第一行的格式作為模板
        format_template = None
        if len(table.rows) > 0:
            format_template = []
            for cell in table.rows[0].cells:
                cell_formats = []
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        cell_formats.append({
                            'font_name': run.font.name,
                            'font_size': run.font.size,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline
                        })
                format_template.append(cell_formats)
        
        # 清空現有表格內容（保留第一行標題）
        if len(table.rows) > 1:
            # 刪除除第一行外的所有行
            for i in range(len(table.rows) - 1, 0, -1):
                table._element.remove(table.rows[i]._element)
        
        # 添加新的成分資料行
        for ingredient in ingredients:
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) >= 4:
                # 替換內容並保持格式
                ingredient_data = [
                    ingredient.get('INCI名稱', ''),
                    ingredient.get('CAS號碼', '') or '-',
                    ingredient.get('含量', ''),
                    ingredient.get('功能', '')
                ]
                
                for i, (cell, content) in enumerate(zip(cells, ingredient_data)):
                    # 清空現有內容
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = ""
                    
                    # 添加新內容
                    if cell.paragraphs:
                        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
                        run.text = content
                        
                        # 應用格式模板（如果存在）
                        if format_template and i < len(format_template) and format_template[i]:
                            template = format_template[i][0]  # 使用第一個run的格式
                            if template['font_name']:
                                run.font.name = template['font_name']
                            if template['font_size']:
                                run.font.size = template['font_size']
                            if template['bold'] is not None:
                                run.bold = template['bold']
                            if template['italic'] is not None:
                                run.italic = template['italic']
                            if template['underline'] is not None:
                                run.underline = template['underline']
                        
                        # 添加淺藍色背景標記
                        try:
                            from docx.oxml import OxmlElement
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:val'), 'clear')
                            shading.set(qn('w:color'), 'auto')
                            shading.set(qn('w:fill'), 'B4D4E6')
                            run._element.get_or_add_rPr().append(shading)
                        except Exception as e:
                            logger.warning(f"添加背景標記失敗: {e}")
    
    def _replace_ingredients_tables_in_doc(self, doc_path: Path, data: Dict[str, Any]) -> None:
        """在Word文檔中替換所有成分表表格"""
        try:
            from docx import Document
            doc = Document(str(doc_path))
            
            ingredients = data.get('成分表', [])
            if not ingredients:
                logger.info("沒有成分表資料，跳過替換")
                return
            
            logger.info(f"開始替換成分表，共 {len(ingredients)} 個成分")
            
            # 檢查所有表格
            for table_idx, table in enumerate(doc.tables):
                if self._is_ingredients_table(table):
                    logger.info(f"替換表格 {table_idx} 的成分表")
                    self._replace_ingredients_table(table, ingredients)
            
            # 保存修改後的文檔
            doc.save(str(doc_path))
            logger.info("成分表替換完成")
            
        except Exception as e:
            logger.error(f"替換成分表失敗: {e}")
            # 不影響主要功能，只記錄錯誤
    
    def _highlight_inserted_data(self, output_path: Path, data: Dict[str, Any]) -> None:
        """為插入的資料添加淺藍底色標記"""
        try:
            from docx import Document
            from docx.shared import RGBColor
            from docx.oxml.shared import qn
            
            doc = Document(str(output_path))
            
            # 收集所有需要標記的資料值
            data_values = set()
            
            def collect_values(obj: Any, prefix: str = ""):
                """遞歸收集所有資料值"""
                if isinstance(obj, dict):
                    for k, v in obj.items():
                        collect_values(v, f"{prefix}.{k}" if prefix else k)
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        if isinstance(item, dict):
                            collect_values(item, f"{prefix}[{i}]")
                        else:
                            value = str(item).strip() if item is not None else ""
                            if value and len(value) > 2:  # 只收集長度大於2的值
                                data_values.add(value)
                else:
                    value = str(obj).strip() if obj is not None else ""
                    if value and len(value) > 2:  # 只收集長度大於2的值
                        data_values.add(value)
            
            collect_values(data)
            
            # 特別處理成分表資料
            ingredients_data = data.get('成分表', [])
            if isinstance(ingredients_data, list):
                for ingredient in ingredients_data:
                    if isinstance(ingredient, dict):
                        for field in ['INCI名稱', 'CAS號碼', '含量', '功能']:
                            value = str(ingredient.get(field, '')).strip()
                            if value and len(value) > 1:
                                data_values.add(value)
            
            # 淺藍色背景
            light_blue = RGBColor(173, 216, 230)  # 淺藍色
            
            # 處理段落中的資料
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run_text = run.text.strip()
                    if run_text and any(value in run_text for value in data_values):
                        # 設置淺藍色背景
                        run.font.highlight_color = None  # 清除現有高亮
                        # 使用shading代替highlight（更明顯）
                        try:
                            from docx.oxml import OxmlElement
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:val'), 'clear')
                            shading.set(qn('w:color'), 'auto')
                            shading.set(qn('w:fill'), 'B4D4E6')
                            run._element.get_or_add_rPr().append(shading)
                        except Exception as e:
                            logger.warning(f"添加背景標記失敗: {e}")
            
            # 處理表格中的資料
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run_text = run.text.strip()
                                if run_text and any(value in run_text for value in data_values):
                                    # 設置淺藍色背景
                                    try:
                                        from docx.oxml import OxmlElement
                                        shading = OxmlElement('w:shd')
                                        shading.set(qn('w:val'), 'clear')
                                        shading.set(qn('w:color'), 'auto')
                                        shading.set(qn('w:fill'), 'B4D4E6')
                                        run._element.get_or_add_rPr().append(shading)
                                    except Exception as e:
                                        logger.warning(f"添加背景標記失敗: {e}")
            
            # 保存修改後的文檔
            doc.save(str(output_path))
            logger.info(f"已為插入的資料添加淺藍底色標記，共標記 {len(data_values)} 個資料值")
            
        except Exception as e:
            logger.warning(f"添加淺藍底色標記失敗: {e}")
            # 不影響主要功能，只記錄警告
    
    def _transform_to_pdf(self, data: Dict[str, Any], 
                         template_path: Path, output_path: Path, allow_fallback: bool = False) -> bool:
        """轉換為PDF文檔"""
        try:
            # 先轉換為Word，再轉為PDF
            temp_docx = output_path.with_suffix('.docx')
            
            if self._transform_to_docx(data, template_path, temp_docx, allow_fallback):
                # 轉換為PDF
                return self._docx_to_pdf(temp_docx, output_path)
            else:
                return False
                
        except Exception as e:
            logger.error(f"PDF轉換失敗: {e}")
            # 直接拋出異常，不使用回退機制
            raise Exception(f"PDF轉換失敗: {e}")
    
    def _docx_to_pdf(self, docx_path: Path, pdf_path: Path) -> bool:
        """將Word文檔轉換為PDF"""
        try:
            # 嘗試使用docx2pdf
            try:
                from docx2pdf import convert
                convert(str(docx_path), str(pdf_path))
                return True
            except ImportError:
                logger.warning("docx2pdf未安裝，嘗試其他方法")
            
            # 嘗試使用LibreOffice
            try:
                import subprocess
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', str(pdf_path.parent),
                    str(docx_path)
                ], capture_output=True, text=True)
                
                if result.returncode == 0:
                    return True
                else:
                    logger.error(f"LibreOffice轉換失敗: {result.stderr}")
                    return False
                    
            except FileNotFoundError:
                logger.error("LibreOffice未安裝")
                return False
                
        except Exception as e:
            logger.error(f"PDF轉換失敗: {e}")
            # 直接拋出異常，不使用回退機制
            raise Exception(f"PDF轉換失敗: {e}")
    
    def _transform_to_json(self, data: Dict[str, Any], output_path: Path) -> bool:
        """轉換為JSON檔案"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"JSON檔案已生成: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"JSON轉換失敗: {e}")
            # 直接拋出異常，不使用回退機制
            raise Exception(f"JSON轉換失敗: {e}")
    
    def get_available_templates(self) -> List[Dict[str, str]]:
        """獲取可用的模板列表"""
        templates = []
        
        if self.template_dir.exists():
            for template_file in self.template_dir.iterdir():
                if template_file.is_file() and template_file.suffix.lower() in [".docx", ".doc", ".pdf"]:
                    templates.append({
                        "name": template_file.stem,
                        "filename": template_file.name,
                        "path": str(template_file),
                        "type": template_file.suffix.lower().lstrip(".")
                    })
        
        return templates
    
    def create_template(self, template_name: str, 
                       structure: Dict[str, Any]) -> bool:
        """
        創建新模板
        
        Args:
            template_name: 模板名稱
            structure: 文檔結構定義
            
        Returns:
            創建是否成功
        """
        try:
            # 這裡可以實現模板創建邏輯
            # 暫時返回False
            logger.info(f"模板創建功能待實現: {template_name}")
            return False
            
        except Exception as e:
            logger.error(f"模板創建失敗: {e}")
            return False


