#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
é€šç”¨æ–‡æª”æå–å™¨
æ”¯æ´å¤šç¨®æ–‡æª”æ ¼å¼çš„æ™ºèƒ½æå–
"""

import os
import sys
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
import pdfplumber
from docx import Document

from utils.ai_client import AIClient
from utils.file_handler import FileHandler
from .profile_manager import ProfileManager
from .prompt_parser import IntelligentPromptParser

logger = logging.getLogger(__name__)

class DocumentExtractor:
    """é€šç”¨æ–‡æª”æå–å™¨"""
    
    def __init__(self, profile_path: Optional[str] = None, work_id: Optional[str] = None, 
                 work_data: Optional[Dict[str, Any]] = None, ai_provider: Optional[str] = None,
                 ai_model: Optional[str] = None, user_prompt: Optional[str] = None):
        """
        åˆå§‹åŒ–æ–‡æª”æå–å™¨
        
        Args:
            profile_path: Profileæª”æ¡ˆè·¯å¾‘ï¼ˆèˆŠç‰ˆç›¸å®¹ï¼‰
            work_id: å·¥ä½œIDï¼ˆæ–°ç‰ˆåˆ†å±¤Profileï¼‰
            work_data: å·¥ä½œè³‡æ–™ï¼ˆæ–°ç‰ˆåˆ†å±¤Profileï¼‰
            ai_provider: AIæä¾›è€…ï¼ˆå„ªå…ˆä½¿ç”¨ï¼Œå¦å‰‡ä½¿ç”¨é è¨­ï¼‰
            ai_model: AIæ¨¡å‹ï¼ˆå„ªå…ˆä½¿ç”¨ï¼Œå¦å‰‡ä½¿ç”¨é è¨­ï¼‰
            user_prompt: ä½¿ç”¨è€…æŒ‡å®šçš„æç¤ºè©
        """
        self.profile_manager = ProfileManager()
        self.file_handler = FileHandler()
        self.prompt_parser = IntelligentPromptParser()
        
        # è¨­å®šAIæä¾›è€…å’Œæ¨¡å‹ï¼ˆå„ªå…ˆä½¿ç”¨å‚³å…¥åƒæ•¸ï¼Œå¦å‰‡ä½¿ç”¨é è¨­ï¼‰
        self.ai_provider = ai_provider
        self.ai_model = ai_model
        
        # ä¿å­˜ä½¿ç”¨è€…æç¤ºè©
        self.user_prompt = user_prompt
        
        # å»¶é²åˆå§‹åŒ–AIClientï¼Œé¿å…åœ¨æ²’æœ‰APIé‡‘é‘°æ™‚å‡ºéŒ¯
        self.ai_client = None
        
        # å¦‚æœæŒ‡å®šäº†AIè¨­å®šï¼Œè¨˜éŒ„æ—¥èªŒ
        if ai_provider and ai_model:
            logger.info(f"DocumentExtractor using specified AI settings: {ai_provider}/{ai_model}")
        else:
            logger.info(f"DocumentExtractor using default AI settings")
        
        # è¼‰å…¥Profileï¼ˆåš´æ ¼æŒ‰ç…§ä½¿ç”¨è€…æŒ‡å®šï¼‰
        if work_id and work_data:
            # æ–°ç‰ˆï¼šä½¿ç”¨åˆ†å±¤Profile
            self.profile = self.profile_manager.load_work_profile(work_id, work_data)
            logger.info(f"Document extractor initialized with hierarchical profile: {work_id}")
        elif profile_path:
            # èˆŠç‰ˆï¼šç›´æ¥è¼‰å…¥Profileæª”æ¡ˆ
            self.profile = self.profile_manager.load_profile(profile_path)
            logger.info(f"Document extractor initialized with profile: {self.profile.get('name', 'default')}")
        else:
            # å¦‚æœæ²’æœ‰æŒ‡å®šProfileï¼Œç›´æ¥å ±éŒ¯
            raise ValueError("å¿…é ˆæŒ‡å®šProfileè·¯å¾‘æˆ–å·¥ä½œè³‡æ–™ã€‚ç³»çµ±ä¸æœƒä½¿ç”¨é è¨­Profileã€‚")
    
    def _get_ai_client(self):
        """å»¶é²åˆå§‹åŒ–AIClient"""
        if self.ai_client is None:
            try:
                from utils.multi_ai_client import MultiAIClient
                from utils.settings_manager import SettingsManager
                
                # ç²å–è¨­å®š
                settings_manager = SettingsManager()
                settings = settings_manager.get_all_settings()
                
                # å¦‚æœæŒ‡å®šäº†AIè¨­å®šï¼Œè¦†è“‹é è¨­è¨­å®š
                if self.ai_provider and self.ai_model:
                    settings['ai_provider'] = self.ai_provider
                    settings['ai_model'] = self.ai_model
                    logger.info(f"Using specified AI settings: {self.ai_provider}/{self.ai_model}")
                else:
                    logger.info("Using default AI settings")
                
                self.ai_client = MultiAIClient(settings)
                # å¾MultiAIClientç²å–å¯¦éš›çš„AIæä¾›è€…å’Œæ¨¡å‹
                self.ai_provider = self.ai_client.current_provider
                self.ai_model = self.ai_client.current_model
                logger.info("MultiAIClient lazy initialization successful")
            except Exception as e:
                logger.warning(f"MultiAIClient initialization failed: {e}")
                # å‰µå»ºä¸€å€‹å‡çš„AIClientï¼Œé¿å…å¾ŒçºŒéŒ¯èª¤
                self.ai_client = type('MockAIClient', (), {
                    'extract_data': lambda *args, **kwargs: {"error": "AIå®¢æˆ¶ç«¯æœªæ­£ç¢ºåˆå§‹åŒ–"},
                    'generate_content': lambda *args, **kwargs: "AIå®¢æˆ¶ç«¯æœªæ­£ç¢ºåˆå§‹åŒ–"
                })()
        return self.ai_client
    
    def extract(self, file_path: Union[str, Path], output_format: str = "json", 
                user_prompt: Optional[str] = None, selected_pages: List[int] = None) -> Dict[str, Any]:
        """
        æå–æ–‡æª”çµæ§‹åŒ–è³‡æ–™
        
        Args:
            file_path: æ–‡æª”æª”æ¡ˆè·¯å¾‘
            output_format: è¼¸å‡ºæ ¼å¼ (json, dict)
            user_prompt: ä½¿ç”¨è€…è‡ªå®šç¾©æç¤ºè©
            selected_pages: ç”¨æˆ¶é¸æ“‡çš„é é¢åˆ—è¡¨
            
        Returns:
            æå–çš„çµæ§‹åŒ–è³‡æ–™
        """
        try:
            # å­˜å„²ä½¿ç”¨è€…æç¤ºè©åˆ°å¯¦ä¾‹è®Šæ•¸
            self.user_prompt = user_prompt
            
            # è®€å–æ–‡æª”å…§å®¹
            content = self._read_document(file_path)
            if not content:
                raise ValueError(f"ç„¡æ³•è®€å–æ–‡æª”: {file_path}")
            
            # å¦‚æœæœ‰ä½¿ç”¨è€…æç¤ºè©ï¼Œä½¿ç”¨æ™ºèƒ½è§£æ
            if user_prompt:
                structured_data = self._extract_with_intelligent_prompt(
                    content, file_path, user_prompt, selected_pages
                )
            else:
                # æ™ºèƒ½åˆ†é è™•ç†
                if self.profile.get('use_page_extraction', True):
                    structured_data = self._extract_by_pages(content, file_path, selected_pages)
                else:
                    structured_data = self._extract_whole_document(content)
            
            # å¾Œè™•ç†
            structured_data = self._post_process(structured_data)
            
            # æ ¼å¼åŒ–è¼¸å‡º
            if output_format == "json":
                return structured_data
            elif output_format == "dict":
                return structured_data
            else:
                raise ValueError(f"ä¸æ”¯æ´çš„è¼¸å‡ºæ ¼å¼: {output_format}")
                
        except Exception as e:
            logger.error(f"Document extraction failed: {e}")
            raise
    
    def _extract_with_intelligent_prompt(self, content: str, file_path: Path, 
                                       user_prompt: str, selected_pages: List[int] = None) -> Dict[str, Any]:
        """
        ä½¿ç”¨æ™ºèƒ½æç¤ºè©æå–æ–‡æª”
        
        Args:
            content: æ–‡æª”å…§å®¹
            file_path: æ–‡æª”è·¯å¾‘
            user_prompt: ä½¿ç”¨è€…æç¤ºè©
            
        Returns:
            æå–çš„çµæ§‹åŒ–è³‡æ–™
        """
        try:
            # è¨­ç½®ä½¿ç”¨è€…æç¤ºè©åˆ°å¯¦ä¾‹è®Šæ•¸
            self.user_prompt = user_prompt
            
            # æ ¹æ“šæ˜¯å¦æœ‰é¸å®šé é¢æ±ºå®šä½¿ç”¨å“ªç¨®æ–¹æ³•
            if selected_pages:
                logger.info("ğŸ” ä½¿ç”¨åˆ†é æå–æ–¹æ³•ï¼Œç¢ºä¿ä½¿ç”¨è€…æç¤ºè©ä¸è¢«è¦†è“‹")
                return self._extract_by_pages(content, file_path, selected_pages)
            else:
                logger.info("ğŸ” ä½¿ç”¨æ•´ä»½æ–‡æª”æå–æ–¹æ³•ï¼Œç¢ºä¿ä½¿ç”¨è€…æç¤ºè©ä¸è¢«è¦†è“‹")
                return self._extract_whole_document(content)
            
        except Exception as e:
            logger.error(f"æ™ºèƒ½æç¤ºè©æå–å¤±æ•—: {e}")
            # ç›´æ¥æ‹‹å‡ºç•°å¸¸ï¼Œä¸ä½¿ç”¨å›é€€æ©Ÿåˆ¶
            raise Exception(f"æ™ºèƒ½æç¤ºè©æå–å¤±æ•—ï¼Œè«‹æª¢æŸ¥æç¤ºè©å…§å®¹æˆ–AIæ¨¡å‹è¨­å®š: {e}")
    
    def _get_template_info(self) -> str:
        """ç²å–æ¨¡æ¿è³‡è¨Š"""
        # é€™è£¡å¯ä»¥å¾Profileä¸­ç²å–æ¨¡æ¿è³‡è¨Š
        template_info = "ä½¿ç”¨é è¨­æ¨¡æ¿æ ¼å¼"
        
        if 'template_info' in self.profile:
            template_info = self.profile['template_info']
        
        return template_info
    
    def _read_document(self, file_path: Union[str, Path]) -> str:
        """è®€å–æ–‡æª”å…§å®¹"""
        file_path = Path(file_path)
        file_type = file_path.suffix.lower()
        
        if file_type == '.pdf':
            return self._read_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            return self._read_word(file_path)
        elif file_type == '.txt':
            return self._read_text(file_path)
        else:
            raise ValueError(f"ä¸æ”¯æ´çš„æª”æ¡ˆæ ¼å¼: {file_type}")
    
    def _read_pdf(self, file_path: Path) -> str:
        """è®€å–PDFæª”æ¡ˆ"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"PDFè®€å–å¤±æ•—: {e}")
            return ""
    
    def _read_pdf_by_pages(self, file_path: Path, selected_pages: List[int] = None) -> List[Dict[str, Any]]:
        """æŒ‰é é¢è®€å–PDFæª”æ¡ˆï¼Œåªè®€å–é¸å®šçš„é é¢"""
        try:
            pages_data = []
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"ğŸ” PDFç¸½é æ•¸: {len(pdf.pages)}")
                
                # åªè®€å–é¸å®šçš„é é¢
                pages_to_read = selected_pages if selected_pages else range(1, len(pdf.pages) + 1)
                
                for page_num in pages_to_read:
                    if page_num > len(pdf.pages):
                        logger.warning(f"ğŸ” ç¬¬ {page_num} é è¶…å‡ºç¯„åœï¼Œè·³é")
                        continue
                        
                    page = pdf.pages[page_num - 1]  # é é¢ç´¢å¼•å¾0é–‹å§‹
                    # å„ªå…ˆæå–è¡¨æ ¼å…§å®¹
                    tables = page.extract_tables()
                    page_text = page.extract_text()
                    
                    # å¦‚æœæœ‰è¡¨æ ¼ï¼Œå°‡è¡¨æ ¼å…§å®¹è½‰æ›ç‚ºæ–‡æœ¬
                    table_text = ""
                    if tables:
                        for table in tables:
                            for row in table:
                                if row:  # ç¢ºä¿è¡Œä¸ç‚ºç©º
                                    table_text += " | ".join([cell or "" for cell in row]) + "\n"
                    
                    # åˆä½µæ–‡æœ¬å’Œè¡¨æ ¼å…§å®¹
                    combined_content = ""
                    if page_text:
                        combined_content += page_text + "\n"
                    if table_text:
                        combined_content += "\n=== è¡¨æ ¼å…§å®¹ ===\n" + table_text
                    
                    
                    if combined_content.strip():
                        # æå–åœ–ç‰‡ä½ç½®ï¼ˆå¦‚æœæœ‰ï¼‰
                        images = []
                        if hasattr(page, 'images'):
                            images = page.images
                        
                        
                        pages_data.append({
                            'page_number': page_num,
                            'content': combined_content,
                            'tables': tables or [],
                            'images': images or [],
                            'paragraphs': self._split_into_paragraphs(combined_content)
                        })
                    else:
                        logger.warning(f"ğŸ” ç¬¬ {page_num} é æ²’æœ‰æå–åˆ°ä»»ä½•å…§å®¹")
            logger.info(f"ğŸ” æˆåŠŸè®€å– {len(pages_data)} å€‹æœ‰å…§å®¹çš„é é¢")
            return pages_data
        except Exception as e:
            logger.error(f"PDFåˆ†é è®€å–å¤±æ•—: {e}")
            return []
    
    def _split_into_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """å°‡æ–‡å­—åˆ†å‰²ç‚ºæ®µè½"""
        paragraphs = []
        lines = text.split('\n')
        current_para = []
        
        for line_num, line in enumerate(lines, 1):
            line = line.strip()
            if line:
                current_para.append(line)
            else:
                if current_para:
                    paragraphs.append({
                        'paragraph_number': len(paragraphs) + 1,
                        'line_number': line_num - len(current_para),
                        'content': '\n'.join(current_para)
                    })
                    current_para = []
        
        # è™•ç†æœ€å¾Œä¸€å€‹æ®µè½
        if current_para:
            paragraphs.append({
                'paragraph_number': len(paragraphs) + 1,
                'line_number': len(lines) - len(current_para),
                'content': '\n'.join(current_para)
            })
        
        return paragraphs
    
    def _read_word(self, file_path: Path) -> str:
        """è®€å–Wordæª”æ¡ˆ"""
        try:
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Wordè®€å–å¤±æ•—: {e}")
            return ""
    
    def _read_docx_by_pages(self, file_path: Path, selected_pages: List[int] = None) -> List[Dict[str, Any]]:
        """æŒ‰é é¢è®€å–DOCXæª”æ¡ˆï¼Œåªè®€å–é¸å®šçš„é é¢ï¼ˆæ¨¡æ“¬åˆ†é ï¼‰"""
        try:
            doc = Document(file_path)
            pages_data = []
            
            # ç”±æ–¼DOCXæ²’æœ‰æ˜ç¢ºçš„é é¢æ¦‚å¿µï¼Œæˆ‘å€‘æŒ‰æ®µè½åˆ†çµ„ä¾†æ¨¡æ“¬åˆ†é 
            current_page = 1
            current_content = []
            paragraphs_per_page = 10  # æ¯é å¤§ç´„10å€‹æ®µè½
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    current_content.append(paragraph.text)
                
                # æ¯10å€‹æ®µè½æˆ–é‡åˆ°åˆ†é ç¬¦æ™‚å‰µå»ºæ–°é é¢
                if (i + 1) % paragraphs_per_page == 0 or i == len(doc.paragraphs) - 1:
                    if current_content:
                        page_text = '\n'.join(current_content)
                        pages_data.append({
                            'page_number': current_page,
                            'content': page_text,
                            'tables': [],  # DOCXè¡¨æ ¼è™•ç†è¼ƒè¤‡é›œï¼Œæš«æ™‚ç•™ç©º
                            'images': [],  # åœ–ç‰‡è™•ç†è¼ƒè¤‡é›œï¼Œæš«æ™‚ç•™ç©º
                            'paragraphs': self._split_into_paragraphs(page_text)
                        })
                        current_page += 1
                        current_content = []
            
            # å¦‚æœæ²’æœ‰å…§å®¹ï¼Œè‡³å°‘è¿”å›ä¸€é 
            if not pages_data:
                pages_data.append({
                    'page_number': 1,
                    'content': 'æª”æ¡ˆå…§å®¹ç‚ºç©º',
                    'tables': [],
                    'images': [],
                    'paragraphs': []
                })
            
            return pages_data
        except Exception as e:
            logger.error(f"DOCXåˆ†é è®€å–å¤±æ•—: {e}")
            return []
    
    def _read_text(self, file_path: Path) -> str:
        """è®€å–æ–‡å­—æª”æ¡ˆ"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"æ–‡å­—æª”æ¡ˆè®€å–å¤±æ•—: {e}")
            return ""
    
    def _extract_by_pages(self, content: str, file_path: Path, selected_pages: List[int] = None) -> Dict[str, Any]:
        """åˆ†é æå–ï¼ˆæ™ºèƒ½è­˜åˆ¥é é¢é¡å‹ï¼‰"""
        try:
            # ç¢ºä¿ user_prompt å·²è¨­ç½®ï¼ˆå¾èª¿ç”¨è€…å‚³éï¼‰
            if not hasattr(self, 'user_prompt') or not self.user_prompt:
                raise ValueError("åˆ†é æå–éœ€è¦ä½¿ç”¨è€…æç¤ºè©ï¼Œä½†æœªæä¾›")
            
            # æª¢æŸ¥æ˜¯å¦å•Ÿç”¨åˆ†é æå–
            # å¦‚æœç”¨æˆ¶é¸æ“‡äº†é é¢ï¼Œå¼·åˆ¶ä½¿ç”¨åˆ†é æå–
            use_page_extraction = self.profile.get('use_page_extraction', True)
            if not use_page_extraction and not selected_pages:
                logger.info("åˆ†é æå–æœªå•Ÿç”¨ä¸”æœªé¸æ“‡é é¢ï¼Œä½¿ç”¨æ•´ä»½æ–‡æª”æå–")
                return self._extract_whole_document(content)
            
            # è®€å–é é¢ï¼šæœ‰é¸å®šå°±æŒ‰é¸æ“‡ï¼Œæ²’é¸å®šå°±å…¨éƒ¨
            file_type = file_path.suffix.lower()
            if file_type == '.pdf':
                important_pages = self._read_pdf_by_pages(file_path, selected_pages)
            elif file_type in ['.docx', '.doc']:
                important_pages = self._read_docx_by_pages(file_path, selected_pages)
            else:
                logger.warning(f"ä¸æ”¯æ´åˆ†é è™•ç†çš„æª”æ¡ˆæ ¼å¼: {file_type}")
                return self._extract_whole_document(content)
            
            if not important_pages:
                logger.error("ç„¡æ³•è®€å–é é¢ï¼Œè«‹æª¢æŸ¥æ–‡æª”æ ¼å¼æˆ–æª”æ¡ˆå®Œæ•´æ€§")
                raise Exception("ç„¡æ³•è®€å–æ–‡æª”é é¢ï¼Œè«‹æª¢æŸ¥æ–‡æª”æ ¼å¼æ˜¯å¦æ­£ç¢º")
            
            if selected_pages:
                logger.info(f"ç”¨æˆ¶é¸æ“‡äº† {len(important_pages)} å€‹é é¢é€²è¡Œè™•ç†: {selected_pages}")
            else:
                logger.info(f"ç”¨æˆ¶æœªé¸æ“‡é é¢ï¼Œä½¿ç”¨å…¨éƒ¨ {len(important_pages)} å€‹é é¢")
            
            # è®€å–è™•ç†è¨­å®š
            try:
                from utils.settings_manager import SettingsManager
                proc_cfg = SettingsManager().get_processing_config()
            except Exception:
                proc_cfg = {"page_selection": {"mode": "manual", "truncate_chars": 0}, "merge_strategy": {"ingredients": "copy_as_is", "strings": "append", "dict": "prefer_non_empty"}, "prompt": {"inject_page_meta": False, "simplify_profile": False}}

            # important_pages å·²ç¶“åœ¨è®€å–é é¢æ™‚å®šç¾©äº†ï¼Œé€™è£¡åªéœ€è¦è¨˜éŒ„æ—¥èªŒ
            if selected_pages:
                logger.info(f"ç”¨æˆ¶é¸æ“‡äº† {len(important_pages)} å€‹é é¢é€²è¡Œè™•ç†: {selected_pages}")
            else:
                logger.info(f"ç”¨æˆ¶æœªé¸æ“‡é é¢ï¼Œä½¿ç”¨å…¨éƒ¨ {len(important_pages)} å€‹é é¢")
            
            # ä¼°ç®—æˆæœ¬
            estimated_cost = self._estimate_processing_cost(important_pages)
            logger.info(f"é ä¼°è™•ç†æˆæœ¬: ${estimated_cost:.2f}")
            
            # é å…ˆæª¢æŸ¥æ˜¯å¦éœ€è¦åˆ†æ®µè™•ç†ï¼ˆåŸºæ–¼ä¸åŒAIæ¨¡å‹çš„ç‰¹æ€§ï¼‰
            # ç¢ºä¿AIå®¢æˆ¶ç«¯å·²åˆå§‹åŒ–ï¼Œä»¥ä¾¿ç²å–æ­£ç¢ºçš„ai_provider
            self._get_ai_client()
            ai_provider = getattr(self, 'ai_provider', '')
            total_content_length = sum(len(page_data.get('content', '')) for page_data in important_pages)
            should_chunk = self._should_use_chunking(ai_provider, total_content_length)
            
            if should_chunk:
                logger.info(f"æª¢æ¸¬åˆ°é•·æ–‡æª”ï¼ˆ{total_content_length} å­—ç¬¦ï¼‰ï¼Œä½¿ç”¨åˆ†æ®µè™•ç†ç­–ç•¥")
                # ä½¿ç”¨ç°¡åŒ–çš„åˆä½µé‚è¼¯å‰µå»ºå…§å®¹ç”¨æ–¼åˆ†æ®µè™•ç†
                merged_content = self._simple_merge_pages(important_pages, file_type)
                merged_result = self._extract_document_in_chunks(merged_content)
                merged_result['_processed_pages'] = [p['page_number'] for p in important_pages]
            else:
                # ç°¡åŒ–é é¢åˆä½µç­–ç•¥
                logger.info(f"ä½¿ç”¨ç°¡åŒ–é é¢åˆä½µç­–ç•¥ï¼Œåˆä½µ {len(important_pages)} å€‹é é¢")
                
                # ä½¿ç”¨ç°¡åŒ–çš„åˆä½µé‚è¼¯
                merged_content = self._simple_merge_pages(important_pages, file_type)
                
                # æ§‹å»ºæç¤ºè©
                merged_prompt = self._build_extraction_prompt(merged_content)
                
                try:
                    # èª¿ç”¨AI APIè™•ç†åˆä½µå…§å®¹
                    ai_client = self._get_ai_client()
                    merged_response = ai_client.extract_data(merged_prompt)
                    merged_result = self._parse_ai_response(merged_response)
                    merged_result['_raw_response'] = merged_response
                    merged_result['_processed_pages'] = [p['page_number'] for p in important_pages]
                    
                except Exception as e:
                    logger.error(f"é é¢è™•ç†å¤±æ•—: {e}")
                    raise e
            
            # å¾Œè™•ç†åˆä½µçµæœ
            return self._post_process_results(merged_result)
            
        except Exception as e:
            logger.error(f"åˆ†é æå–å¤±æ•—: {e}")
            # ç›´æ¥æ‹‹å‡ºç•°å¸¸ï¼Œä¸ä½¿ç”¨å›é€€æ©Ÿåˆ¶
            raise Exception(f"åˆ†é æå–å¤±æ•—ï¼Œè«‹æª¢æŸ¥æ–‡æª”æ ¼å¼æˆ–é é¢é¸æ“‡: {e}")
    
    def _extract_whole_document(self, content: str) -> Dict[str, Any]:
        """æ•´ä»½æ–‡æª”æå–"""
        # æª¢æŸ¥æ˜¯å¦éœ€è¦åˆ†æ®µè™•ç†ï¼ˆåŸºæ–¼ä¸åŒAIæ¨¡å‹çš„ç‰¹æ€§ï¼‰
        # ç¢ºä¿AIå®¢æˆ¶ç«¯å·²åˆå§‹åŒ–ï¼Œä»¥ä¾¿ç²å–æ­£ç¢ºçš„ai_provider
        self._get_ai_client()
        ai_provider = getattr(self, 'ai_provider', '')
        should_chunk = self._should_use_chunking(ai_provider, len(content))
        if should_chunk:
            logger.info(f"æª¢æ¸¬åˆ°é•·æ–‡æª”ï¼ˆ{len(content)} å­—ç¬¦ï¼‰ï¼Œä½¿ç”¨åˆ†æ®µè™•ç†ç­–ç•¥")
            return self._extract_document_in_chunks(content)
        
        # æ§‹å»ºæç¤ºè©
        prompt = self._build_extraction_prompt(content)
        
        # èª¿ç”¨AI API
        ai_client = self._get_ai_client()
        response = ai_client.extract_data(prompt)
        
        # è§£æå›æ‡‰
        structured_data = self._parse_ai_response(response)
        structured_data['_raw_response'] = response  # ä¿å­˜åŸå§‹éŸ¿æ‡‰
        return structured_data
    
    def _extract_document_in_chunks(self, content: str, chunk_size: int = 40000) -> Dict[str, Any]:
        """åˆ†æ®µè™•ç†é•·æ–‡æª”"""
        try:
            # å°‡æ–‡æª”åˆ†æˆå¤šå€‹æ®µè½
            chunks = self._split_content_into_chunks(content, chunk_size)
            logger.info(f"æ–‡æª”å·²åˆ†æˆ {len(chunks)} å€‹æ®µè½é€²è¡Œè™•ç†")
            
            # è™•ç†æ¯å€‹æ®µè½
            all_results = []
            successful_chunks = 0
            
            for i, chunk in enumerate(chunks):
                logger.info(f"æ­£åœ¨è™•ç†ç¬¬ {i+1}/{len(chunks)} å€‹æ®µè½")
                try:
                    # æ§‹å»ºæ®µè½æç¤ºè©
                    prompt = self._build_chunk_extraction_prompt(chunk, i+1, len(chunks))
                    
                    # èª¿ç”¨AI API
                    ai_client = self._get_ai_client()
                    response = ai_client.extract_data(prompt)
                    
                    # è§£æå›æ‡‰
                    chunk_result = self._parse_ai_response(response)
                    # ä¿å­˜åŸå§‹éŸ¿æ‡‰
                    chunk_result['_raw_response'] = response
                    all_results.append(chunk_result)
                    successful_chunks += 1
                    logger.info(f"ç¬¬ {i+1} å€‹æ®µè½è™•ç†æˆåŠŸ")
                    
                except Exception as e:
                    logger.warning(f"ç¬¬ {i+1} å€‹æ®µè½è™•ç†å¤±æ•—: {e}")
                    # å‰µå»ºç©ºçš„çµæœä»¥ä¿æŒçµæ§‹
                    empty_result = {
                        '_raw_response': '',
                        '_chunk_index': i,
                        '_chunk_error': str(e)
                    }
                    all_results.append(empty_result)
                    continue
            
            # æª¢æŸ¥æ˜¯å¦æœ‰æˆåŠŸçš„æ®µè½
            if successful_chunks == 0:
                logger.error("æ‰€æœ‰æ®µè½è™•ç†éƒ½å¤±æ•—äº†")
                raise Exception("æ‰€æœ‰æ®µè½è™•ç†éƒ½å¤±æ•—äº†")
            elif successful_chunks < len(chunks):
                logger.warning(f"åªæœ‰ {successful_chunks}/{len(chunks)} å€‹æ®µè½è™•ç†æˆåŠŸ")
            
            # åˆä½µæ‰€æœ‰çµæœ
            return self._merge_chunk_results(all_results)
                
        except Exception as e:
            logger.error(f"åˆ†æ®µè™•ç†å¤±æ•—: {e}")
            raise
    
    def _split_content_into_chunks(self, content: str, chunk_size: int) -> List[str]:
        """å°‡å…§å®¹åˆ†æˆå¤šå€‹æ®µè½ï¼Œæ™ºèƒ½é¿å…åˆ†å‰²æˆåˆ†è¡¨"""
        # æª¢æ¸¬æˆåˆ†è¡¨å€åŸŸ
        lines = content.split('\n')
        ingredient_section_start = -1
        ingredient_section_end = -1
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            # æª¢æ¸¬æˆåˆ†è¡¨é–‹å§‹ - æ”¾å¯¬æª¢æ¸¬æ¢ä»¶
            if ('inci' in line_lower or 'composizione' in line_lower or 'ingredient' in line_lower or 'æˆåˆ†' in line_lower):
                ingredient_section_start = i
            # æª¢æ¸¬æˆåˆ†è¡¨çµæŸï¼ˆé‡åˆ°ç©ºè¡Œæˆ–æ–°ç« ç¯€ï¼‰
            elif ingredient_section_start >= 0 and (line.strip() == '' or line.startswith('PIF:') or line.startswith('PARTE')):
                ingredient_section_end = i
                break
        
        # å¦‚æœæ‰¾åˆ°æˆåˆ†è¡¨å€åŸŸï¼Œç¢ºä¿ä¸è¢«åˆ†å‰²
        if ingredient_section_start >= 0:
            logger.info(f"æª¢æ¸¬åˆ°æˆåˆ†è¡¨å€åŸŸï¼šè¡Œ {ingredient_section_start} åˆ° {ingredient_section_end}")
            
            # å°‡æˆåˆ†è¡¨å€åŸŸä½œç‚ºä¸€å€‹æ•´é«”æ®µè½
            ingredient_lines = lines[ingredient_section_start:ingredient_section_end] if ingredient_section_end > 0 else lines[ingredient_section_start:]
            ingredient_chunk = '\n'.join(ingredient_lines)
            
            # è™•ç†æˆåˆ†è¡¨ä¹‹å‰çš„å…§å®¹
            before_lines = lines[:ingredient_section_start]
            before_chunk = '\n'.join(before_lines)
            
            # è™•ç†æˆåˆ†è¡¨ä¹‹å¾Œçš„å…§å®¹
            after_lines = lines[ingredient_section_end:] if ingredient_section_end > 0 else []
            after_chunk = '\n'.join(after_lines)
            
            # åˆ†åˆ¥è™•ç†å„å€‹éƒ¨åˆ†
            chunks = []
            if before_chunk.strip():
                chunks.extend(self._split_content_simple(before_chunk, chunk_size))
            
            if ingredient_chunk.strip():
                chunks.append(ingredient_chunk)
            
            if after_chunk.strip():
                chunks.extend(self._split_content_simple(after_chunk, chunk_size))
            
            return chunks
        
        # å¦‚æœæ²’æœ‰æª¢æ¸¬åˆ°æˆåˆ†è¡¨ï¼Œä½¿ç”¨ç°¡å–®åˆ†å‰²
        return self._split_content_simple(content, chunk_size)
    
    def _split_content_simple(self, content: str, chunk_size: int) -> List[str]:
        """ç°¡å–®çš„å…§å®¹åˆ†å‰²"""
        chunks = []
        start = 0
        
        while start < len(content):
            end = start + chunk_size
            
            # å¦‚æœä¸æ˜¯æœ€å¾Œä¸€å€‹æ®µè½ï¼Œå˜—è©¦åœ¨å¥è™Ÿæˆ–æ›è¡Œè™•åˆ†å‰²
            if end < len(content):
                # å°‹æ‰¾æœ€è¿‘çš„å¥è™Ÿæˆ–æ›è¡Œ
                for i in range(end, start + chunk_size - 1000, -1):
                    if content[i] in ['ã€‚', '\n', '\r']:
                        end = i + 1
                        break
            
            chunk = content[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end
        
        return chunks
    
    def _build_extraction_prompt(self, content: str) -> str:
        """æ§‹å»ºæå–æç¤ºè©"""
        # åš´æ ¼ä½¿ç”¨ä½¿ç”¨è€…æŒ‡å®šçš„æç¤ºè©
        if not hasattr(self, 'user_prompt') or not self.user_prompt:
            raise ValueError("å¿…é ˆæä¾›ä½¿ç”¨è€…æç¤ºè©ã€‚ç³»çµ±ä¸æœƒä½¿ç”¨é è¨­æç¤ºè©ã€‚")
        
        prompt_template = self.user_prompt
        
        # ç¢ºä¿æç¤ºè©æ¨¡æ¿æœ‰æ­£ç¢ºçš„çµæ§‹åˆ†éš”
        if "{{content}}" in prompt_template:
            # æ›¿æ›ä½”ä½ç¬¦ï¼Œä½†ä¿æŒcontentåœ¨æœ€å¾Œï¼ˆé¿å…AIæ··æ·†ï¼‰
            prompt = prompt_template.replace("{{content}}", "[æ–‡æª”å…§å®¹å°‡åœ¨ä¸‹æ–¹æä¾›]")
            
            # è™•ç†Profileä½”ä½ç¬¦ï¼ˆè€ƒæ…®settingsè¨­å®šï¼‰
            if "{{profile}}" in prompt:
                if self._should_attach_profile(prompt_template):
                    prompt = prompt.replace("{{profile}}", f"\n\n=== Profileé…ç½® ===\n{json.dumps(self.profile, ensure_ascii=False, indent=2)}\n\n=== Profileé…ç½®çµæŸ ===")
                else:
                    prompt = prompt.replace("{{profile}}", "[Profileé…ç½®å·²çœç•¥]")
            
            # åœ¨æœ€å¾Œé™„åŠ å¯¦éš›å…§å®¹
            prompt = prompt + f"\n\n=== éœ€è¦è™•ç†çš„æ–‡æª”å…§å®¹ ===\n{content}\n\n=== æ–‡æª”å…§å®¹çµæŸ ==="
        else:
            # å¦‚æœæ²’æœ‰æ¨¡æ¿è®Šæ•¸ï¼Œåœ¨æœ€å¾Œé™„åŠ å…§å®¹
            prompt = prompt_template + f"\n\n=== éœ€è¦è™•ç†çš„æ–‡æª”å…§å®¹ ===\n{content}\n\n=== æ–‡æª”å…§å®¹çµæŸ ==="
        
        # æ·»åŠ èª¿è©¦æ—¥èªŒ
        logger.info(f"ğŸ” æ§‹å»ºçš„æç¤ºè©é•·åº¦: {len(prompt)} å­—ç¬¦")
        logger.info(f"ğŸ” æç¤ºè©å‰500å­—ç¬¦: {prompt[:500]}...")
        
        # æª¢æŸ¥Profileå…§å®¹
        logger.info(f"ğŸ” Profileå…§å®¹: {json.dumps(self.profile, ensure_ascii=False, indent=2)[:200]}...")
        
        return prompt
    
    def _build_chunk_extraction_prompt(self, chunk: str, chunk_num: int, total_chunks: int) -> str:
        """æ§‹å»ºæ®µè½æå–æç¤ºè©"""
        return self._build_extraction_prompt(chunk)
    
    def _build_extraction_prompt_with_filename(self, filename: str, file_path: Path = None, content: str = None) -> str:
        """æ§‹å»ºæå–æç¤ºè© - è§£æ±ºAIæ··æ·†å•é¡Œ"""
        # åš´æ ¼ä½¿ç”¨ä½¿ç”¨è€…æŒ‡å®šçš„æç¤ºè©
        if not hasattr(self, 'user_prompt') or not self.user_prompt or self.user_prompt.strip() == '':
            raise ValueError("å¿…é ˆæä¾›ä½¿ç”¨è€…æç¤ºè©ã€‚ç³»çµ±ä¸æœƒä½¿ç”¨é è¨­æç¤ºè©ã€‚")
        
        prompt_template = self.user_prompt
        
        # ç­–ç•¥ï¼šä½¿ç”¨æˆåŠŸé©—è­‰çš„åˆ†éš”æ¨™è¨˜ï¼ˆåŸºæ–¼2025-10-01æˆåŠŸå¯¦ç¾ï¼‰
        if "{{content}}" in prompt_template:
            # æ›¿æ›ä½”ä½ç¬¦ï¼Œä½¿ç”¨æˆåŠŸé©—è­‰çš„åˆ†éš”æ¨™è¨˜
            prompt = prompt_template.replace("{{content}}", 
                f"\n\n=== æ–‡æª”å…§å®¹ ===\n"
                f"æª”æ¡ˆåç¨±ï¼š{filename}\n"
                f"è³‡æ–™å…§å®¹ï¼š\n{content}\n"
                f"=== æ–‡æª”å…§å®¹çµæŸ ===\n")
        else:
            # æ²’æœ‰ä½”ä½ç¬¦æ™‚ï¼Œåœ¨æœ€å¾Œé™„åŠ å…§å®¹ï¼Œä½¿ç”¨æˆåŠŸé©—è­‰çš„åˆ†éš”æ¨™è¨˜
            prompt = prompt_template + f"\n\n=== æ–‡æª”å…§å®¹ ===\næª”æ¡ˆåç¨±ï¼š{filename}\nè³‡æ–™å…§å®¹ï¼š\n{content}\n=== æ–‡æª”å…§å®¹çµæŸ ===\n"
        
        # Profile è™•ç†ï¼šæª¢æŸ¥æ˜¯å¦éœ€è¦é™„åŠ 
        if "{{profile}}" in prompt:
            essential_profile = {
                "name": self.profile.get('name', 'default'),
                "fields": self.profile.get('fields', [])
            }
            prompt = prompt.replace("{{profile}}", 
                f"\n\n=== Profile é…ç½® ===\n"
                f"é…ç½®è³‡æ–™ï¼š\n{json.dumps(essential_profile, ensure_ascii=False, indent=2)}\n"
                f"=== Profile é…ç½®çµæŸ ===\n")
        else:
            # æª¢æŸ¥æ˜¯å¦éœ€è¦é™„åŠ Profile
            if self._should_attach_profile(prompt_template):
                essential_profile = {
                    "name": self.profile.get('name', 'default'),
                    "fields": self.profile.get('fields', [])
                }
                prompt += f"\n\n=== Profile é…ç½® ===\né…ç½®è³‡æ–™ï¼š\n{json.dumps(essential_profile, ensure_ascii=False, indent=2)}\n=== Profile é…ç½®çµæŸ ===\n"
        
        return prompt
    
    def _should_use_chunking(self, ai_provider: str, content_length: int) -> bool:
        """åˆ¤æ–·æ˜¯å¦éœ€è¦ä½¿ç”¨åˆ†æ®µè™•ç†ï¼ˆåŸºæ–¼ä¸åŒAIæ¨¡å‹çš„ç‰¹æ€§ï¼‰"""
        try:
            # æ ¹æ“šä¸åŒAIæ¨¡å‹çš„å¯¦éš›ä¸Šä¸‹æ–‡é™åˆ¶è¨­ç½®é–¾å€¼ï¼ˆè®“AIæ¨¡å‹è™•ç†ï¼Œä¸è¦æ”¬éä¾†ï¼‰
            chunking_thresholds = {
                # OpenAI GPT ç³»åˆ— - æ ¹æ“šå¯¦éš›ä¸Šä¸‹æ–‡é™åˆ¶
                'openai': 96000,  # 128K tokens â‰ˆ 96K å­—ç¬¦
                'gpt-4': 96000,  # GPT-4: 128K tokens
                'gpt-4o': 96000,  # GPT-4o: 128K tokens
                'gpt-4o-mini': 96000,  # GPT-4o-mini: 128K tokens
                'gpt-4-turbo': 96000,  # GPT-4 Turbo: 128K tokens
                'gpt-3.5-turbo': 12000,  # GPT-3.5: 16K tokens â‰ˆ 12K å­—ç¬¦
                
                # Anthropic Claude ç³»åˆ— - æ ¹æ“šå¯¦éš›ä¸Šä¸‹æ–‡é™åˆ¶ (200K tokens â‰ˆ 150K å­—ç¬¦)
                'anthropic': 150000,
                'claude': 150000,
                'claude-3': 150000,
                'claude-3-sonnet': 150000,
                'claude-3-haiku': 150000,
                'claude-3-opus': 150000,
                'claude-3-5-sonnet-20241022': 150000,  # Claude 3.5 Sonnet: 200K tokens
                'claude-3-5-haiku-20241022': 150000,  # Claude 3.5 Haiku: 200K tokens
                'claude-3-opus-20240229': 150000,  # Claude 3 Opus: 200K tokens
                
                # Google Gemini ç³»åˆ— - æ ¹æ“šå¯¦éš›ä¸Šä¸‹æ–‡é™åˆ¶
                'google': 1500000,
                'gemini': 1500000,
                'gemini-pro': 23000,  # Gemini Pro: 30K tokens â‰ˆ 23K å­—ç¬¦
                'gemini-2.5-pro': 1500000,  # Gemini 2.5 Pro: 2M tokens
                'gemini-2.5-flash': 750000,  # Gemini 2.5 Flash: 1M tokens â‰ˆ 750K å­—ç¬¦
                'gemini-2.0-flash': 1000000,  # Gemini 2.0 Flash: 1M tokens
                'gemini-2.0-flash-lite': 1000000,  # Gemini 2.0 Flash Lite: 1M tokens
                
                # xAI Grok ç³»åˆ— - æ ¹æ“šå¯¦éš›ä¸Šä¸‹æ–‡é™åˆ¶ (128K tokens â‰ˆ 96K å­—ç¬¦)
                'xai': 96000,
                'grok': 96000,
                'grok-2': 96000,  # Grok-2: 128K tokens
                'grok-beta': 96000,  # Grok Beta: 128K tokens
                
                # Microsoft Copilot ç³»åˆ— - æ ¹æ“šå¯¦éš›ä¸Šä¸‹æ–‡é™åˆ¶ (128K tokens â‰ˆ 96K å­—ç¬¦)
                'microsoft': 96000,
                'copilot': 96000,
                'copilot-gpt-4': 96000,  # Copilot GPT-4: 128K tokens
                'copilot-gpt-4-turbo': 96000,  # Copilot GPT-4 Turbo: 128K tokens
                
                # é è¨­é–¾å€¼
                'default': 150000
            }
            
            # ç²å–å°æ‡‰çš„é–¾å€¼
            provider_key = ai_provider.lower() if ai_provider else 'default'
            threshold = chunking_thresholds.get(provider_key, chunking_thresholds['default'])
            
            # æª¢æŸ¥æ˜¯å¦éœ€è¦åˆ†æ®µè™•ç†
            should_chunk = content_length > threshold
            
            if should_chunk:
                logger.info(f"AIæ¨¡å‹ {ai_provider} çš„å…§å®¹é•·åº¦ {content_length} è¶…éé–¾å€¼ {threshold}ï¼Œå»ºè­°ä½¿ç”¨åˆ†æ®µè™•ç†")
            else:
                logger.debug(f"AIæ¨¡å‹ {ai_provider} çš„å…§å®¹é•·åº¦ {content_length} åœ¨é–¾å€¼ {threshold} å…§ï¼Œä½¿ç”¨å–®æ¬¡è™•ç†")
            
            return should_chunk
            
        except Exception as e:
            logger.warning(f"åˆ¤æ–·åˆ†æ®µè™•ç†ç­–ç•¥å¤±æ•—: {e}")
            # é è¨­ä½¿ç”¨è¼ƒä¿å®ˆçš„é–¾å€¼
            return content_length > 30000
    
    def _simple_merge_pages(self, pages_data: List[Dict[str, Any]], file_type: str) -> str:
        """ç°¡åŒ–çš„é é¢åˆä½µé‚è¼¯"""
        merged = []
        for page_data in pages_data:
            if file_type == '.pdf':
                merged.append(f"=== ç¬¬ {page_data['page_number']} é  ===")
            else:
                merged.append(f"=== æ®µè½ {page_data['page_number']} ===")
            merged.append(page_data['content'])
            merged.append("")  # ç©ºè¡Œåˆ†éš”
        
        return "\n".join(merged)
    
    def _should_attach_profile(self, prompt_template: str) -> bool:
        """è®“ä½¿ç”¨è€…æ±ºå®šæ˜¯å¦é™„åŠ Profile"""
        try:
            from utils.settings_manager import SettingsManager
            settings = SettingsManager().get_all_settings()
            profile_strategy = settings.get('prompt', {}).get('profile_strategy', 'auto')
            
            if profile_strategy == 'never':
                return False
            elif profile_strategy == 'always':
                return True
            elif profile_strategy == 'auto':
                # è‡ªå‹•åˆ¤æ–·ï¼šæª¢æŸ¥æç¤ºè©æ˜¯å¦å·²åŒ…å«Profileè³‡è¨Š
                profile_indicators = ['profile', 'é…ç½®', 'fields', 'æ¬„ä½', 'çµæ§‹', 'schema']
                # ä½¿ç”¨æ›´ç²¾ç¢ºçš„åŒ¹é…ï¼Œé¿å…éƒ¨åˆ†åŒ¹é…
                has_profile_info = any(f' {indicator}' in f' {prompt_template.lower()} ' for indicator in profile_indicators)
                return not has_profile_info
            else:
                return True  # é è¨­é™„åŠ 
        except Exception as e:
            logger.warning(f"æª¢æŸ¥Profileç­–ç•¥å¤±æ•—: {e}")
            return True  # é è¨­é™„åŠ 
    
    def _build_extraction_prompt_with_content(self, content: str, file_path: Path = None) -> str:
        """æ§‹å»ºæå–æç¤ºè©ï¼ˆä½¿ç”¨å…§å®¹æ›¿æ›ç­–ç•¥ï¼‰"""
        # åš´æ ¼ä½¿ç”¨ä½¿ç”¨è€…æŒ‡å®šçš„æç¤ºè©
        if not hasattr(self, 'user_prompt') or not self.user_prompt or self.user_prompt.strip() == '':
            raise ValueError("å¿…é ˆæä¾›ä½¿ç”¨è€…æç¤ºè©ã€‚ç³»çµ±ä¸æœƒä½¿ç”¨é è¨­æç¤ºè©ã€‚")
        
        prompt_template = self.user_prompt
        
        # è™•ç†ä½”ä½ç¬¦ï¼šæ›¿æ›ç‚ºå¯¦éš›å…§å®¹ï¼Œä½¿ç”¨æ˜ç¢ºçš„åˆ†éš”æ¨™è¨˜
        if "{{content}}" in prompt_template:
            # æ˜ç¢ºå‘Šè¨´AIé€™ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯æ–‡æª”å…§å®¹
            prompt = prompt_template.replace("{{content}}", f"\n\nã€é‡è¦ï¼šä»¥ä¸‹ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯è¦è™•ç†çš„æ–‡æª”å…§å®¹ã€‘\n{content}\n\nã€æ–‡æª”å…§å®¹çµæŸï¼Œè«‹é–‹å§‹æå–è³‡æ–™ã€‘")
            
            if "{{profile}}" in prompt:
                # æ˜ç¢ºå‘Šè¨´AIé€™ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯Profileé…ç½®
                prompt = prompt.replace("{{profile}}", f"\n\nã€é‡è¦ï¼šä»¥ä¸‹ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯Profileé…ç½®ã€‘\n{json.dumps(self.profile, ensure_ascii=False, indent=2)}\n\nã€Profileé…ç½®çµæŸã€‘")
            
            # åœ¨æœ€å¾Œé¢é™„åŠ å®Œæ•´å…§å®¹
            prompt += f"\n\nã€é‡è¦ï¼šä»¥ä¸‹ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯Profileé…ç½®ã€‘\n{json.dumps(self.profile, ensure_ascii=False, indent=2)}\n\nã€Profileé…ç½®çµæŸã€‘"
            prompt += f"\n\nã€é‡è¦ï¼šä»¥ä¸‹ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯è¦è™•ç†çš„æ–‡æª”å…§å®¹ã€‘\n{content}\n\nã€æ–‡æª”å…§å®¹çµæŸï¼Œè«‹é–‹å§‹æå–è³‡æ–™ã€‘"
        else:
            # æ²’æœ‰ä½”ä½ç¬¦æ™‚ï¼šåœ¨æç¤ºè©å¾Œé¢é™„åŠ profileå’Œcontent
            prompt = prompt_template
            prompt += f"\n\nã€é‡è¦ï¼šä»¥ä¸‹ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯Profileé…ç½®ã€‘\n{json.dumps(self.profile, ensure_ascii=False, indent=2)}\n\nã€Profileé…ç½®çµæŸã€‘"
            prompt += f"\n\nã€é‡è¦ï¼šä»¥ä¸‹ä¸æ˜¯æç¤ºè©ï¼Œæ˜¯è¦è™•ç†çš„æ–‡æª”å…§å®¹ã€‘\n{content}\n\nã€æ–‡æª”å…§å®¹çµæŸï¼Œè«‹é–‹å§‹æå–è³‡æ–™ã€‘"
        
        return prompt
        
    
    def _merge_chunk_results(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """åˆä½µå¤šå€‹æ®µè½çš„çµæœ"""
        merged = {}
        successful_results = [r for r in results if not r.get('_chunk_error')]
        
        logger.info(f"åˆä½µ {len(successful_results)}/{len(results)} å€‹æˆåŠŸæ®µè½çš„çµæœ")
        
        # åˆä½µæ‰€æœ‰åŸå§‹éŸ¿æ‡‰
        raw_responses = []
        for result in results:
            if isinstance(result, dict) and not result.get('_chunk_error'):
                if '_raw_response' in result and result['_raw_response']:
                    raw_responses.append(result['_raw_response'])
        
        # åˆä½µåŸå§‹éŸ¿æ‡‰
        if raw_responses:
            merged['_raw_response'] = '\n\n--- æ®µè½åˆ†éš”ç·š ---\n\n'.join(raw_responses)
            logger.info(f"åˆä½µäº† {len(raw_responses)} å€‹åŸå§‹éŸ¿æ‡‰")
        
        # æ”¶é›†æ‰€æœ‰æˆåˆ†è¡¨ï¼Œç”¨æ–¼å»é‡
        all_ingredients = []
        
        for result in results:
            if isinstance(result, dict) and not result.get('_chunk_error'):
                for key, value in result.items():
                    if key.startswith('_') and key != '_raw_response':  # è·³éå…§éƒ¨æ¬„ä½ï¼Œä½†ä¿ç•™ _raw_response
                        continue
                    
                    # ç‰¹æ®Šè™•ç†æˆåˆ†è¡¨
                    if key == 'æˆåˆ†è¡¨' and isinstance(value, list):
                        all_ingredients.extend(value)
                        continue
                        
                    if key not in merged:
                        merged[key] = value
                    elif isinstance(value, list) and isinstance(merged[key], list):
                        # å°æ–¼ä¸€èˆ¬åˆ—è¡¨ï¼Œåˆä½µ
                        merged[key].extend(value)
                    elif isinstance(value, dict) and isinstance(merged[key], dict):
                        # åˆä½µå­—å…¸
                        merged[key].update(value)
                    elif value and not merged[key]:
                        # å¦‚æœåŸå€¼ç‚ºç©ºï¼Œä½¿ç”¨æ–°å€¼
                        merged[key] = value
        
        # å»é‡æˆåˆ†è¡¨
        if all_ingredients:
            unique_ingredients = self._deduplicate_ingredients(all_ingredients)
            merged['æˆåˆ†è¡¨'] = unique_ingredients
            logger.info(f"æˆåˆ†è¡¨å»é‡ï¼šåŸå§‹ {len(all_ingredients)} å€‹ï¼Œå»é‡å¾Œ {len(unique_ingredients)} å€‹")
        
        # æ·»åŠ è™•ç†çµ±è¨ˆè³‡è¨Š
        merged['_processing_stats'] = {
            'total_chunks': len(results),
            'successful_chunks': len(successful_results),
            'failed_chunks': len(results) - len(successful_results)
        }
        
        return merged
    
    def _deduplicate_ingredients(self, ingredients: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """å»é‡æˆåˆ†è¡¨ï¼ŒåŸºæ–¼INCIåç¨±å’ŒCASè™Ÿç¢¼"""
        seen = set()
        unique_ingredients = []
        
        for ingredient in ingredients:
            if not isinstance(ingredient, dict):
                continue
                
            # å‰µå»ºå”¯ä¸€æ¨™è­˜ç¬¦
            inci_name = ingredient.get('INCIåç¨±', '').strip().upper()
            cas_number = ingredient.get('CASè™Ÿç¢¼', '').strip()
            
            # ä½¿ç”¨INCIåç¨±ä½œç‚ºä¸»è¦æ¨™è­˜ç¬¦ï¼ŒCASè™Ÿç¢¼ä½œç‚ºè¼”åŠ©
            identifier = f"{inci_name}|{cas_number}"
            
            if identifier not in seen:
                seen.add(identifier)
                unique_ingredients.append(ingredient)
            else:
                logger.debug(f"è·³éé‡è¤‡æˆåˆ†: {inci_name} ({cas_number})")
        
        return unique_ingredients
    
    
    def _load_prompt_template(self, template_name: str) -> str:
        """è¼‰å…¥æç¤ºè©æ¨¡æ¿"""
        # å°æ–¼æ‰“åŒ…ç‰ˆæœ¬ï¼Œä½¿ç”¨å·¥ä½œç©ºé–“çš„ prompts ç›®éŒ„
        if getattr(sys, 'frozen', False):
            try:
                from utils.desktop_manager import DesktopManager
                dm = DesktopManager()
                prompts_dir = dm.workspace_dir / "prompts"
            except:
                prompts_dir = Path("prompts")
        else:
            prompts_dir = Path("prompts")
        
        prompt_file = prompts_dir / f"{template_name}.md"
        try:
            with open(prompt_file, 'r', encoding='utf-8') as f:
                return f.read()
        except FileNotFoundError:
            logger.error(f"æç¤ºè©æ¨¡æ¿ä¸å­˜åœ¨: {prompt_file}")
            raise ValueError(f"å¿…é ˆæä¾›æœ‰æ•ˆçš„æç¤ºè©æ¨¡æ¿: {template_name}ã€‚ç³»çµ±ä¸æœƒä½¿ç”¨é è¨­æç¤ºè©ã€‚")
    
    def _get_default_prompt(self) -> str:
        """ç²å–é è¨­æç¤ºè©"""
        return """
è«‹å¾ä»¥ä¸‹æ–‡æª”å…§å®¹ä¸­æå–çµæ§‹åŒ–è³‡æ–™ï¼Œä¸¦ä»¥JSONæ ¼å¼è¼¸å‡ºï¼š

æ–‡æª”å…§å®¹ï¼š
{{content}}

è«‹æŒ‰ç…§ä»¥ä¸‹æ ¼å¼è¼¸å‡ºï¼š
{
  "åŸºæœ¬è³‡è¨Š": {
    "æ¨™é¡Œ": "æ–‡æª”æ¨™é¡Œ",
    "é¡å‹": "æ–‡æª”é¡å‹"
  },
  "ä¸»è¦å…§å®¹": {
    "æ‘˜è¦": "å…§å®¹æ‘˜è¦",
    "é—œéµè³‡è¨Š": "é‡è¦è³‡è¨Š"
  }
}
"""
    
    def _is_document_content_response(self, response: str) -> bool:
        """æª¢æŸ¥AIæ˜¯å¦å¿½ç•¥äº†æç¤ºï¼Œç›´æ¥å›æ‡‰äº†æ–‡æª”å…§å®¹"""
        # æª¢æŸ¥æ˜¯å¦åŒ…å«æ–‡æª”æ ¼å¼çš„é—œéµå­—
        document_indicators = [
            # PDFç›¸é—œ
            'PDF', 'Portable Document Format', 'Adobe', 'Acrobat',
            # Wordç›¸é—œ
            'Microsoft Word', 'DOCX', 'DOC', 'Office', 'Word Document',
            # æ–‡æª”çµæ§‹ç›¸é—œ
            'Page', 'Header', 'Footer', 'Table of Contents', 'Index',
            # æ ¼å¼ç›¸é—œ
            'Font', 'Size', 'Bold', 'Italic', 'Underline', 'Color',
            # æ–‡æª”å…§å®¹ç›¸é—œ
            'Document Properties', 'Metadata', 'Author', 'Title', 'Subject'
        ]
        
        # æª¢æŸ¥æ˜¯å¦åŒ…å«æ–‡æª”æ ¼å¼çš„æ¨™è¨˜
        if any(indicator in response for indicator in document_indicators):
            return True
            
        # æª¢æŸ¥æ˜¯å¦å›æ‡‰äº†åŸå§‹æ–‡æª”å…§å®¹è€Œéæå–çš„çµæ§‹åŒ–è³‡æ–™
        if len(response) > 5000 and not any(marker in response for marker in ['{', '}', 'json', '```']):
            return True
            
        return False
    
    def _is_file_output_response(self, response: str) -> bool:
        """æª¢æŸ¥AIæ˜¯å¦ç›´æ¥è¼¸å‡ºäº†æª”æ¡ˆï¼ˆPDF/Wordç­‰ï¼‰"""
        # æª¢æŸ¥æ˜¯å¦åŒ…å«æª”æ¡ˆè¼¸å‡ºçš„æŒ‡ç¤º
        file_output_indicators = [
            # ç›´æ¥æª”æ¡ˆè¼¸å‡ºæŒ‡ç¤º
            'ç›´æ¥è¼¸å‡ºPDF', 'ç›´æ¥è¼¸å‡ºWord', 'ç”ŸæˆPDFæª”æ¡ˆ', 'ç”ŸæˆWordæª”æ¡ˆ',
            'è¼¸å‡ºPDFæ ¼å¼', 'è¼¸å‡ºWordæ ¼å¼', 'PDFæª”æ¡ˆ', 'Wordæª”æ¡ˆ',
            'ä¸‹è¼‰PDF', 'ä¸‹è¼‰Word', 'PDFä¸‹è¼‰', 'Wordä¸‹è¼‰', 'ä¸‹è¼‰', 'download',
            # æª”æ¡ˆæ ¼å¼æŒ‡ç¤º
            'application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword', 'Content-Type: application/pdf', 'Content-Type: application/msword',
            # äºŒé€²åˆ¶æ•¸æ“šæŒ‡ç¤º
            'base64', 'binary', 'binary data', 'æª”æ¡ˆæ•¸æ“š', 'æª”æ¡ˆå…§å®¹',
            # æª”æ¡ˆä¸‹è¼‰æŒ‡ç¤º
            'attachment', 'æª”æ¡ˆä¸‹è¼‰', 'ä¸‹è¼‰æª”æ¡ˆ'
        ]
        
        # æª¢æŸ¥æ˜¯å¦åŒ…å«æª”æ¡ˆè¼¸å‡ºçš„é—œéµå­—
        if any(indicator in response for indicator in file_output_indicators):
            return True
            
        # æª¢æŸ¥æ˜¯å¦åŒ…å«base64ç·¨ç¢¼çš„æª”æ¡ˆæ•¸æ“š
        if 'base64' in response and len(response) > 1000:
            return True
            
        # æª¢æŸ¥æ˜¯å¦åŒ…å«æª”æ¡ˆæ ¼å¼çš„MIMEé¡å‹
        if any(mime_type in response for mime_type in [
            'application/pdf', 'application/msword', 'application/vnd.openxmlformats'
        ]):
            return True
            
        return False
    
    def _extract_file_from_response(self, response: str) -> Optional[Dict[str, Any]]:
        """å¾AIå›æ‡‰ä¸­æå–æª”æ¡ˆ"""
        try:
            import base64
            import re
            
            # æª¢æŸ¥æ˜¯å¦åŒ…å«base64ç·¨ç¢¼çš„æª”æ¡ˆ
            base64_match = re.search(r'base64[:\s]*([A-Za-z0-9+/=]+)', response, re.IGNORECASE)
            if base64_match:
                base64_data = base64_match.group(1)
                try:
                    file_data = base64.b64decode(base64_data)
                    
                    # æª¢æ¸¬æª”æ¡ˆé¡å‹
                    file_type = self._detect_file_type(file_data)
                    
                    return {
                        'file_data': file_data,
                        'file_type': file_type,
                        'file_size': len(file_data),
                        'extraction_method': 'base64'
                    }
                except Exception as e:
                    logger.warning(f"Base64è§£ç¢¼å¤±æ•—: {e}")
            
            # æª¢æŸ¥æ˜¯å¦åŒ…å«æª”æ¡ˆä¸‹è¼‰é€£çµ
            download_match = re.search(r'(?:ä¸‹è¼‰|download)[:\s]+([^\s]+\.(?:pdf|docx?|doc))', response, re.IGNORECASE)
            if download_match:
                file_path = download_match.group(1)
                return {
                    'file_path': file_path,
                    'file_type': self._get_file_type_from_extension(file_path),
                    'extraction_method': 'download_link'
                }
            
            return None
            
        except Exception as e:
            logger.error(f"æª”æ¡ˆæå–å¤±æ•—: {e}")
            return None
    
    def _detect_file_type(self, file_data: bytes) -> str:
        """æª¢æ¸¬æª”æ¡ˆé¡å‹"""
        # PDFæª”æ¡ˆæ¨™è¨˜
        if file_data.startswith(b'%PDF'):
            return 'pdf'
        
        # Wordæª”æ¡ˆæ¨™è¨˜
        if file_data.startswith(b'PK') and b'word/' in file_data[:1000]:
            return 'docx'
        
        # èˆŠç‰ˆWordæª”æ¡ˆæ¨™è¨˜
        if file_data.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
            return 'doc'
        
        return 'unknown'
    
    def _get_file_type_from_extension(self, file_path: str) -> str:
        """å¾æª”æ¡ˆå‰¯æª”åç²å–æª”æ¡ˆé¡å‹"""
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            return 'pdf'
        elif ext in ['docx', 'doc']:
            return 'docx' if ext == 'docx' else 'doc'
        return 'unknown'
    
    def _parse_ai_response(self, response: str) -> Dict[str, Any]:
        """è§£æAIå›æ‡‰"""
        try:
            import re
            
            # è¨˜éŒ„åŸå§‹å›æ‡‰ä»¥ä¾¿èª¿è©¦
            logger.info(f"AIåŸå§‹å›æ‡‰é•·åº¦: {len(response)} å­—ç¬¦")
            logger.info(f"AIåŸå§‹å›æ‡‰å‰500å­—ç¬¦: {response[:500]}...")
            
            # æª¢æŸ¥AIæ˜¯å¦ç›´æ¥è¼¸å‡ºäº†æª”æ¡ˆ
            if self._is_file_output_response(response):
                logger.info("ğŸ“ AIç›´æ¥è¼¸å‡ºäº†æª”æ¡ˆï¼Œå˜—è©¦æå–æª”æ¡ˆ")
                file_info = self._extract_file_from_response(response)
                if file_info:
                    logger.info(f"âœ… æˆåŠŸæå–æª”æ¡ˆï¼Œé¡å‹: {file_info['file_type']}, å¤§å°: {file_info.get('file_size', 'unknown')}")
                    return {
                        '_file_output': True,
                        '_file_info': file_info,
                        '_raw_response': response,
                        '_extraction_method': 'file_output'
                    }
                else:
                    logger.warning("âš ï¸ æª¢æ¸¬åˆ°æª”æ¡ˆè¼¸å‡ºæŒ‡ç¤ºä½†ç„¡æ³•æå–æª”æ¡ˆ")
            
            # æª¢æŸ¥AIæ˜¯å¦å¿½ç•¥äº†æç¤ºï¼Œç›´æ¥å›æ‡‰äº†æ–‡æª”å…§å®¹
            if self._is_document_content_response(response):
                logger.warning("âš ï¸ AIå¿½ç•¥äº†æç¤ºï¼Œç›´æ¥å›æ‡‰äº†æ–‡æª”å…§å®¹")
                raise Exception("AIå¿½ç•¥äº†æç¤ºï¼Œç›´æ¥å›æ‡‰äº†æ–‡æª”å…§å®¹è€ŒéJSONæ ¼å¼ã€‚è«‹æª¢æŸ¥æç¤ºè©æˆ–å˜—è©¦å…¶ä»–AIæ¨¡å‹ã€‚")
            
            # é¦–å…ˆå˜—è©¦ç›´æ¥è§£ææ•´å€‹å›æ‡‰
            try:
                parsed_data = json.loads(response)
                logger.info(f"âœ… ç›´æ¥JSONè§£ææˆåŠŸï¼Œæ•¸æ“šçµæ§‹: {list(parsed_data.keys())}")
                return parsed_data
            except json.JSONDecodeError as e:
                logger.warning(f"âŒ ç›´æ¥JSONè§£æå¤±æ•—: {e}")
                logger.info(f"å˜—è©¦å…¶ä»–è§£ææ–¹æ³•...")
            
            # å˜—è©¦æå–ä»£ç¢¼å¡Šä¸­çš„JSON
            # å…ˆç§»é™¤ä»£ç¢¼å¡Šæ¨™è¨˜ï¼Œç„¶å¾Œæå–JSON
            if '```json' in response:
                # ç§»é™¤ ```json å’Œ ``` æ¨™è¨˜
                json_start = response.find('```json') + 7
                json_end = response.rfind('```')
                if json_end > json_start:
                    json_str = response[json_start:json_end].strip()
                    try:
                        parsed_data = json.loads(json_str)
                        logger.info(f"âœ… ä»£ç¢¼å¡ŠJSONè§£ææˆåŠŸï¼Œæ•¸æ“šçµæ§‹: {list(parsed_data.keys())}")
                        return parsed_data
                    except json.JSONDecodeError as e:
                        logger.warning(f"âŒ ä»£ç¢¼å¡ŠJSONè§£æå¤±æ•—: {e}")
                        logger.info(f"ä»£ç¢¼å¡Šå…§å®¹å‰100å­—ç¬¦: {json_str[:100]}...")
            elif '```' in response:
                # è™•ç†æ²’æœ‰jsonæ¨™è¨˜çš„ä»£ç¢¼å¡Š
                json_start = response.find('```') + 3
                json_end = response.rfind('```')
                if json_end > json_start:
                    json_str = response[json_start:json_end].strip()
                    try:
                        parsed_data = json.loads(json_str)
                        logger.info(f"âœ… ä»£ç¢¼å¡ŠJSONè§£ææˆåŠŸï¼Œæ•¸æ“šçµæ§‹: {list(parsed_data.keys())}")
                        return parsed_data
                    except json.JSONDecodeError as e:
                        logger.warning(f"âŒ ä»£ç¢¼å¡ŠJSONè§£æå¤±æ•—: {e}")
                        logger.info(f"ä»£ç¢¼å¡Šå…§å®¹å‰100å­—ç¬¦: {json_str[:100]}...")
            
            # å˜—è©¦æå–JSONå°è±¡ï¼ˆæ›´å¯¬é¬†çš„åŒ¹é…ï¼‰
            # ä½¿ç”¨æ›´å¼·çš„æ–¹æ³•ä¾†åŒ¹é…å¤šå±¤åµŒå¥—çš„JSON
            json_match = re.search(r'\{[^{}]*(?:\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}[^{}]*)*\}', response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                try:
                    parsed_data = json.loads(json_str)
                    logger.info(f"âœ… JSONå°è±¡è§£ææˆåŠŸï¼Œæ•¸æ“šçµæ§‹: {list(parsed_data.keys())}")
                    return parsed_data
                except json.JSONDecodeError as e:
                    logger.warning(f"âŒ JSONå°è±¡è§£æå¤±æ•—: {e}")
                    logger.info(f"JSONå°è±¡å…§å®¹å‰100å­—ç¬¦: {json_str[:100]}...")
            
            # å¦‚æœæ‰€æœ‰JSONè§£æéƒ½å¤±æ•—ï¼Œå˜—è©¦å‰µå»ºä¸€å€‹åŸºæœ¬çš„å›æ‡‰çµæ§‹
            logger.warning("ç„¡æ³•å¾AIå›æ‡‰ä¸­æå–JSONï¼Œå‰µå»ºåŸºæœ¬å›æ‡‰çµæ§‹")
            
            # æª¢æŸ¥æ˜¯å¦ç‚ºGPT-4oçš„ç‰¹æ®Šå›æ‡‰æ ¼å¼ï¼ˆå°‡æç¤ºè©ç•¶ä½œJSONéµåï¼‰
            if "è™•ç†çµæœ" in response and any(keyword in response for keyword in [
                "ä½ æ˜¯ä¸€å€‹å°ˆæ¥­çš„é¦™æ°´ç”¢å“è³‡è¨Šæª”æ¡ˆ", "å°‡AZMANå®˜æ–¹è‹±æ–‡", "INCI", "é—œéµ", "æ‰€æœ‰", "Profile"
            ]):
                logger.warning("âš ï¸ æª¢æ¸¬åˆ°GPT-4oç‰¹æ®Šå›æ‡‰æ ¼å¼ï¼Œå˜—è©¦ä¿®å¾©...")
                logger.warning(f"âš ï¸ GPT-4oå›æ‡‰å…§å®¹: {response[:200]}...")
                
                # é€™è¡¨æ˜GPT-4oå°‡æç¤ºè©å…§å®¹ç•¶ä½œäº†JSONéµå
                # æˆ‘å€‘éœ€è¦å‰µå»ºä¸€å€‹ç©ºçš„å›æ‡‰çµæ§‹ï¼Œè®“ç³»çµ±çŸ¥é“æ²’æœ‰æå–åˆ°å¯¦éš›è³‡æ–™
                return {
                    "è™•ç†çµæœ": {
                        "ç”¢å“åç¨±": "",
                        "ç”¢å“é¡åˆ¥": "",
                        "ç”¢å“åŠ‘å‹": "",
                        "ç”¢å“ç”¨é€”": "",
                        "è£½é€ å•†åç¨±": "",
                        "è£½é€ å•†åœ°å€": "",
                        "è£½é€ å•†è¯çµ¡æ–¹å¼": "",
                        "è¼¸å…¥å•†åç¨±": "",
                        "è¼¸å…¥å•†åœ°å€": "",
                        "è¼¸å…¥å•†é›»è©±": "",
                        "åŸç”¢åœ°": "",
                        "å®¹é‡": "",
                        "è£½é€ æ—¥æœŸ": "",
                        "æœ‰æ•ˆæœŸé™": "",
                        "æ‰¹è™Ÿ": "",
                        "ä½¿ç”¨éƒ¨ä½": "",
                        "ä½¿ç”¨æ–¹å¼": "",
                        "é©ç”¨å°è±¡": "",
                        "æ³¨æ„äº‹é …": "",
                        "å®‰å…¨è©•ä¼°çµæœ": "",
                        "æˆåˆ†å®‰å…¨æ€§": "",
                        "ä½¿ç”¨é™åˆ¶": "",
                        "æˆåˆ†è¡¨": [],
                        "ä¸»è¦æˆåˆ†": "",
                        "_raw_response": response,
                        "_parse_method": "gpt4o_error",
                        "_extraction_status": "failed",
                        "_error_reason": "GPT-4oå°‡æç¤ºè©å…§å®¹ç•¶ä½œJSONéµåï¼Œç„¡æ³•æ­£ç¢ºæå–è³‡æ–™"
                    }
                }
            
            # å˜—è©¦å¾å›æ‡‰ä¸­æå–ä¸€äº›åŸºæœ¬ä¿¡æ¯
            basic_info = {}
            
            # æå–ç”¢å“åç¨±
            product_match = re.search(r'(?:ç”¢å“åç¨±|Product Name)[:ï¼š]\s*(.+)', response, re.IGNORECASE)
            if product_match:
                basic_info['ç”¢å“åç¨±'] = product_match.group(1).strip()
            
            # æå–è£½é€ å•†
            manufacturer_match = re.search(r'(?:è£½é€ æ¥­è€…|Manufacturer)[:ï¼š]\s*(.+)', response, re.IGNORECASE)
            if manufacturer_match:
                basic_info['è£½é€ æ¥­è€…'] = manufacturer_match.group(1).strip()
            
            # å¦‚æœæ‰¾åˆ°ä»»ä½•åŸºæœ¬ä¿¡æ¯ï¼Œè¿”å›å®ƒå€‘
            if basic_info:
                basic_info['_raw_response'] = response
                basic_info['_extraction_method'] = 'text_parsing'
                return basic_info
            
            # ç›´æ¥æ‹‹å‡ºç•°å¸¸ï¼Œä¸ä½¿ç”¨å›é€€æ©Ÿåˆ¶
            raise Exception(f"ç„¡æ³•è§£æAIå›æ‡‰ç‚ºJSONæ ¼å¼ã€‚åŸå§‹å›æ‡‰: {response[:200]}...")
            
        except Exception as e:
            logger.error(f"AIå›æ‡‰è§£æå¤±æ•—: {e}")
            logger.debug(f"AIå›æ‡‰å…§å®¹: {response[:500]}...")
            # ç›´æ¥æ‹‹å‡ºç•°å¸¸ï¼Œä¸ä½¿ç”¨å›é€€æ©Ÿåˆ¶
            raise Exception(f"AIå›æ‡‰è§£æå¤±æ•—: {e}")
    
    def _post_process(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """å¾Œè™•ç†æå–çš„è³‡æ–™"""
        # é€™è£¡å¯ä»¥æ·»åŠ è³‡æ–™æ¸…ç†ã€é©—è­‰ç­‰é‚è¼¯
        return data
    
    def learn_from_corrections(self, original_data: Dict[str, Any], 
                             corrected_data: Dict[str, Any]) -> bool:
        """
        å¾ä¿®æ­£ä¸­å­¸ç¿’
        
        Args:
            original_data: åŸå§‹æå–è³‡æ–™
            corrected_data: ä¿®æ­£å¾Œçš„è³‡æ–™
            
        Returns:
            å­¸ç¿’æ˜¯å¦æˆåŠŸ
        """
        try:
            # åˆ†æå·®ç•°
            differences = self._analyze_differences(original_data, corrected_data)
            
            # æ›´æ–°Profileè¦å‰‡
            self.profile_manager.update_profile_from_learning(
                self.profile, differences
            )
            
            logger.info("å­¸ç¿’å®Œæˆï¼ŒProfileå·²æ›´æ–°")
            return True
            
        except Exception as e:
            logger.error(f"å­¸ç¿’å¤±æ•—: {e}")
            return False
    
    def _analyze_differences(self, original: Dict[str, Any], 
                           corrected: Dict[str, Any]) -> List[Dict[str, Any]]:
        """åˆ†æåŸå§‹è³‡æ–™èˆ‡ä¿®æ­£è³‡æ–™çš„å·®ç•°"""
        differences = []
        
        def compare_dict(orig: Dict, corr: Dict, path: str = ""):
            for key in corr:
                current_path = f"{path}.{key}" if path else key
                
                if key not in orig:
                    differences.append({
                        "type": "added",
                        "path": current_path,
                        "original": None,
                        "corrected": corr[key]
                    })
                elif isinstance(corr[key], dict) and isinstance(orig[key], dict):
                    compare_dict(orig[key], corr[key], current_path)
                elif orig[key] != corr[key]:
                    differences.append({
                        "type": "modified",
                        "path": current_path,
                        "original": orig[key],
                        "corrected": corr[key]
                    })
        
        compare_dict(original, corrected)
        return differences
    
    def _build_page_extraction_prompt(self, page_content: str, page_num: int, page_data: Dict[str, Any], file_path: Path = None) -> str:
        """æ§‹å»ºå–®é æå–æç¤ºè©ï¼ˆçµ±ä¸€ä½¿ç”¨æª”åæ›¿æ›ç­–ç•¥ï¼‰"""
        # çµ±ä¸€ä½¿ç”¨æª”åæ›¿æ›ç­–ç•¥ï¼Œé¿å…ç›´æ¥åµŒå…¥å…§å®¹
        filename = f"page_{page_num}"
        return self._build_extraction_prompt(page_content)
    
    def _identify_page_type(self, page_content: str, page_data: Dict[str, Any]) -> str:
        """è­˜åˆ¥é é¢é¡å‹"""
        page_identification = self.profile.get('page_identification', {})
        
        for page_type, config in page_identification.items():
            keywords = config.get('keywords', [])
            for keyword in keywords:
                if keyword.lower() in page_content.lower():
                    return page_type
        
        return "general"
    
    def _merge_page_results(self, merged_result: Dict[str, Any], page_result: Dict[str, Any], page_num: int) -> Dict[str, Any]:
        """åˆä½µé é¢çµæœ"""
        if not merged_result:
            merged_result = page_result.copy()
            merged_result['_metadata'] = {
                'extraction_method': 'page_by_page',
                'total_pages': 0,
                'processed_pages': []
            }
        
        # æ›´æ–°å…ƒè³‡æ–™
        if '_metadata' not in merged_result:
            merged_result['_metadata'] = {'extraction_method': 'page_by_page', 'total_pages': 0, 'processed_pages': []}
        
        merged_result['_metadata']['processed_pages'].append(page_num)
        merged_result['_metadata']['total_pages'] = len(merged_result['_metadata']['processed_pages'])
        
        # åˆä½µè³‡æ–™æ¬„ä½
        for key, value in page_result.items():
            if key == '_metadata':
                continue
                
            if key not in merged_result:
                merged_result[key] = value
            elif isinstance(value, list) and isinstance(merged_result[key], list):
                # æ™ºèƒ½åˆä½µåˆ—è¡¨ï¼ˆç‰¹åˆ¥è™•ç†æˆåˆ†è¡¨ï¼‰
                if key == 'æˆåˆ†è¡¨':
                    merged_result[key] = self._merge_ingredients(merged_result[key], value)
                else:
                    merged_result[key].extend(value)
            elif isinstance(value, dict) and isinstance(merged_result[key], dict):
                # åˆä½µå­—å…¸
                merged_result[key] = self._merge_dicts(merged_result[key], value)
            elif isinstance(value, str) and isinstance(merged_result[key], str):
                # åˆä½µå­—ä¸²ï¼ˆé¿å…é‡è¤‡ï¼Œå„ªå…ˆä½¿ç”¨éç©ºå€¼ï¼‰
                if value and value.strip() and (not merged_result[key] or not merged_result[key].strip()):
                    merged_result[key] = value
                elif value and value.strip() and value not in merged_result[key]:
                    merged_result[key] += f"\n{value}"
        
        return merged_result
    
    def _merge_ingredients(self, existing_ingredients: List[Dict[str, Any]], new_ingredients: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """æ™ºèƒ½åˆä½µæˆåˆ†è¡¨ï¼Œé¿å…é‡è¤‡"""
        # å‰µå»ºæˆåˆ†ç´¢å¼•ï¼ˆåŸºæ–¼INCIåç¨±å’ŒCASè™Ÿç¢¼ï¼‰
        ingredient_index = {}
        
        # è™•ç†ç¾æœ‰æˆåˆ†
        for ingredient in existing_ingredients:
            key = self._get_ingredient_key(ingredient)
            if key:
                ingredient_index[key] = ingredient
        
        # è™•ç†æ–°æˆåˆ†
        for ingredient in new_ingredients:
            key = self._get_ingredient_key(ingredient)
            if key:
                if key in ingredient_index:
                    # åˆä½µç¾æœ‰æˆåˆ†ï¼ˆå„ªå…ˆä½¿ç”¨æ›´å®Œæ•´çš„è³‡æ–™ï¼‰
                    ingredient_index[key] = self._merge_ingredient_data(ingredient_index[key], ingredient)
                else:
                    # æ·»åŠ æ–°æˆåˆ†
                    ingredient_index[key] = ingredient
        
        # è¿”å›å»é‡å¾Œçš„æˆåˆ†åˆ—è¡¨
        return list(ingredient_index.values())
    
    def _get_ingredient_key(self, ingredient: Dict[str, Any]) -> str:
        """ç²å–æˆåˆ†çš„å”¯ä¸€éµ"""
        inci_name = ingredient.get('INCIåç¨±', '').strip().upper()
        cas_number = ingredient.get('CASè™Ÿç¢¼', '').strip()
        
        if inci_name:
            return f"{inci_name}_{cas_number}" if cas_number else inci_name
        return ""
    
    def _merge_ingredient_data(self, existing: Dict[str, Any], new: Dict[str, Any]) -> Dict[str, Any]:
        """åˆä½µæˆåˆ†è³‡æ–™ï¼Œå„ªå…ˆä½¿ç”¨éç©ºå€¼"""
        merged = existing.copy()
        
        for key, value in new.items():
            if value and str(value).strip() and (not existing.get(key) or not str(existing.get(key)).strip()):
                merged[key] = value
            elif key == 'åŠŸèƒ½' and value and str(value).strip():
                # åˆä½µåŠŸèƒ½æè¿°
                existing_func = existing.get(key, '')
                if existing_func and str(existing_func).strip():
                    if str(value).strip() not in str(existing_func):
                        merged[key] = f"{existing_func}, {value}"
                else:
                    merged[key] = value
        
        return merged
    
    def _merge_dicts(self, existing: Dict[str, Any], new: Dict[str, Any]) -> Dict[str, Any]:
        """åˆä½µå­—å…¸ï¼Œå„ªå…ˆä½¿ç”¨éç©ºå€¼"""
        merged = existing.copy()
        
        for key, value in new.items():
            if value and str(value).strip() and (not existing.get(key) or not str(existing.get(key)).strip()):
                merged[key] = value
            elif isinstance(value, str) and isinstance(existing.get(key), str):
                # åˆä½µå­—ä¸²
                existing_val = existing.get(key, '')
                if value.strip() and value.strip() not in existing_val:
                    merged[key] = f"{existing_val}\n{value}".strip()
        
        return merged
    
    def _select_important_pages(self, pages_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """é¸æ“‡é‡è¦é é¢é€²è¡Œè™•ç†"""
        max_pages = self.profile.get('max_pages', 10)
        
        # ç­–ç•¥1: åŸºæ–¼Profileé…ç½®çš„é é¢è­˜åˆ¥è¦å‰‡
        if self.profile.get('page_identification'):
            return self._select_pages_by_profile_rules(pages_data, max_pages)
        
        # ç­–ç•¥2: åŸºæ–¼å…§å®¹å¯†åº¦çš„æ™ºèƒ½é¸æ“‡
        return self._select_pages_by_content_density(pages_data, max_pages)
    
    def _select_pages_by_profile_rules(self, pages_data: List[Dict[str, Any]], max_pages: int) -> List[Dict[str, Any]]:
        """åŸºæ–¼Profileè¦å‰‡é¸æ“‡é é¢"""
        page_identification = self.profile.get('page_identification', {})
        selected_pages = []
        
        # æŒ‰å„ªå…ˆç´šè™•ç†é é¢é¡å‹
        for page_type, config in sorted(page_identification.items(), key=lambda x: x[1].get('priority', 999)):
            keywords = config.get('keywords', [])
            
            for page_data in pages_data:
                if len(selected_pages) >= max_pages:
                    break
                    
                content = page_data.get('content', '').lower()
                if any(keyword.lower() in content for keyword in keywords):
                    if page_data not in selected_pages:
                        selected_pages.append(page_data)
        
        # å¦‚æœé¸ä¸­çš„é é¢ä¸è¶³ï¼Œè£œå……å‰å¹¾é 
        if len(selected_pages) < max_pages:
            for page_data in pages_data:
                if len(selected_pages) >= max_pages:
                    break
                if page_data not in selected_pages:
                    selected_pages.append(page_data)
        
        return selected_pages[:max_pages]
    
    def _select_pages_by_content_density(self, pages_data: List[Dict[str, Any]], max_pages: int) -> List[Dict[str, Any]]:
        """åŸºæ–¼å…§å®¹å¯†åº¦é¸æ“‡é é¢"""
        # è¨ˆç®—æ¯é çš„å…§å®¹å¯†åº¦
        page_densities = []
        for page_data in pages_data:
            content = page_data.get('content', '')
            # å…§å®¹å¯†åº¦ = æ–‡å­—é•·åº¦ + è¡¨æ ¼æ•¸é‡ + é—œéµè©å¯†åº¦
            density = len(content)
            
            # æª¢æŸ¥æ˜¯å¦æœ‰è¡¨æ ¼
            if page_data.get('tables'):
                density += len(page_data['tables']) * 100
            
            # æª¢æŸ¥æ˜¯å¦æœ‰çµæ§‹åŒ–å…§å®¹ï¼ˆå¦‚æˆåˆ†è¡¨ï¼‰
            if any(keyword in content.lower() for keyword in ['æˆåˆ†', 'ingredients', 'cas', 'inci']):
                density += 200
            
            page_densities.append((density, page_data))
        
        # æŒ‰å¯†åº¦æ’åº
        page_densities.sort(key=lambda x: x[0], reverse=True)
        
        # é¸æ“‡å¯†åº¦æœ€é«˜çš„é é¢
        selected_pages = [page_data for _, page_data in page_densities[:max_pages]]
        
        return selected_pages
    
    def _calculate_page_importance(self, page_data: Dict[str, Any]) -> int:
        """è¨ˆç®—é é¢é‡è¦æ€§åˆ†æ•¸"""
        content = page_data.get('content', '').lower()
        score = 0
        
        # é—œéµè©è©•åˆ†
        important_keywords = {
            'ç”¢å“åç¨±': 10,
            'product name': 10,
            'æˆåˆ†': 8,
            'ingredients': 8,
            'è£½é€ æ¥­è€…': 7,
            'manufacturer': 7,
            'ä½¿ç”¨æ–¹æ³•': 6,
            'usage': 6,
            'å®‰å…¨': 5,
            'safety': 5,
            'cas': 4,
            'inci': 4,
            'å«é‡': 3,
            'concentration': 3
        }
        
        for keyword, points in important_keywords.items():
            if keyword in content:
                score += points
        
        # é é¢ä½ç½®è©•åˆ†ï¼ˆå‰å¹¾é é€šå¸¸æ›´é‡è¦ï¼‰
        page_num = page_data.get('page_number', 0)
        if page_num <= 3:
            score += 5
        elif page_num <= 10:
            score += 2
        
        # å…§å®¹é•·åº¦è©•åˆ†ï¼ˆå…§å®¹è¶Šå¤šå¯èƒ½è¶Šé‡è¦ï¼‰
        content_length = len(content)
        if content_length > 500:
            score += 3
        elif content_length > 200:
            score += 1
        
        return score
    
    def _estimate_processing_cost(self, pages_data: List[Dict[str, Any]]) -> float:
        """ä¼°ç®—è™•ç†æˆæœ¬ï¼ˆèˆ‡å‰ç«¯è¨ˆç®—é‚è¼¯ä¸€è‡´ï¼‰"""
        try:
            # ç²å–ç•¶å‰AIæ¨¡å‹çš„å®šåƒ¹è³‡è¨Š
            ai_provider = getattr(self, 'ai_provider', 'openai')
            ai_model = getattr(self, 'ai_model', 'gpt-4o')
            
            # æ ¹æ“šAIæ¨¡å‹ç²å–å®šåƒ¹
            pricing_info = self._get_model_pricing(ai_provider, ai_model)
            input_cost_per_1k = pricing_info.get('input_per_1k', 0.03)
            output_cost_per_1k = pricing_info.get('output_per_1k', 0.06)
            
            total_input_tokens = 0
            total_output_tokens = 0
            
            for page_data in pages_data:
                content = page_data.get('content', '')
                if not content:
                    continue
                
                # æ›´æº–ç¢ºçš„tokenä¼°ç®—ï¼šä¸­æ–‡å­—ç¬¦ç´„1.8å€‹tokenï¼Œè‹±æ–‡å­—ç¬¦ç´„0.75å€‹token
                chinese_chars = len([c for c in content if '\u4e00' <= c <= '\u9fff'])
                english_chars = len(content) - chinese_chars
                page_tokens = int(chinese_chars * 1.8 + english_chars * 0.75)
                
                # è¼¸å‡ºtokensä¼°ç®—ï¼šåŸºæ–¼å¯¦éš›ç¶“é©—ï¼ŒPIFè½‰æ›è¼¸å‡ºé€šå¸¸æ˜¯è¼¸å…¥çš„30-50%
                output_tokens = int(page_tokens * 0.4)
                
                total_input_tokens += page_tokens
                total_output_tokens += output_tokens
            
            # æ·»åŠ ç³»çµ±æç¤ºè©å’Œæ¨¡æ¿æç¤ºè©çš„ä¼°ç®—
            system_prompt_tokens = 2000  # ç³»çµ±æç¤ºè©ç´„2000 tokens
            template_prompt_tokens = 1000  # æ¨¡æ¿æç¤ºè©ç´„1000 tokens
            total_system_tokens = system_prompt_tokens + template_prompt_tokens
            
            # è¨ˆç®—ç¸½æˆæœ¬
            total_input_tokens_with_system = total_input_tokens + total_system_tokens
            input_cost = (total_input_tokens_with_system / 1000) * input_cost_per_1k
            output_cost = (total_output_tokens / 1000) * output_cost_per_1k
            
            total_cost = input_cost + output_cost
            
            logger.debug(f"æˆæœ¬ä¼°ç®—è©³æƒ…: è¼¸å…¥tokens={total_input_tokens_with_system}, è¼¸å‡ºtokens={total_output_tokens}, "
                        f"è¼¸å…¥æˆæœ¬=${input_cost:.4f}, è¼¸å‡ºæˆæœ¬=${output_cost:.4f}, ç¸½æˆæœ¬=${total_cost:.4f}")
            
            return total_cost
            
        except Exception as e:
            logger.error(f"æˆæœ¬ä¼°ç®—å¤±æ•—: {e}")
            # ç›´æ¥æ‹‹å‡ºç•°å¸¸ï¼Œä¸ä½¿ç”¨å›é€€æ©Ÿåˆ¶
            raise Exception(f"æˆæœ¬ä¼°ç®—å¤±æ•—ï¼Œè«‹æª¢æŸ¥AIæ¨¡å‹è¨­å®šæˆ–æ–‡æª”å…§å®¹: {e}")
    
    def _get_model_pricing(self, ai_provider: str, ai_model: str) -> Dict[str, float]:
        """ç²å–AIæ¨¡å‹çš„å®šåƒ¹è³‡è¨Š"""
        # æ ¹æ“šä¸åŒAIæ¨¡å‹è¿”å›å®šåƒ¹è³‡è¨Š
        pricing_map = {
            # OpenAI GPT ç³»åˆ—
            'gpt-4o': {'input_per_1k': 0.005, 'output_per_1k': 0.015},
            'gpt-4-turbo': {'input_per_1k': 0.01, 'output_per_1k': 0.03},
            'gpt-4': {'input_per_1k': 0.03, 'output_per_1k': 0.06},
            'gpt-3.5-turbo': {'input_per_1k': 0.0005, 'output_per_1k': 0.0015},
            
            # Claude ç³»åˆ—
            'claude-3-opus': {'input_per_1k': 0.015, 'output_per_1k': 0.075},
            'claude-3-sonnet': {'input_per_1k': 0.003, 'output_per_1k': 0.015},
            'claude-3-haiku': {'input_per_1k': 0.00025, 'output_per_1k': 0.00125},
            
            # Gemini ç³»åˆ—
            'gemini-2.5-pro': {'input_per_1k': 0.00125, 'output_per_1k': 0.01},
            'gemini-2.5-flash': {'input_per_1k': 0.0003, 'output_per_1k': 0.0025},
            'gemini-2.5-flash-lite': {'input_per_1k': 0.0001, 'output_per_1k': 0.0004},
            'gemini-2.0-flash': {'input_per_1k': 0.0001, 'output_per_1k': 0.0004},
            'gemini-2.0-flash-lite': {'input_per_1k': 0.000075, 'output_per_1k': 0.0003},
            'gemini-pro': {'input_per_1k': 0.0005, 'output_per_1k': 0.0015},
            
            # Grok ç³»åˆ—
            'grok-2': {'input_per_1k': 0.0003, 'output_per_1k': 0.001},
            'grok-beta': {'input_per_1k': 0.0002, 'output_per_1k': 0.0015},
            
            # Microsoft Copilot ç³»åˆ—
            'copilot-gpt-4': {'input_per_1k': 0.005, 'output_per_1k': 0.015},
            'copilot-gpt-4-turbo': {'input_per_1k': 0.01, 'output_per_1k': 0.03},
            
            # é è¨­å®šåƒ¹
            'default': {'input_per_1k': 0.03, 'output_per_1k': 0.06}
        }
        
        # å˜—è©¦åŒ¹é…ç²¾ç¢ºæ¨¡å‹åç¨±
        if ai_model in pricing_map:
            return pricing_map[ai_model]
        
        # å˜—è©¦åŒ¹é…æä¾›è€…
        provider_key = f"{ai_provider.lower()}-default"
        if provider_key in pricing_map:
            return pricing_map[provider_key]
        
        # è¿”å›é è¨­å®šåƒ¹
        return pricing_map['default']
    
    def _post_process_results(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """å¾Œè™•ç†çµæœ"""
        post_process_config = self.profile.get('post_process', {})
        
        if post_process_config.get('clean_text', False):
            result = self._clean_text_fields(result)
        
        if post_process_config.get('normalize_format', False):
            result = self._normalize_format(result)
        
        if post_process_config.get('validate_data', False):
            result = self._validate_data(result)
        
        if post_process_config.get('merge_similar_fields', False):
            result = self._merge_similar_fields(result)
        
        return result
    
    def _clean_text_fields(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """æ¸…ç†æ–‡å­—æ¬„ä½"""
        for key, value in data.items():
            if isinstance(value, str):
                # ç§»é™¤å¤šé¤˜ç©ºç™½
                data[key] = ' '.join(value.split())
            elif isinstance(value, dict):
                data[key] = self._clean_text_fields(value)
            elif isinstance(value, list):
                data[key] = [self._clean_text_fields(item) if isinstance(item, dict) else item for item in value]
        
        return data
    
    def _normalize_format(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """æ¨™æº–åŒ–æ ¼å¼"""
        # é€™è£¡å¯ä»¥æ·»åŠ æ ¼å¼æ¨™æº–åŒ–é‚è¼¯
        return data
    
    def _validate_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """é©—è­‰è³‡æ–™"""
        # é€™è£¡å¯ä»¥æ·»åŠ è³‡æ–™é©—è­‰é‚è¼¯
        return data
    
    def _merge_similar_fields(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """åˆä½µç›¸ä¼¼æ¬„ä½"""
        # é€™è£¡å¯ä»¥æ·»åŠ æ¬„ä½åˆä½µé‚è¼¯
        return data


